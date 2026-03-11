#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
نظام الترجمة الشامل عالي الجودة المحسن - النسخة المطورة
Complete High-Quality Translation System using Gemini - Enhanced Version
ضمان ترجمة كاملة لكل المحتوى مع تحسينات شاملة
"""

import os
import json
import time
import logging
import asyncio
import aiohttp
import re
import sqlite3
import contextlib
import hashlib
import traceback
import unicodedata
import argparse
import sys
from logging.handlers import RotatingFileHandler
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any, Tuple
from collections import deque
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import structlog
import pybreaker
from rich.console import Console
from rich.progress import Progress, TextColumn, BarColumn, TimeRemainingColumn, TaskProgressColumn
from rich.logging import RichHandler
from rich.table import Table
from rich.panel import Panel

# تهيئة Rich Console للطباعة المنسقة في الطرفية
console = Console()

# ============= تحسين 1: نظام السجلات المحسن الذكي باستخدام Structlog =============
def setup_comprehensive_logging():
    """إعداد نظام سجلات شامل وذكي مع structlog و rich، متوافق مع الاستعلام وقاعدة البيانات"""
    log_dir = Path("translation_logs")
    log_dir.mkdir(exist_ok=True)
    
    # تنظيف السجلات القديمة: الاحتفاظ بآخر 5 ملفات فقط لكل نوع (تجنب التراكم الذي ذكره المستخدم)
    def clean_old_logs(prefix: str, keep: int = 5):
        logs = sorted(log_dir.glob(f"{prefix}_*.log"), key=os.path.getmtime, reverse=True)
        for old_log in logs[keep:]:
            try:
                old_log.unlink()
            except Exception:
                pass

    clean_old_logs("main_translation")
    clean_old_logs("quality_control")

    # أسماء ملفات السجلات الجديدة بصيغة JSON
    main_log = log_dir / f'main_translation_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
    quality_log = log_dir / f'quality_control_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
    
    # إعداد إعدادات logging القياسية للتوجيه لملفات باستخدام JSON
    shared_processors = [
        structlog.stdlib.add_log_level,
        structlog.stdlib.add_logger_name,
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.StackInfoRenderer(),
        structlog.processors.ExceptionRenderer(),
        structlog.processors.UnicodeDecoder(),
    ]

    structlog.configure(
        processors=shared_processors + [
            structlog.stdlib.PositionalArgumentsFormatter(),
            structlog.stdlib.ProcessorFormatter.wrap_for_formatter,
        ],
        logger_factory=structlog.stdlib.LoggerFactory(),
        wrapper_class=structlog.stdlib.BoundLogger,
        cache_logger_on_first_use=True,
    )

    # إعداد Formatter لملفات السجل (JSON) و Formatter للطرفية (Rich)
    json_formatter = structlog.stdlib.ProcessorFormatter(
        processor=structlog.processors.JSONRenderer(ensure_ascii=False),
        foreign_pre_chain=shared_processors,
    )

    console_formatter = structlog.stdlib.ProcessorFormatter(
        processor=structlog.dev.ConsoleRenderer(colors=True),
        foreign_pre_chain=shared_processors,
    )

    # إعداد logger الرئيسي مع rotation
    main_logger_std = logging.getLogger('main')
    main_logger_std.setLevel(logging.INFO)
    main_logger_std.handlers.clear() # تفريغ لتجنب التكرار
    
    # Rotating file handler للملف الرئيسي بصيغة JSON
    main_handler = RotatingFileHandler(
        main_log,
        maxBytes=10*1024*1024,  # 10MB
        backupCount=5,
        encoding='utf-8'
    )
    main_handler.setFormatter(json_formatter)
    main_logger_std.addHandler(main_handler)

    # Console handler لطباعة السجلات الملونة والمنسقة في الطرفية باستخدام Rich
    rich_handler = RichHandler(console=console, show_time=False, show_path=False, markup=True)
    rich_handler.setFormatter(console_formatter)
    main_logger_std.addHandler(rich_handler)

    # logger منفصل لمراقبة الجودة مع rotation
    quality_logger_std = logging.getLogger('quality_control')
    quality_logger_std.setLevel(logging.INFO)
    quality_logger_std.handlers.clear()
    
    quality_handler = RotatingFileHandler(
        quality_log,
        maxBytes=5*1024*1024,  # 5MB
        backupCount=3,
        encoding='utf-8'
    )
    quality_handler.setFormatter(json_formatter)
    quality_logger_std.addHandler(quality_handler)

    quality_rich_handler = RichHandler(console=console, show_time=False, show_path=False, markup=True)
    quality_rich_handler.setFormatter(console_formatter)
    quality_logger_std.addHandler(quality_rich_handler)

    # استخدام structlog للحصول على واجهة استخدام محسنة (تتيح إرسال kwargs)
    main_logger = structlog.get_logger('main')
    quality_logger = structlog.get_logger('quality_control')
    
    return main_logger, quality_logger

logger, quality_logger = setup_comprehensive_logging()

# ============= تحسين 2: نظام Rate Limiting محسن (Tokens & Requests) =============

# --- تقدير التوكنز الذكي المدرك للغة ---
def _estimate_tokens_smart(text: str) -> int:
    """
    تقدير دقيق للتوكنز مع دعم نصوص مختلطة العربية/الإنجليزية.
    يستخدم tiktoken إذا كان متاحاً (دقة عالية)، وإلا يستخدم
    إحصاءات لغوية محسّنة تُراعي أن النصوص العربية تحتاج 20-40% توكنز أكثر.
    """
    if not text:
        return 0

    # المحاولة الأولى: tiktoken (دقة عالية، يعمل محلياً)
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return max(1, len(enc.encode(text)))
    except (ImportError, Exception):
        pass

    # تحليل لغوي للنص
    total_chars = len(text)
    if total_chars == 0:
        return 0

    # حساب نسبة الأحرف العربية (U+0600–U+06FF)
    arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
    arabic_ratio = arabic_chars / total_chars

    # معادلة التقدير المحسّنة:
    # • نص عربي بحت  → ~2.3 حرف/توكن
    # • نص مختلط     → ~3.0 حرف/توكن
    # • نص إنجليزي   → ~4.0 حرف/توكن
    if arabic_ratio >= 0.50:
        chars_per_token = 2.3
    elif arabic_ratio >= 0.20:
        chars_per_token = 3.0
    else:
        chars_per_token = 4.0

    return max(1, int(total_chars / chars_per_token))


# --- ملف إعدادات Rate Limiter الخارجي ---
_RATE_LIMITER_CONFIG_PATH = Path("rate_limiter_config.json")

_RATE_LIMITER_DEFAULT_CONFIG = {
    "default": {
        "max_rpm": 5,
        "max_tpm": 32000,
        "max_rpd": 25
    },
    "adaptive": {
        "enabled": True,
        "min_rpm_floor": 1,
        "error_threshold_per_hour": 3,
        "reduction_factor": 0.5,
        "recovery_factor": 0.1
    },
    "keys": {}
}


def _load_rate_limiter_config() -> dict:
    """
    تحميل إعدادات Rate Limiter من ملف JSON خارجي.
    إذا لم يوجد الملف، يتم إنشاؤه بالإعدادات الافتراضية تلقائياً.
    البنية:
      {
        "default": { "max_rpm": 5, "max_tpm": 32000, "max_rpd": 25 },
        "adaptive": { "enabled": true, ... },
        "keys": {
          "AIzaSy...": { "max_rpm": 5, "max_tpm": 60000, "max_rpd": 100 }
        }
      }
    """
    if _RATE_LIMITER_CONFIG_PATH.exists():
        try:
            with open(_RATE_LIMITER_CONFIG_PATH, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
            # دمج مع الإعدادات الافتراضية (لدعم الإضافات المستقبلية)
            merged = _RATE_LIMITER_DEFAULT_CONFIG.copy()
            merged["default"].update(loaded.get("default", {}))
            merged["adaptive"].update(loaded.get("adaptive", {}))
            merged["keys"].update(loaded.get("keys", {}))
            return merged
        except Exception:
            pass  # fallback to defaults

    # إنشاء الملف الافتراضي للمرة الأولى
    try:
        with open(_RATE_LIMITER_CONFIG_PATH, 'w', encoding='utf-8') as f:
            json.dump(_RATE_LIMITER_DEFAULT_CONFIG, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    return _RATE_LIMITER_DEFAULT_CONFIG.copy()


class TokenRateLimiter:
    """
    نظام Rate Limiting ذكي ومتكيّف لكل مفتاح API.
    
    المميزات:
      ✅ إعدادات قابلة للتهيئة من ملف خارجي (rate_limiter_config.json)
      ✅ دعم إعدادات مختلفة لكل مفتاح API على حدة (خطط مختلفة)
      ✅ تقدير دقيق للتوكنز يراعي النصوص العربية (دعم tiktoken)
      ✅ نظام تكيّفي: يتعلم من أخطاء 429 ويخفّض RPM تلقائياً
        في الساعات ذات الضغط العالي ثم يعود تدريجياً للحد الطبيعي
    """

    def __init__(self, max_rpm: int = 5, max_tpm: int = 32000, max_rpd: int = 25,
                 key_id: Optional[str] = None):
        """
        Args:
            max_rpm: الحد الأقصى للطلبات في الدقيقة (قابل للتجاوز من الملف الخارجي).
            max_tpm: الحد الأقصى للتوكنز في الدقيقة.
            max_rpd: الحد الأقصى للطلبات في اليوم.
            key_id: مفتاح API للبحث عن إعداداته المخصصة في ملف الإعداد.
        """
        # --- تحميل الإعدادات من الملف الخارجي ---
        config = _load_rate_limiter_config()
        default_cfg = config.get("default", {})
        adaptive_cfg = config.get("adaptive", {})
        key_cfg = config.get("keys", {}).get(key_id, {}) if key_id else {}

        # الأولوية: إعدادات المفتاح المخصصة ← ملف الإعداد ← المعاملات الممررة
        self._base_max_rpm = key_cfg.get("max_rpm", default_cfg.get("max_rpm", max_rpm))
        self._base_max_tpm = key_cfg.get("max_tpm", default_cfg.get("max_tpm", max_tpm))
        self._base_max_rpd = key_cfg.get("max_rpd", default_cfg.get("max_rpd", max_rpd))

        # الحدود الفعّالة الحالية (قد تُقلَّص مؤقتاً من النظام التكيّفي)
        self.max_rpm = self._base_max_rpm
        self.max_tpm = self._base_max_tpm
        self.max_rpd = self._base_max_rpd

        # --- طوابير الطلبات ---
        self.requests: deque = deque()        # (timestamp, tokens)
        self.daily_requests: deque = deque()  # timestamps خلال 24 ساعة

        # --- النظام التكيّفي ---
        self._adaptive_enabled: bool = adaptive_cfg.get("enabled", True)
        self._min_rpm_floor: int = max(1, adaptive_cfg.get("min_rpm_floor", 1))
        self._error_threshold: int = adaptive_cfg.get("error_threshold_per_hour", 3)
        self._reduction_factor: float = adaptive_cfg.get("reduction_factor", 0.5)
        self._recovery_factor: float = adaptive_cfg.get("recovery_factor", 0.1)

        # عداد أخطاء 429 مُجمَّعة بالساعة {hour_of_day: count}
        self._errors_by_hour: Dict[int, int] = {}
        # الساعة التي آخر مرة حدث فيها تعديل تكيّفي
        self._last_adaptive_hour: Optional[int] = None

    # ------------------------------------------------------------------ #
    #  تقدير التوكنز                                                       #
    # ------------------------------------------------------------------ #

    @staticmethod
    def estimate_tokens(text: str) -> int:
        """تقدير دقيق لعدد توكنز النص يعتمد على _estimate_tokens_smart."""
        return _estimate_tokens_smart(text)

    # ------------------------------------------------------------------ #
    #  النظام التكيّفي                                                     #
    # ------------------------------------------------------------------ #

    def record_429_error(self):
        """
        تسجيل خطأ 429 (Rate Limit) من الـ API الفعلي.
        بعد تجاوز _error_threshold في الساعة الحالية، يُخفَّض max_rpm
        بمعامل _reduction_factor حتى لا يقل عن _min_rpm_floor.
        """
        if not self._adaptive_enabled:
            return

        current_hour = datetime.now().hour
        self._errors_by_hour[current_hour] = self._errors_by_hour.get(current_hour, 0) + 1

        if self._errors_by_hour[current_hour] >= self._error_threshold:
            new_rpm = max(self._min_rpm_floor,
                          int(self.max_rpm * (1.0 - self._reduction_factor)))
            if new_rpm < self.max_rpm:
                self.max_rpm = new_rpm
                logger.warning(
                    f"[AdaptiveRateLimiter] ارتفاع أخطاء 429 في الساعة {current_hour}:00 "
                    f"→ تخفيض RPM إلى {self.max_rpm}/{self._base_max_rpm}"
                )

        self._last_adaptive_hour = current_hour

    def _try_recover_rpm(self):
        """
        محاولة استعادة RPM تدريجياً بعد ساعة هادئة (بدون أخطاء 429 جديدة).
        تُستدعى داخلياً عند كل طلب ناجح.
        """
        if not self._adaptive_enabled:
            return
        if self.max_rpm >= self._base_max_rpm:
            return

        current_hour = datetime.now().hour
        hour_errors = self._errors_by_hour.get(current_hour, 0)

        # إذا لم تحدث أخطاء في الساعة الحالية، استعد جزءاً من الحد الأصلي
        if hour_errors == 0:
            recovered = max(1, int(self._base_max_rpm * self._recovery_factor))
            self.max_rpm = min(self._base_max_rpm, self.max_rpm + recovered)

    # ------------------------------------------------------------------ #
    #  الواجهة الرئيسية (تتوافق 100% مع الإصدار السابق)                   #
    # ------------------------------------------------------------------ #

    def _purge_old_entries(self, now: float):
        """إزالة الإدخالات القديمة خارج نافذة الدقيقة والـ 24 ساعة."""
        while self.requests and self.requests[0][0] < now - 60:
            self.requests.popleft()
        while self.daily_requests and self.daily_requests[0] < now - 86400:
            self.daily_requests.popleft()

    def can_make_request(self, estimated_tokens: int = 0) -> bool:
        """هل يمكن إرسال طلب جديد الآن دون تجاوز أي حد؟"""
        now = time.time()
        self._purge_old_entries(now)

        if len(self.requests) >= self.max_rpm:
            return False
        if sum(t for _, t in self.requests) + estimated_tokens > self.max_tpm:
            return False
        if len(self.daily_requests) >= self.max_rpd:
            return False

        return True

    def add_request(self, estimated_tokens: int = 0):
        """تسجيل طلب جديد وتحديث النظام التكيّفي للاسترداد."""
        now = time.time()
        self.requests.append((now, estimated_tokens))
        self.daily_requests.append(now)
        # محاولة استعادة RPM تدريجياً عند كل طلب ناجح
        self._try_recover_rpm()

    def time_until_next_request(self, estimated_tokens: int = 0) -> float:
        """وقت الانتظار (بالثواني) حتى يمكن إرسال طلب جديد."""
        if self.can_make_request(estimated_tokens):
            return 0.0

        now = time.time()
        self._purge_old_entries(now)

        wait_times = []

        # انتظار بسبب الحد اليومي
        if len(self.daily_requests) >= self.max_rpd:
            wait_times.append((self.daily_requests[0] + 86400) - now)

        # انتظار بسبب حد الطلبات في الدقيقة
        if len(self.requests) >= self.max_rpm:
            wait_times.append((self.requests[0][0] + 60) - now)

        # انتظار بسبب حد التوكنز في الدقيقة
        current_tpm = sum(t for _, t in self.requests)
        if current_tpm + estimated_tokens > self.max_tpm:
            tokens_to_free = (current_tpm + estimated_tokens) - self.max_tpm
            freed = 0
            for req_time, tokens in self.requests:
                freed += tokens
                if freed >= tokens_to_free:
                    wait_times.append((req_time + 60) - now)
                    break

        return max(wait_times) if wait_times else 0.0

    def get_status(self) -> Dict[str, Any]:
        """
        حالة Rate Limiter الحالية (مفيد للتشخيص والسجلات).
        """
        now = time.time()
        self._purge_old_entries(now)
        return {
            "effective_rpm": self.max_rpm,
            "base_rpm": self._base_max_rpm,
            "max_tpm": self.max_tpm,
            "max_rpd": self.max_rpd,
            "current_rpm_usage": len(self.requests),
            "current_tpm_usage": sum(t for _, t in self.requests),
            "current_rpd_usage": len(self.daily_requests),
            "adaptive_enabled": self._adaptive_enabled,
            "errors_by_hour": dict(self._errors_by_hour),
        }

# ============= تحسين 3: إحصائيات محسنة للمفاتيح ونظام التنبيه الذكي =============
class KeyStatistics:
    """
    إحصائيات متقدمة وذكية لكل مفتاح API مع:
      ✅ تخزين دائم في SQLite بين الجلسات (لا يصفّر التاريخ عند إغلاق البرنامج)
      ✅ تصنيف دقيق لأنواع الفشل: network_error / rate_limit / invalid_key /
         content_blocked / server_error — كل نوع يحمل وزناً مختلفاً
      ✅ نقاط صحة مبنية على بيانات حقيقية لا عقوبات تقديرية ثابتة
      ✅ نموذج تنبؤ ساعي بسيط يتعلم متى يكون المفتاح في أفضل حالاته
    """

    _DB_PATH: Path = Path("key_statistics.db")
    _db_initialized: bool = False  # يُهيَّأ مرة واحدة على مستوى الكلاس

    # ---- الأنواع القياسية للفشل ----
    FAILURE_TYPES = frozenset([
        'network_error',    # timeout / اتصال مقطوع / استثناء شبكي
        'rate_limit',       # 429 - تجاوز حد الطلبات
        'invalid_key',      # 401/403 - مفتاح منتهٍ أو محظور
        'content_blocked',  # رد غير متوقع أو محجوب من Gemini Safety
        'server_error',     # 5xx - خطأ في سيرفر Gemini
        'general',          # أخطاء لا تندرج تحت فئة محددة
    ])

    # خريطة التوحيد: raw error_type الموجود في الكود → canonical type
    _ERROR_CODE_MAP: Dict[str, str] = {
        'timeout':          'network_error',
        'exception':        'network_error',
        'network_error':    'network_error',
        'rate_limit':       'rate_limit',
        'invalid_key':      'invalid_key',
        'api_error':        'invalid_key',       # 401/403 في make_precision_request
        'content_blocked':  'content_blocked',
        'invalid_response': 'content_blocked',   # رد Gemini غير متوقع
        'server_error':     'server_error',
        'general':          'general',
    }

    # أوزان تأثير كل نوع فشل على نقطة الصحة
    # (مشتقة من طبيعة كل خطأ، ليست تقديرية عشوائية)
    FAILURE_WEIGHTS: Dict[str, float] = {
        'invalid_key':      3.0,   # الأشد خطورة — يعني المفتاح معطَّل نهائياً
        'server_error':     1.5,   # خطير لكن مؤقت
        'network_error':    0.8,   # مؤقت، ناجم عن ظروف الشبكة
        'rate_limit':       0.4,   # متوقع في الاستخدام المكثف
        'content_blocked':  0.2,   # لا علاقة له بصحة المفتاح نفسه
        'general':          1.0,
    }

    def __init__(self, key_id: str = ""):
        self.key_id   = key_id
        # تجزئة للهوية في DB (لا نخزّن المفتاح الحقيقي)
        self.key_hash = hashlib.md5(key_id.encode()).hexdigest()[:16] if key_id else "unknown"

        # ---- عدادات في الذاكرة ----
        self.total_requests:        int   = 0
        self.successful_requests:   int   = 0
        self.failed_requests:       int   = 0
        self.consecutive_failures:  int   = 0
        self.last_error_time:   Optional[datetime] = None
        self.last_success_time: Optional[datetime] = None
        self.average_response_time: float = 0.0
        self.response_times: deque = deque(maxlen=100)

        # ---- عدادات مصنَّفة لأنواع الفشل ----
        self.failure_counts: Dict[str, int] = {ft: 0 for ft in self.FAILURE_TYPES}

        # ---- بيانات ساعية للتنبؤ: {0..23: {'success': N, 'total': N}} ----
        self.hourly_data: Dict[int, Dict[str, int]] = {
            h: {'success': 0, 'total': 0} for h in range(24)
        }

        # ---- تحكم في كثافة الكتابة إلى DB ----
        self._save_counter: int = 0
        self._SAVE_EVERY:   int = 5   # حفظ كل 5 طلبات لتجنب I/O زائد

        # تهيئة DB وتحميل البيانات المحفوظة إن وُجدت
        if key_id:
            self._ensure_db()
            self._load_from_db()

    # ------------------------------------------------------------------ #
    #  إدارة قاعدة البيانات                                               #
    # ------------------------------------------------------------------ #

    @classmethod
    def _ensure_db(cls):
        """إنشاء جدول SQLite إن لم يكن موجوداً (يُنفَّذ مرة واحدة على مستوى الكلاس)."""
        if cls._db_initialized:
            return
        try:
            with sqlite3.connect(cls._DB_PATH) as conn:
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS key_statistics (
                        key_hash            TEXT PRIMARY KEY,
                        total_requests      INTEGER NOT NULL DEFAULT 0,
                        successful_requests INTEGER NOT NULL DEFAULT 0,
                        failed_requests     INTEGER NOT NULL DEFAULT 0,
                        failure_counts      TEXT    NOT NULL DEFAULT '{}',
                        hourly_data         TEXT    NOT NULL DEFAULT '{}',
                        avg_response_time   REAL    NOT NULL DEFAULT 0.0,
                        response_times      TEXT    NOT NULL DEFAULT '[]',
                        last_error_time     TEXT,
                        last_success_time   TEXT,
                        updated_at          TEXT    NOT NULL DEFAULT (datetime('now'))
                    )
                """)
                conn.commit()
            cls._db_initialized = True
            logger.info("[KeyStatistics] SQLite DB initialized successfully")
        except Exception as e:
            logger.warning(f"[KeyStatistics] DB init failed: {e}")

    def _load_from_db(self):
        """تحميل الإحصائيات المحفوظة من SQLite إلى الذاكرة عند بدء الجلسة."""
        try:
            with sqlite3.connect(self._DB_PATH) as conn:
                conn.row_factory = sqlite3.Row
                row = conn.execute(
                    "SELECT * FROM key_statistics WHERE key_hash = ?",
                    (self.key_hash,)
                ).fetchone()

            if not row:
                return  # مفتاح جديد — لا توجد بيانات سابقة

            self.total_requests        = row['total_requests']
            self.successful_requests   = row['successful_requests']
            self.failed_requests       = row['failed_requests']
            self.average_response_time = row['avg_response_time']

            # تحميل عدادات الفشل المصنَّفة
            stored_fc = json.loads(row['failure_counts'] or '{}')
            for ft in self.FAILURE_TYPES:
                self.failure_counts[ft] = stored_fc.get(ft, 0)

            # تحميل البيانات الساعية
            stored_hd = json.loads(row['hourly_data'] or '{}')
            for h in range(24):
                entry = stored_hd.get(str(h), {})
                self.hourly_data[h]['success'] = entry.get('success', 0)
                self.hourly_data[h]['total']   = entry.get('total',   0)

            # تحميل آخر أوقات الاستجابة
            stored_rt = json.loads(row['response_times'] or '[]')
            self.response_times = deque(stored_rt, maxlen=100)
            self._update_average_response_time()

            # تحميل أوقات آخر خطأ/نجاح
            if row['last_error_time']:
                try:
                    self.last_error_time = datetime.fromisoformat(row['last_error_time'])
                except Exception:
                    pass
            if row['last_success_time']:
                try:
                    self.last_success_time = datetime.fromisoformat(row['last_success_time'])
                except Exception:
                    pass

            logger.info(
                f"[KeyStatistics] Loaded history for {self.key_hash}: "
                f"{self.total_requests} requests, "
                f"success_rate={self.get_success_rate():.1f}%"
            )
        except Exception as e:
            logger.warning(f"[KeyStatistics] Load from DB failed ({self.key_hash}): {e}")

    def _save_to_db(self, force: bool = False):
        """
        حفظ الإحصائيات في SQLite.
        يُكتب فعلياً كل _SAVE_EVERY طلبات (أو فوراً إذا force=True)
        لتجنب I/O زائد على كل طلب.
        """
        self._save_counter += 1
        if not force and (self._save_counter % self._SAVE_EVERY != 0):
            return
        try:
            with sqlite3.connect(self._DB_PATH) as conn:
                conn.execute("""
                    INSERT INTO key_statistics
                        (key_hash, total_requests, successful_requests, failed_requests,
                         failure_counts, hourly_data, avg_response_time, response_times,
                         last_error_time, last_success_time, updated_at)
                    VALUES (?,?,?,?,?,?,?,?,?,?,datetime('now'))
                    ON CONFLICT(key_hash) DO UPDATE SET
                        total_requests      = excluded.total_requests,
                        successful_requests = excluded.successful_requests,
                        failed_requests     = excluded.failed_requests,
                        failure_counts      = excluded.failure_counts,
                        hourly_data         = excluded.hourly_data,
                        avg_response_time   = excluded.avg_response_time,
                        response_times      = excluded.response_times,
                        last_error_time     = excluded.last_error_time,
                        last_success_time   = excluded.last_success_time,
                        updated_at          = excluded.updated_at
                """, (
                    self.key_hash,
                    self.total_requests,
                    self.successful_requests,
                    self.failed_requests,
                    json.dumps(self.failure_counts, ensure_ascii=False),
                    json.dumps(self.hourly_data,    ensure_ascii=False),
                    self.average_response_time,
                    json.dumps(list(self.response_times)),
                    self.last_error_time.isoformat()   if self.last_error_time   else None,
                    self.last_success_time.isoformat() if self.last_success_time else None,
                ))
                conn.commit()
        except Exception as e:
            logger.warning(f"[KeyStatistics] Save to DB failed ({self.key_hash}): {e}")

    # ------------------------------------------------------------------ #
    #  تسجيل الأحداث                                                      #
    # ------------------------------------------------------------------ #

    def record_success(self, response_time: float):
        """تسجيل طلب ناجح، تحديث البيانات الساعية، وإعادة تعيين عداد الإخفاقات."""
        self.successful_requests += 1
        self.total_requests      += 1
        self.consecutive_failures = 0
        self.last_success_time    = datetime.now()
        self.response_times.append(response_time)
        self._update_average_response_time()

        # تحديث البيانات الساعية (أساس نموذج التنبؤ)
        hour = self.last_success_time.hour
        self.hourly_data[hour]['success'] += 1
        self.hourly_data[hour]['total']   += 1

        self._save_to_db()

    def record_failure(self, error_type: str = "general") -> bool:
        """
        تسجيل طلب فاشل مع تصنيف دقيق لنوع الفشل.
        يُوحَّد error_type إلى النوع القياسي المناسب تلقائياً.
        يرجع True إذا وصل عدد الإخفاقات المتتالية إلى 3 (للتنبيه في المستدعي).
        """
        # توحيد النوع الخام إلى النوع القياسي
        canonical = self._ERROR_CODE_MAP.get(error_type, 'general')

        self.failed_requests      += 1
        self.total_requests       += 1
        self.last_error_time       = datetime.now()
        self.failure_counts[canonical] = self.failure_counts.get(canonical, 0) + 1

        # content_blocked و rate_limit لا يعكسان خللاً في المفتاح نفسه
        # → لا تزيد عداد الإخفاقات المتتالية (لتجنب التأثير الخاطئ على get_health_score)
        if canonical not in ('content_blocked', 'rate_limit'):
            self.consecutive_failures += 1

        # تحديث البيانات الساعية (فشل → total فقط، بدون success)
        hour = self.last_error_time.hour
        self.hourly_data[hour]['total'] += 1

        self._save_to_db(force=(self.consecutive_failures >= 3))
        return self.consecutive_failures >= 3

    # ------------------------------------------------------------------ #
    #  خصائص التوافق مع الكود القديم (للحفاظ على الهيكلة)               #
    # ------------------------------------------------------------------ #

    @property
    def rate_limit_hits(self) -> int:
        """توافق مع الكود القديم — يقرأ من failure_counts المصنَّفة."""
        return self.failure_counts.get('rate_limit', 0)

    @property
    def server_errors(self) -> int:
        """توافق مع الكود القديم — يقرأ من failure_counts المصنَّفة."""
        return self.failure_counts.get('server_error', 0)

    # ------------------------------------------------------------------ #
    #  الحسابات والتحليل                                                  #
    # ------------------------------------------------------------------ #

    def _update_average_response_time(self):
        """تحديث متوسط وقت الاستجابة من نافذة آخر 100 طلب."""
        if self.response_times:
            self.average_response_time = sum(self.response_times) / len(self.response_times)

    def get_success_rate(self) -> float:
        """حساب معدل النجاح الإجمالي (0-100)."""
        if self.total_requests == 0:
            return 0.0
        return (self.successful_requests / self.total_requests) * 100

    def get_health_score(self) -> float:
        """
        نقاط الصحة (0-100) مبنية على بيانات حقيقية لا عقوبات تقديرية ثابتة.

        المكوّنات:
          base_score     = معدل النجاح الإجمالي
          failure_penalty = مجموع(عدد_فشل_نوع × وزن_النوع) / إجمالي_الطلبات × 20
                           (حد أقصى 40 نقطة — مشتق من نسبة الفشل الحقيقية)
          consec_penalty  = 5 × consecutive_failures (حد أقصى 25)
          recency_bonus   = +5 إذا كان آخر طلب ناجح منذ أقل من 30 دقيقة

        حالات خاصة:
          • مفتاح جديد (0 طلبات)  → 100 (فرصة كاملة)
          • 3+ أخطاء invalid_key   → 0  (المفتاح معطَّل، يُزال من الدوران)
        """
        if self.total_requests == 0:
            return 100.0

        # مفتاح معطَّل نهائياً
        if self.failure_counts.get('invalid_key', 0) >= 3:
            return 0.0

        # 1. النقطة الأساسية
        base_score = self.get_success_rate()

        # 2. عقوبة الفشل المرجَّحة (مبنية على البيانات الحقيقية)
        weighted_failures = sum(
            self.failure_counts.get(ft, 0) * self.FAILURE_WEIGHTS.get(ft, 1.0)
            for ft in self.FAILURE_TYPES
        )
        failure_penalty = min(
            (weighted_failures / max(self.total_requests, 1)) * 20.0,
            40.0  # سقف الخصم
        )

        # 3. عقوبة الإخفاقات المتتالية الحالية (مؤشر آني)
        consec_penalty = min(self.consecutive_failures * 5.0, 25.0)

        # 4. مكافأة النشاط الأخير السليم
        recency_bonus = 0.0
        if (self.last_success_time and
                (self.last_error_time is None or self.last_success_time > self.last_error_time)):
            mins_since_success = (datetime.now() - self.last_success_time).total_seconds() / 60
            if mins_since_success < 30:
                recency_bonus = 5.0

        health = base_score - failure_penalty - consec_penalty + recency_bonus
        return max(0.0, min(100.0, health))

    # ------------------------------------------------------------------ #
    #  نموذج التنبؤ الساعي                                                #
    # ------------------------------------------------------------------ #

    def get_predicted_performance(self) -> float:
        """
        تنبؤ بالأداء المتوقع للمفتاح في الساعة الحالية بناءً على البيانات التاريخية.

        الخوارزمية:
          • إذا توفرت ≥5 طلبات تاريخية لهذه الساعة بالذات:
              predicted = 70% × معدل_النجاح_الساعي + 30% × معدل_النجاح_الإجمالي
            (الوزن الأكبر للساعي لأنه الأدق لهذه اللحظة من اليوم)
          • وإلا (بيانات غير كافية):
              predicted = get_health_score()  (fallback آمن)

        يُستدعى من get_optimal_api_key لاختيار المفتاح الأنسب لكل لحظة.
        """
        current_hour = datetime.now().hour
        hourly   = self.hourly_data.get(current_hour, {'success': 0, 'total': 0})
        global_rate = self.get_success_rate()

        if hourly['total'] >= 5:
            hourly_rate = (hourly['success'] / hourly['total']) * 100
            predicted   = 0.7 * hourly_rate + 0.3 * global_rate
        else:
            # بيانات ساعية غير كافية → نقطة الصحة الحالية كـ fallback
            predicted = self.get_health_score()

        return max(0.0, min(100.0, predicted))

    def get_failure_breakdown(self) -> Dict[str, Any]:
        """
        تفصيل مصنَّف لجميع أنواع الفشل (مفيد للتشخيص والسجلات).
        يُعرض في get_statistics_summary.
        """
        return {
            canonical: {
                'count':      self.failure_counts.get(canonical, 0),
                'percentage': round(
                    self.failure_counts.get(canonical, 0) / max(self.total_requests, 1) * 100, 1
                ),
                'weight':     self.FAILURE_WEIGHTS.get(canonical, 1.0),
            }
            for canonical in self.FAILURE_TYPES
        }

# ============= استثناء داخلي لـ Circuit Breaker =============
class _CircuitBreakerKeyError(Exception):
    """
    استثناء داخلي يُرفع داخل _api_call ليُحسب ضمن فشل Circuit Breaker.
    يُمثّل فشلاً حقيقياً في المفتاح (server error، invalid key، network error).
    لا يُستخدم لأخطاء Rate Limit (429) التي تُعالج بـ blocked_keys.
    """
    def __init__(self, status: int = 0, message: str = "", error_type: str = "general"):
        self.status     = status
        self.message    = message
        self.error_type = error_type
        super().__init__(f"[{status}] {message}")


# ============= الفئة المحسنة الرئيسية (مع المفاتيح كما هي) =============
class EnhancedGeminiAPI:
    """إدارة محسنة لـ Gemini API مع مفاتيح متعددة"""
    
    def __init__(self, api_keys: List[str] = None):
        # المفاتيح كما كانت في الكود الأصلي
        self.api_keys = [
            "AIzaSyA9HAxV5Q-3MSkb1GZVy28ie_TI",
            

        ]
        
        if isinstance(api_keys, list):
            self.api_keys.extend([key for key in api_keys if key not in self.api_keys])
        
        # Rate limiters لكل مفتاح - إعدادات قابلة للتهيئة من ملف خارجي
        self.rate_limiters = {
            key: TokenRateLimiter(max_rpm=5, max_tpm=32000, max_rpd=25, key_id=key)
            for key in self.api_keys
        }
        
        # إحصائيات متقدمة لكل مفتاح (مع تمرير key_id لتفعيل SQLite والتنبؤ)
        self.key_stats = {key: KeyStatistics(key_id=key) for key in self.api_keys}
        
        # مفاتيح محظورة مؤقتاً
        self.blocked_keys = {}  # {key: unblock_time}
        
        # التوزيع الدائري
        self.current_key_index = 0

        # إعدادات الAPI لـ Gemini 2.5 Flash
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
        self.max_retries = 3
        # gemini-2.5-pro free tier: نافذة RPM = 60 ثانية
        # يجب الانتظار 60+ ثانية بعد كل 429 حتى تنتهي النافذة الحالية
        self.retry_delays = [65, 90, 120]
        
        # Connection pool للأداء الأفضل
        self.connector = None
        self.session = None
        
        logger.info(f"Gemini API initialized with {len(self.api_keys)} keys")

        # ---- Generation Profiles: إعدادات توليد مخصصة حسب نوع الطلب ----
        # temperature يأتي من المستدعي؛ هنا نضبط topK و topP بدقة حسب النوع
        self.GENERATION_PROFILES: Dict[str, Dict[str, Any]] = {
            # الترجمة الأدبية والسياقية — تنوع عالٍ لضمان جودة أسلوبية
            "translation":               {"topK": 20, "topP": 0.9},
            "complete_translation":      {"topK": 20, "topP": 0.9},
            "contextual_translation":    {"topK": 20, "topP": 0.9},
            "chapter_title_translation": {"topK": 20, "topP": 0.9},
            # ترجمة القطع الكبيرة chunk-based — نفس إعدادات الترجمة الأدبية
            "chunk_translation":         {"topK": 20, "topP": 0.9},
            "subchunk_translation":      {"topK": 20, "topP": 0.9},
            # استخراج المصطلحات — دقة عالية، تنوع منخفض
            "terminology_extraction":    {"topK": 10, "topP": 0.7},
            "term_extraction":           {"topK": 10, "topP": 0.7},
            # المراجعة والتصحيح النهائي — أقصى دقة وأدنى عشوائية
            "completion_review":         {"topK": 8,  "topP": 0.6},
            "comprehensive_correction":  {"topK": 8,  "topP": 0.6},
            "final_completion":          {"topK": 8,  "topP": 0.6},
            "final_cleanup":             {"topK": 8,  "topP": 0.6},
            "final_review":              {"topK": 8,  "topP": 0.6},
            # التصحيح المستهدف للجمل المشكلة — دقة قصوى
            "targeted_correction":       {"topK": 8,  "topP": 0.6},
        }
        # الإعدادات الافتراضية لأي نوع طلب غير مصنَّف
        self._default_profile: Dict[str, Any] = {"topK": 12, "topP": 0.8}

        # ---- Circuit Breakers: تطبيق داخلي بديل عن pybreaker لتجنب تعارض async ----
        # بعد 5 فشل حقيقي متتالٍ → يُوقَف المفتاح 5 دقائق (300 ثانية)
        # _cb_open_until: {key → timestamp_until_open}
        self._cb_open_until: Dict[str, float] = {}
        # نُبقي circuit_breakers للتوافق مع الكود الخارجي (القراءة فقط)
        self.circuit_breakers: Dict[str, Any] = {}
        logger.info(f"Internal Circuit Breakers initialized: fail_max=5, timeout=300s")
    
    async def _ensure_session(self):
        """التأكد من وجود جلسة HTTP نشطة مع connection pooling"""
        if not self.session or self.session.closed:
            self.connector = aiohttp.TCPConnector(
                limit=100,  # الحد الأقصى للاتصالات
                ttl_dns_cache=300,  # cache DNS لـ 5 دقائق
                enable_cleanup_closed=True
            )
            timeout = aiohttp.ClientTimeout(total=300)
            self.session = aiohttp.ClientSession(
                connector=self.connector,
                timeout=timeout
            )
    
    def _unblock_keys(self):
        """إلغاء حظر المفاتيح التي انتهت فترة حظرها"""
        current_time = time.time()
        keys_to_unblock = []
        
        for key, unblock_time in self.blocked_keys.items():
            if current_time >= unblock_time:
                keys_to_unblock.append(key)
        
        for key in keys_to_unblock:
            del self.blocked_keys[key]
            logger.info(f"Unblocked key: {key[:10]}...")
    
    def estimate_tokens(self, text: str) -> int:
        """تقدير دقيق لعدد التوكنز: يدعم tiktoken للعربية والإنجليزية."""
        return _estimate_tokens_smart(text)

    def _get_generation_profile(self, request_type: str) -> Dict[str, Any]:
        """
        إرجاع إعدادات topK و topP المناسبة لنوع الطلب.
        تعود إلى الإعدادات الافتراضية إذا لم يكن النوع معروفاً.
        """
        return self.GENERATION_PROFILES.get(request_type, self._default_profile)

    async def get_optimal_api_key(self, estimated_tokens: int = 0) -> Optional[str]:
        """الحصول على المفتاح التالي المتاح باستخدام Round-Robin مع انتظار ذكي"""
        while True:
            self._unblock_keys()
            
            num_keys = len(self.api_keys)
            if num_keys == 0:
                return None

            checked_keys = 0
            
            # البحث عن مفتاح متاح
            while checked_keys < num_keys:
                key = self.api_keys[self.current_key_index]
                self.current_key_index = (self.current_key_index + 1) % num_keys
                checked_keys += 1

                # تخطي المفاتيح المحظورة
                if key in self.blocked_keys:
                    continue

                # تخطي المفاتيح التي فتح دائرتها (Circuit Breaker داخلي)
                if key in self._cb_open_until:
                    if time.time() < self._cb_open_until[key]:
                        continue
                    else:
                        # انتهت فترة الإيقاف → HALF_OPEN: نسمح بمرور طلب واحد
                        del self._cb_open_until[key]

                # التحقق من rate limit
                if not self.rate_limiters[key].can_make_request(estimated_tokens):
                    continue

                # التحقق من صحة المفتاح باستخدام التنبؤ الذكي
                # (يدمج معدل النجاح الساعي التاريخي + الصحة الكلية)
                predicted_perf = self.key_stats[key].get_predicted_performance()
                if predicted_perf < 10:
                    continue

                return key

            # إذا وصلنا هنا، يعني أن كل المفاتيح مشغولة، يجب الانتظار الذكي
            min_wait = float('inf')
            all_daily_exhausted = True

            for key in self.api_keys:
                if key not in self.blocked_keys:
                    # التحقق إذا كان المفتاح استنفد الحد اليومي
                    if len(self.rate_limiters[key].daily_requests) < self.rate_limiters[key].max_rpd:
                        all_daily_exhausted = False

                    wait_time = self.rate_limiters[key].time_until_next_request(estimated_tokens)
                    if wait_time < min_wait:
                        min_wait = wait_time

            if all_daily_exhausted:
                logger.error("All keys have exhausted their daily limit! Must wait until the next day or add new keys.")
                # الانتظار لمدة طويلة ثم المحاولة مجدداً عبر الحلقة (تجنب الاستدعاء المتكرر recursive)
                await asyncio.sleep(60)
                continue

            if min_wait < float('inf') and min_wait > 0:
                logger.info(f"All keys are currently busy, smartly waiting for {min_wait:.1f} seconds...")
                await asyncio.sleep(min_wait + 0.5)
                continue
            
            # إذا فشلت كل المحاولات، أعد تعيين المفاتيح المحظورة وانتظر
            # انتظار نافذة RPM الكاملة (60 ثانية) ثم إعادة المحاولة
            logger.warning("All keys are blocked, waiting for RPM window to reset (65s)...")
            await asyncio.sleep(65)
            self.blocked_keys.clear()
            # Loop will continue
    
    async def make_precision_request(self, prompt: str, system_instruction: str = "", 
                                   temperature: float = 0.05, max_tokens: int = 8192,
                                   request_type: str = "translation") -> Tuple[Optional[str], float, Optional[str]]:
        """
        إرسال طلب دقيق مع تحسينات شاملة:
          ✅ systemInstruction كحقل مستقل في payload (وزن أعلى لدى النموذج)
          ✅ إعدادات توليد مخصصة (topK، topP) حسب نوع الطلب
          ✅ maxOutputTokens يُحسب ديناميكياً من طول النص المُدخَل
          ✅ Circuit Breaker لكل مفتاح (5 فشل → إيقاف 5 دقائق ثم اختبار واحد)
        """
        # التأكد من وجود جلسة نشطة
        await self._ensure_session()

        # --- حساب التوكنز وإعداد maxOutputTokens ---
        # gemini-2.5-pro free tier: max 32,000 TPM per key
        # 800-word chunk ≈ 1200 input tokens + 8192 max output = ~9392 total per request
        # نستخدم max_tokens كاملاً — النموذج يتوقف عند الانتهاء الطبيعي قبل الحد
        dynamic_max_tokens = max_tokens  # 8192 — يحترم حصة TPM الحرة
        estimated_input_tokens = self.estimate_tokens(prompt + system_instruction)
        estimated_output_tokens = dynamic_max_tokens
        total_estimated_tokens  = estimated_input_tokens + estimated_output_tokens

        # --- profile للتوليد حسب نوع الطلب ---
        profile = self._get_generation_profile(request_type)

        for attempt in range(self.max_retries):
            api_key = await self.get_optimal_api_key(total_estimated_tokens)
            if not api_key:
                logger.error("No API keys available")
                return None
            
            # تسجيل الطلب والتوكنز المقدرة
            self.rate_limiters[api_key].add_request(total_estimated_tokens)
            request_start = time.time()
            
            headers = {
                'Content-Type': 'application/json',
                'User-Agent': 'Professional-Translation-System/Enhanced'
            }
            
            # --- بناء الـ payload مع systemInstruction كحقل مستقل ---
            payload: Dict[str, Any] = {
                "contents": [
                    {
                        "role": "user",
                        "parts": [{"text": prompt}]
                    }
                ],
                "generationConfig": {
                    "temperature":      temperature,
                    "topK":             profile["topK"],
                    "topP":             profile["topP"],
                    "maxOutputTokens":  dynamic_max_tokens,
                    "candidateCount":   1,
                    "stopSequences":    ["###TRANSLATION_END###", "###END###"]
                },
                "safetySettings": [
                    {"category": "HARM_CATEGORY_HARASSMENT",        "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_HATE_SPEECH",       "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
                ]
            }
            # إضافة systemInstruction كحقل مستقل فقط إذا كان غير فارغ
            if system_instruction and system_instruction.strip():
                payload["systemInstruction"] = {
                    "parts": [{"text": system_instruction.strip()}]
                }

            url = f"{self.base_url}?key={api_key}"

            # --- تنفيذ طلب HTTP مع Circuit Breaker داخلي ---
            result_holder: Dict[str, Any] = {}

            async def _api_call():
                """
                دالة داخلية تنفذ طلب HTTP وترفع _CircuitBreakerKeyError
                عند أي فشل حقيقي (server/network/invalid key).
                لا ترفع استثناءً للـ 429 (rate limit) — يُعالج بشكل منفصل.
                """
                try:
                    async with self.session.post(url, json=payload, headers=headers) as response:
                        result_holder["status"] = response.status
                        if response.status == 200:
                            result_holder["json"] = await response.json()
                        elif response.status == 429:
                            # Rate limit ليس فشلاً حقيقياً للمفتاح — لا يفتح الدائرة
                            result_holder["rate_limited"] = True
                        elif response.status in [401, 403]:
                            text = await response.text()
                            raise _CircuitBreakerKeyError(response.status, text, "invalid_key")
                        elif response.status >= 500:
                            raise _CircuitBreakerKeyError(
                                response.status, f"Server error {response.status}", "server_error"
                            )
                        else:
                            text = await response.text()
                            raise _CircuitBreakerKeyError(response.status, text, "api_error")
                except asyncio.TimeoutError:
                    raise _CircuitBreakerKeyError(0, "Request timed out", "timeout")
                except _CircuitBreakerKeyError:
                    raise
                except Exception as e:
                    raise _CircuitBreakerKeyError(0, str(e), "exception")

            try:
                logger.info(
                    f"Sending {request_type} request (attempt {attempt+1}) "
                    f"key={api_key[:10]}... "
                    f"maxTokens={dynamic_max_tokens} topK={profile['topK']} topP={profile['topP']}"
                )
                await _api_call()

                response_time = time.time() - request_start

                # --- معالجة النتيجة ---
                if result_holder.get("json"):
                    result = result_holder["json"]
                    if (
                        "candidates" in result
                        and len(result["candidates"]) > 0
                        and "content" in result["candidates"][0]
                        and "parts" in result["candidates"][0]["content"]
                        and len(result["candidates"][0]["content"]["parts"]) > 0
                    ):
                        content = result["candidates"][0]["content"]["parts"][0]["text"]
                        # ── كشف التوقف المبكر بسبب MAX_TOKENS ──
                        finish_reason = result["candidates"][0].get("finishReason", "STOP")
                        if finish_reason == "MAX_TOKENS":
                            logger.warning(
                                f"⚠️ MAX_TOKENS reached for {request_type} | "
                                f"key={api_key[:10]}... | output truncated! "
                                f"Consider reducing chunk size below 1800 words."
                            )
                            # نُضيف علامة نهاية مقطوعة ليستطيع المُستدعي اكتشافها
                            content = content.strip() + "\n###TRUNCATED###"
                        self.key_stats[api_key].record_success(response_time)
                        logger.info(f"Request {request_type} succeeded | key={api_key[:10]}... | time={response_time:.2f}s | finishReason={finish_reason}")
                        return content.strip(), response_time, api_key
                    else:
                        logger.warning(f"Unexpected response from Gemini: {result}")
                        should_alert = self.key_stats[api_key].record_failure("invalid_response")
                        if should_alert:
                            logger.warning(f"Intelligence Alert: Key {api_key[:10]}... failed 3 consecutive times", key_status="consecutive_failures")
                            console.print(f"[bold yellow]⚠️ Alert: Key {api_key[:10]} failed 3 consecutive times but remains in use.[/bold yellow]")

                elif result_holder.get("rate_limited"):
                    logger.warning(f"Rate limit exceeded for key {api_key[:10]}... waiting")
                    should_alert = self.key_stats[api_key].record_failure("rate_limit")
                    if should_alert:
                        logger.warning(f"Intelligence Alert: Key {api_key[:10]}... failed 3 consecutive times (Rate Limit)", key_status="consecutive_failures")
                        console.print(f"[bold yellow]⚠️ Alert: Key {api_key[:10]} failed 3 consecutive times due to rate limits.[/bold yellow]")
                    # إبلاغ النظام التكيّفي
                    self.rate_limiters[api_key].record_429_error()
                    block_duration = self.retry_delays[min(attempt, len(self.retry_delays)-1)]
                    self.blocked_keys[api_key] = time.time() + block_duration
                    await asyncio.sleep(block_duration)

            except _CircuitBreakerKeyError as e:
                # فشل حقيقي — سجّل وعالج حسب نوع الخطأ
                response_time = time.time() - request_start
                error_type = e.error_type

                should_alert = self.key_stats[api_key].record_failure(error_type)
                if should_alert:
                    logger.warning(
                        f"Intelligence Alert: Key {api_key[:10]}... failed 3 consecutive times ({error_type})",
                        key_status="consecutive_failures"
                    )
                    console.print(
                        f"[bold red]⚠️ Alert: Key {api_key[:10]} is facing consecutive {error_type} errors![/bold red]"
                    )

                # Circuit Breaker داخلي: فتح الدائرة عند 5 إخفاقات متتالية (300 ثانية)
                consec = self.key_stats[api_key].consecutive_failures
                if consec >= 5 and api_key not in self._cb_open_until:
                    self._cb_open_until[api_key] = time.time() + 300
                    logger.warning(
                        f"[CircuitBreaker] Circuit OPEN for key {api_key[:10]} "
                        f"({consec} consecutive failures). Suspended for 300s."
                    )
                    console.print(
                        f"[bold red]⚡ Circuit Breaker OPEN: key {api_key[:10]} suspended for 300s[/bold red]"
                    )
                    continue  # جرّب مفتاحاً آخر

                if error_type == "timeout":
                    logger.warning(f"Request {request_type} timed out (attempt {attempt+1})")
                    await asyncio.sleep(self.retry_delays[min(attempt, len(self.retry_delays)-1)])

                elif error_type == "invalid_key":
                    logger.error(f"Gemini API error {e.status} for key {api_key[:10]}: {e.message}")
                    self.blocked_keys[api_key] = time.time() + 3600  # حظر لمدة ساعة
                    # لا داعي للانتظار — انتقل لمفتاح آخر
                    continue

                elif error_type == "server_error":
                    logger.error(f"Gemini server error {e.status}")
                    await asyncio.sleep(self.retry_delays[min(attempt, len(self.retry_delays)-1)])

                else:
                    logger.error(f"Unexpected API error [{e.status}]: {e.message}")
                    await asyncio.sleep(self.retry_delays[min(attempt, len(self.retry_delays)-1)])

            except Exception as e:
                logger.error(f"Error in {request_type} request (attempt {attempt+1}): {str(e)}")
                should_alert = self.key_stats[api_key].record_failure("exception")
                if should_alert:
                    logger.warning(f"Intelligence Alert: Key {api_key[:10]}... failed 3 consecutive times (Exception)", key_status="consecutive_failures")
                await asyncio.sleep(self.retry_delays[min(attempt, len(self.retry_delays)-1)])
        
        logger.error(f"Request {request_type} failed after {self.max_retries} attempts")
        return None, 0.0, None
    
    def get_statistics_summary(self) -> Dict[str, Any]:
        """الحصول على ملخص إحصائيات جميع المفاتيح"""
        summary = {
            'total_keys': len(self.api_keys),
            'active_keys': len([k for k in self.api_keys if k not in self.blocked_keys]),
            'blocked_keys': len(self.blocked_keys),
            'keys_performance': []
        }
        
        for key in self.api_keys:
            stats = self.key_stats[key]
            key_info = {
                'key_preview': f"{key[:10]}...",
                'health_score': stats.get_health_score(),
                'predicted_performance': round(stats.get_predicted_performance(), 1),
                'success_rate': stats.get_success_rate(),
                'total_requests': stats.total_requests,
                'successful_requests': stats.successful_requests,
                'failed_requests': stats.failed_requests,
                'avg_response_time': round(stats.average_response_time, 2),
                'is_blocked': key in self.blocked_keys,
                'failure_breakdown': stats.get_failure_breakdown(),
            }
            summary['keys_performance'].append(key_info)
        
        # ترتيب حسب الأداء المتنبأ به (يدمج الصحة والتاريخ الساعي)
        summary['keys_performance'].sort(
            key=lambda x: x['predicted_performance'], reverse=True
        )
        
        return summary
    
    async def cleanup(self):
        """تنظيف الموارد عند الانتهاء"""
        if self.session and not self.session.closed:
            await self.session.close()
        if self.connector:
            await self.connector.close()
        logger.info("Gemini API resources cleaned up")

class ComprehensiveContentProcessor:
    """
    معالج المحتوى الشامل الذكي - النسخة المطورة
    ─────────────────────────────────────────────
    التحسينات الجوهرية عن النسخة السابقة:

    1. قائمة بيضاء ديناميكية: أي كلمة إنجليزية تتكرر 3+ مرات في الكتاب
       تُضاف تلقائياً (أسماء شخصيات/أماكن مقبولة) بدلاً من القائمة الثابتة الصغيرة.

    2. كشف اللغة الحقيقي: استخدام langdetect (اختياري) لفحص لغة الجملة
       بدلاً من الاعتماد فقط على Regex الذي يُخطئ مع أسماء الأعلام.

    3. استخراج الكيانات المسماة: spaCy (اختياري) لاستخراج أسماء الشخصيات
       والأماكن والمنظمات وإعفائها من علامات "المحتوى الأجنبي".

    4. قياس تغطية ذكي: نسبة الكلمات مع معامل تصحيح لغوي بدلاً من
       مقارنة عدد الجمل التي تُعطي نسب مضللة (جملة → جملتان = 200% زائف).

    5. تحويل أرقام ذكي: يحمي الأنماط الخاصة (URLs، ISBN، MP3، IPv4...)
       قبل التحويل بدلاً من استبدال كل رقم عشوائياً.

    6. علامات ترقيم شاملة: extract_content_segments تدعم ؟ ؛ ! العربية.
    """

    # ── القائمة البيضاء الثابتة الموسّعة (مصطلحات دولية مقبولة دائماً) ──
    STATIC_ACCEPTABLE_ENGLISH: frozenset = frozenset([
        # تقنية وإنترنت
        'OK', 'PDF', 'ISBN', 'URL', 'ID', 'TV', 'PC', 'CD', 'DVD',
        'AI', 'API', 'HTML', 'CSS', 'XML', 'JSON', 'SQL', 'HTTP', 'HTTPS',
        'FBI', 'CIA', 'NATO', 'UN', 'EU', 'US', 'UK', 'UAE',
        'WiFi', 'GPS', 'SMS', 'DNA', 'RNA', 'VIP', 'ATM', 'PIN', 'SIM',
        'MP3', 'MP4', 'USB', 'RAM', 'ROM', 'CPU', 'GPU', 'SSD', 'HDD',
        # ألقاب وتشريفات
        'Dr', 'Mr', 'Mrs', 'Ms', 'Jr', 'Sr', 'Prof', 'PhD', 'MD', 'CEO',
        # وحدات ورموز
        'km', 'kg', 'cm', 'mm', 'Hz', 'MHz', 'GHz', 'GB', 'MB', 'KB',
        'PM', 'AM', 'AD', 'BC',
    ])

    # ── أنماط محمية من تحويل الأرقام ──
    _NUMBER_PROTECTION_RE = re.compile(
        r'https?://[^\s]+'           # URLs كاملة
        r'|ISBN[-\s]?\d[\d\s\-]+'   # أكواد ISBN
        r'|\b[A-Za-z]+\d+\w*\b'     # كلمات مختلطة: MP3، IPv4، B2B
        r'|\b\d+[A-Za-z]+\w*\b'     # كلمات مختلطة معكوسة: 2D، 3G
    )

    def __init__(self):
        # ── القائمة البيضاء الديناميكية (تُبنى من تحليل الكتاب الكامل) ──
        self._dynamic_whitelist: set = set()
        # ── الكيانات المسماة المستخرجة بـ spaCy ──
        self._named_entities: set = set()
        # ── علامة: هل بُنيت القائمة البيضاء للكتاب كاملاً؟ ──
        self._whitelist_built: bool = False

        # ── فحص توفر المكتبات الاختيارية (مرة واحدة عند الإنشاء) ──
        self._langdetect_available: bool = self._check_import('langdetect')
        self._spacy_available: bool = self._check_import('spacy')
        self._rapidfuzz_available: bool = self._check_import('rapidfuzz')
        self._spacy_nlp = None  # تحميل كسول عند الحاجة

    # ════════════════════════════════════════════════════════════
    #  أدوات داخلية مساعدة
    # ════════════════════════════════════════════════════════════

    @staticmethod
    def _check_import(module_name: str) -> bool:
        """فحص وجود مكتبة اختيارية دون رفع استثناء"""
        try:
            __import__(module_name)
            return True
        except ImportError:
            return False

    def _load_spacy_model(self):
        """تحميل نموذج spaCy الإنجليزي بشكل كسول (مرة واحدة فقط)"""
        if self._spacy_nlp is not None or not self._spacy_available:
            return
        try:
            import spacy
            try:
                self._spacy_nlp = spacy.load("en_core_web_sm")
            except OSError:
                # محاولة تنزيل النموذج تلقائياً
                import subprocess
                subprocess.run(
                    ["python", "-m", "spacy", "download", "en_core_web_sm"],
                    capture_output=True, timeout=120
                )
                try:
                    self._spacy_nlp = spacy.load("en_core_web_sm")
                except Exception:
                    self._spacy_available = False
        except Exception:
            self._spacy_available = False

    def _extract_named_entities_spacy(self, text: str) -> set:
        """
        استخراج الكيانات المسماة (أشخاص، أماكن، منظمات) باستخدام spaCy.
        يُعيد set فارغة إذا لم يكن spaCy متاحاً.
        """
        entities: set = set()
        if not self._spacy_available:
            return entities
        self._load_spacy_model()
        if self._spacy_nlp is None:
            return entities
        try:
            # معالجة أقصى 50000 حرف لتجنب الثقل
            doc = self._spacy_nlp(text[:50000])
            target_labels = {'PERSON', 'GPE', 'LOC', 'ORG', 'NORP', 'FAC', 'PRODUCT', 'EVENT'}
            for ent in doc.ents:
                if ent.label_ in target_labels:
                    full_name = ent.text.strip()
                    entities.add(full_name)
                    entities.add(full_name.upper())
                    # أضف كل كلمة منفردة من الاسم المركّب
                    for token in full_name.split():
                        t = token.strip('.,;:!?"\' ')
                        if len(t) >= 2:
                            entities.add(t)
                            entities.add(t.upper())
        except Exception:
            pass
        return entities

    def _detect_sentence_language(self, sentence: str) -> str:
        """
        كشف لغة الجملة باستخدام langdetect.
        يُعيد: 'ar' | 'en' | 'unknown'
        """
        if not self._langdetect_available or len(sentence.strip()) < 8:
            return 'unknown'
        try:
            from langdetect import detect
            return detect(sentence.strip())
        except Exception:
            return 'unknown'

    def _is_acceptable_english(self, word: str) -> bool:
        """
        هل هذه الكلمة الإنجليزية مقبولة ولا تستوجب الترجمة؟
        ترتيب الفحص: ثابتة → ديناميكية → كيانات spaCy
        """
        if word.upper() in self.STATIC_ACCEPTABLE_ENGLISH:
            return True
        w_lower = word.lower()
        if (word in self._dynamic_whitelist
                or w_lower in self._dynamic_whitelist
                or word.upper() in self._dynamic_whitelist):
            return True
        if word in self._named_entities or word.upper() in self._named_entities:
            return True
        return False

    def _calculate_smart_coverage(self, original_text: str, translated_text: str) -> float:
        """
        قياس نسبة التغطية بعدد الكلمات (لا بعدد الجمل).

        المنطق الصحيح:
        ─────────────
        العربية أكثر إسهاباً من الإنجليزية، أي أن الترجمة الكاملة تُنتج عادةً
        كلمات أكثر. لذا نستخدم النسبة المباشرة بدون معامل تصحيح، مع عتبة
        منخفضة (65%) لا تُشغَّل إلا عند فجوة حقيقية في المحتوى.

        أمثلة:
          100 كلمة إنجليزية → 100 عربية : 100%  لا مراجعة ✓
          100 كلمة إنجليزية →  80 عربية :  80%  لا مراجعة ✓ (ترجمة موجزة)
          100 كلمة إنجليزية →  60 عربية :  60%  مراجعة مطلوبة ✓ (فجوة حقيقية)
        """
        if not original_text or not translated_text:
            return 0.0
        original_words = len(re.findall(r'\b\w+\b', original_text))
        translated_words = len(re.findall(r'\b\w+\b', translated_text))
        if original_words == 0:
            return 100.0
        return min((translated_words / original_words) * 100.0, 100.0)

    # ════════════════════════════════════════════════════════════
    #  الواجهة العامة: بناء القائمة البيضاء للكتاب كاملاً
    # ════════════════════════════════════════════════════════════

    def build_book_whitelist(self, full_text: str, min_occurrences: int = 3):
        """
        تحليل النص الكامل للكتاب وبناء قائمة بيضاء ديناميكية.

        الخوارزمية:
        1. تحسب تكرار كل كلمة إنجليزية في النص الكامل.
        2. أي كلمة تتكرر >= min_occurrences هي على الأرجح اسم شخصية/مكان → تُضاف.
        3. إذا كان spaCy متاحاً: تُضاف الكيانات المسماة مباشرة أيضاً.

        يجب استدعاء هذه الدالة مرة واحدة قبل بدء الترجمة لأفضل نتيجة.
        """
        if not full_text:
            return

        # حساب تكرارات الكلمات الإنجليزية
        word_counts: Dict[str, int] = {}
        for word in re.findall(r'\b[A-Za-z]{2,}\b', full_text):
            key = word.lower()
            word_counts[key] = word_counts.get(key, 0) + 1

        # إضافة الكلمات المتكررة بما يكفي
        added_freq = 0
        for word, count in word_counts.items():
            if count >= min_occurrences:
                self._dynamic_whitelist.add(word)           # small
                self._dynamic_whitelist.add(word.upper())   # UPPER
                self._dynamic_whitelist.add(word.capitalize())  # Title
                added_freq += 1

        # إضافة الكيانات المسماة من spaCy (إذا كان متاحاً)
        entities = self._extract_named_entities_spacy(full_text)
        self._named_entities.update(entities)
        self._dynamic_whitelist.update(entities)

        self._whitelist_built = True
        logger.info(
            f"[ContentProcessor] ✅ قائمة بيضاء: {added_freq} كلمة متكررة "
            f"+ {len(self._named_entities)} كيان مسمى (spaCy) "
            f"= {len(self._dynamic_whitelist)} إجمالي"
        )

    # ════════════════════════════════════════════════════════════
    #  الدوال الثابتة (Static) — لا تحتاج حالة الكائن
    # ════════════════════════════════════════════════════════════

    @staticmethod
    def number_to_arabic_text(number: int) -> str:
        """تحويل الرقم إلى كتابة عربية ترتيبية"""
        arabic_numbers = {
            1: "الأول", 2: "الثاني", 3: "الثالث", 4: "الرابع", 5: "الخامس",
            6: "السادس", 7: "السابع", 8: "الثامن", 9: "التاسع", 10: "العاشر",
            11: "الحادي عشر", 12: "الثاني عشر", 13: "الثالث عشر", 14: "الرابع عشر", 15: "الخامس عشر",
            16: "السادس عشر", 17: "السابع عشر", 18: "الثامن عشر", 19: "التاسع عشر", 20: "العشرون",
            21: "الواحد والعشرون", 22: "الثاني والعشرون", 23: "الثالث والعشرون",
            24: "الرابع والعشرون", 25: "الخامس والعشرون", 26: "السادس والعشرون",
            27: "السابع والعشرون", 28: "الثامن والعشرون", 29: "التاسع والعشرون", 30: "الثلاثون"
        }
        if number <= 30:
            return arabic_numbers.get(number, f"الفصل {number}")
        return f"الفصل {number}"

    @staticmethod
    def convert_numbers_to_arabic(text: str) -> str:
        """
        تحويل الأرقام اللاتينية إلى عربية بشكل ذكي.

        يحمي الأنماط التالية من التحويل:
          • URLs  (https://example.com:8080)
          • أكواد ISBN
          • كلمات مختلطة مثل MP3، IPv4، B2B، 2D، 3G
        وذلك لتجنب تخريب النصوص التقنية والأكواد الخاصة.
        """
        if not text:
            return text

        english_to_arabic = {
            '0': '٠', '1': '١', '2': '٢', '3': '٣', '4': '٤',
            '5': '٥', '6': '٦', '7': '٧', '8': '٨', '9': '٩'
        }

        result: List[str] = []
        last_end = 0

        for match in ComprehensiveContentProcessor._NUMBER_PROTECTION_RE.finditer(text):
            # الجزء السابق للنمط المحمي: نحوّل أرقامه
            segment = text[last_end:match.start()]
            for eng, ar in english_to_arabic.items():
                segment = segment.replace(eng, ar)
            result.append(segment)
            # النمط المحمي نفسه: نبقيه كما هو
            result.append(match.group())
            last_end = match.end()

        # الجزء المتبقي بعد آخر نمط محمي
        segment = text[last_end:]
        for eng, ar in english_to_arabic.items():
            segment = segment.replace(eng, ar)
        result.append(segment)

        return ''.join(result)

    # ════════════════════════════════════════════════════════════
    #  الدوال الأساسية (Instance) — تستخدم القوائم الديناميكية
    # ════════════════════════════════════════════════════════════

    def extract_content_segments(self, text: str) -> List[str]:
        """
        استخراج أجزاء المحتوى للمقارنة.
        يدعم علامات الترقيم العربية (؟ ؛ !) والإنجليزية (.!?) معاً.
        """
        sentences = re.split(r'[.!?؟!؛]+', text)
        return [s.strip() for s in sentences if len(s.strip()) > 10]

    def detect_incomplete_translation(self, original_text: str, translated_text: str) -> Dict[str, Any]:
        """
        كشف الترجمة غير المكتملة بطريقة ذكية ثلاثية المراحل:

        المرحلة 1 — قائمة بيضاء مؤقتة من النص الأصلي:
            أي كلمة إنجليزية تتكرر 3+ مرات في الأصل → مقبولة كاسم علم.

        المرحلة 2 — فلترة الكلمات المتبقية:
            يُطبّق _is_acceptable_english (ثابتة + ديناميكية + spaCy).

        المرحلة 3 — قياس التغطية بنسبة الكلمات (لا عدد الجمل):
            يتجنب الـ 200% الزائفة حين تُصبح جملة → جملتين بالعربية.
        """
        issues: Dict[str, Any] = {
            'missing_segments': [],
            'untranslated_english': [],
            'incomplete_phrases': [],
            'missing_names': [],
            'coverage_percentage': 0.0
        }

        # ── المرحلة 1: قائمة بيضاء مؤقتة من تكرارات الأصل ──
        temp_whitelist: set = set()
        if original_text:
            freq: Dict[str, int] = {}
            for w in re.findall(r'\b[A-Za-z]{2,}\b', original_text):
                freq[w.lower()] = freq.get(w.lower(), 0) + 1
            for w, cnt in freq.items():
                if cnt >= 3:
                    temp_whitelist.add(w)
                    temp_whitelist.add(w.upper())
                    temp_whitelist.add(w.capitalize())

        # ── المرحلة 2: فحص الكلمات الإنجليزية في الترجمة ──
        seen: set = set()
        for word in re.findall(r'\b[A-Za-z]{2,}\b', translated_text):
            w_low = word.lower()
            if w_low in seen:
                continue
            seen.add(w_low)
            if self._is_acceptable_english(word):
                continue
            if word in temp_whitelist or w_low in temp_whitelist:
                continue
            # فحص langdetect: إذا الجملة المحيطة عربية → الكلمة على الأرجح مقبولة
            if self._langdetect_available:
                # استخرج الجملة المحيطة بالكلمة (50 حرف يميناً ويساراً)
                idx = translated_text.find(word)
                if idx != -1:
                    surrounding = translated_text[max(0, idx-50): idx+50+len(word)]
                    lang = self._detect_sentence_language(surrounding)
                    if lang == 'ar':
                        continue  # الجملة عربية → الكلمة مقبولة ضمن سياقها
            issues['untranslated_english'].append(word)

        # ── المرحلة 3: قياس التغطية الذكي ──
        issues['coverage_percentage'] = self._calculate_smart_coverage(
            original_text, translated_text
        )

        return issues

    def needs_completion_review(self, original_text: str, translated_text: str) -> bool:
        """تحديد ما إذا كانت الترجمة تحتاج مراجعة لضمان الاكتمال.
        العتبة 65% تمثل فجوة حقيقية في المحتوى لا مجرد فارق أسلوبي."""
        issues = self.detect_incomplete_translation(original_text, translated_text)
        return (
            len(issues['untranslated_english']) > 0
            or issues['coverage_percentage'] < 65.0
        )

    def has_any_foreign_content(self, text: str, original_text: Optional[str] = None) -> bool:
        """
        التحقق من وجود محتوى أجنبي حقيقي في النص.

        طبقات الفلترة بالترتيب:
        1. القائمة الثابتة الموسّعة (مصطلحات تقنية/دولية)
        2. القائمة الديناميكية (أسماء متكررة 3+ مرات في الكتاب)
        3. كلمات موجودة في النص الأصلي (مُمررة كـ original_text) → أسماء أعلام مُبقاة
        4. Heuristic اسم علم: كلمة تبدأ بحرف كبير محاطة بنص عربي → مقبولة
        5. langdetect: إذا السياق المحيط عربي → الكلمة مقبولة
        """
        # قائمة بيضاء مؤقتة: كل كلمة إنجليزية موجودة في الأصل
        temp_whitelist: set = set()
        if original_text:
            temp_whitelist = set(
                w for w in re.findall(r'\b[A-Za-z]{2,}\b', original_text)
            )

        for word in re.findall(r'\b[A-Za-z]{2,}\b', text):
            # 1. القائمة الثابتة والديناميكية
            if self._is_acceptable_english(word):
                continue
            # 2. كلمة موجودة في النص الأصلي (اسم علم مُبقى عمداً)
            if word in temp_whitelist or word.lower() in temp_whitelist:
                continue
            # 3. Heuristic: كلمة مبدوءة بحرف كبير محاطة بنص عربي → اسم علم مُبقى
            if word[0].isupper():
                idx = text.find(word)
                if idx > 5:  # ليست في بداية النص تماماً
                    before = text[max(0, idx - 15):idx]
                    # إذا الحرف السابق مباشرة عربي أو مسافة بعد عربي
                    if any('\u0600' <= c <= '\u06FF' for c in before):
                        continue
            # 4. langdetect للسياق المحيط
            if self._langdetect_available:
                idx = text.find(word)
                if idx != -1:
                    surrounding = text[max(0, idx - 50): idx + 50 + len(word)]
                    lang = self._detect_sentence_language(surrounding)
                    if lang == 'ar':
                        continue
            # كلمة إنجليزية غير مقبولة بأي معيار → محتوى أجنبي حقيقي
            return True

        return False


class CompleteTranslationEngine:
    """
    محرك الترجمة الكاملة - المُحسَّن بالنقاط الست من خريطة التطوير

    التحسينات المطبقة:
    ─────────────────
    ① systemInstruction كحقل مستقل (لا دمج في prompt)
    ② ملف معرفة الكتاب المنظم (شخصيات + أماكن + أحداث + أسلوب)
    ③ كشف النوع/النبرة بالنقاط مع عينات ثلاثية (لا if/elif بأولوية ثابتة)
    ④ سلسلة مراجعة مستهدفة: كشف الجمل المشكلة محلياً → إرسالها فقط للـ API
    ⑤ Few-Shot Prompting بأمثلة حقيقية من الكتاب (بعد الفصلين الأولين)
    ⑥ استخراج المصطلحات ضمن الترجمة نفسها (لا API call منفصل)
    """

    # ── حجم القطعة الآمن — مُحسَّن لـ gemini-2.5-pro free tier ──
    # gemini-2.5-pro: RPM=5, TPM=32000, RPD=25
    # 800 كلمة ≈ 1200 توكن إدخال + ~3000 توكن إخراج = ~4200 توكن/طلب
    # 32000 TPM ÷ 4200 ≈ 7 طلب/دقيقة — ضمن حد الـ RPM=5 بأمان
    SAFE_CHUNK_WORDS: int = 800

    def __init__(self, api_manager: EnhancedGeminiAPI, target_language: str = "Arabic"):
        self.api_manager = api_manager
        self.target_language = target_language
        self.translation_memory = {}
        self.terminology_database: Dict[str, str] = {}   # {term_en: term_ar}
        self.term_frequency: Dict[str, int] = {}         # ① تردد كل مصطلح
        self.context_history: List[str] = []             # للتوافق مع الكود الخارجي
        self.content_processor = ComprehensiveContentProcessor()

        # إعدادات الترجمة السياقية
        self.genre_detection = True
        self.emotional_tone_preservation = True
        self.stylistic_adaptation = True

        # ── ② ملف معرفة الكتاب المنظم ──
        self.book_knowledge: Dict[str, Any] = {
            'title':            '',
            'author':           '',
            'characters':       {},   # {en_name: {'arabic': ..., 'description': ...}}
            'places':           {},   # {en_name: ar_name}
            'events':           [],   # [{'chapter': ..., 'summary': ...}]
            'style_notes':      [],
            'few_shot_examples': []   # ⑤ أمثلة حقيقية بعد الفصلين الأولين
        }
        self.chapters_completed: int = 0  # لتفعيل Few-Shot بعد فصلين

    # ════════════════════════════════════════════════════════════
    #  ③ كشف النوع والنبرة بالنقاط مع عينات ثلاثية
    # ════════════════════════════════════════════════════════════

    def detect_text_genre_and_tone(self, text: str) -> Dict[str, Any]:
        """
        كشف متحسن لنوع النص ونبرته العاطفية.

        المنهج:
        • ثلاث عينات: بداية + وسط + نهاية (لا أول 1000 حرف فقط)
        • نقاط موزونة لكل نوع (لا if/elif بأولوية ثابتة)
        • كشف سياقي للنبرة يتجنب الأخطاء كـ"love" في نص حزين
        • دعم النوع الهجين عبر genre_scores

        يُعيد dict بالمفاتيح المتوافقة مع الكود القديم: genre, tone
        ويُضيف: genre_scores, tone_scores للتشخيص
        """
        total = len(text)
        # عينات من ثلاثة مواضع
        samples = [
            text[:min(1200, total)],
            text[max(0, total // 2 - 600): min(total, total // 2 + 600)],
            text[max(0, total - 1200):]
        ]
        sample = ' '.join(samples).lower()

        # ── نقاط النوع ──
        genre_scores: Dict[str, int] = {
            'poetry':    0,
            'drama':     0,
            'narrative': 0,
            'prose':     2,   # وزن افتراضي
        }

        _poetry_signals = [
            'poem', 'verse', 'stanza', 'rhyme', 'poetry', 'sonnet', 'lyric',
            'ode', 'ballad', 'haiku', 'couplet'
        ]
        _drama_signals = [
            'dialogue', 'scene ', 'act ', ' said,', ' said.', '"—', 'exclaimed',
            'whispered', 'replied', 'shouted', 'muttered', 'cried out'
        ]
        _narrative_signals = [
            'chapter', 'story', 'novel', 'tale', 'narrator', 'once upon',
            'he walked', 'she ran', 'he said', 'she said', 'he thought',
            'she felt', 'he looked', 'she looked'
        ]

        for sig in _poetry_signals:
            if sig in sample:
                genre_scores['poetry'] += 2

        for sig in _drama_signals:
            if sig in sample:
                genre_scores['drama'] += 1

        for sig in _narrative_signals:
            if sig in sample:
                genre_scores['narrative'] += 1

        # علامات الحوار الفعلية (اقتباسات كثيرة = دراما/سرد)
        quote_count = sample.count('"') + sample.count('\u201c') + sample.count('\u2018')
        if quote_count > 8:
            genre_scores['drama'] += 2
        if quote_count > 4:
            genre_scores['narrative'] += 1

        genre = max(genre_scores, key=genre_scores.get)

        # ── نقاط النبرة ──
        tone_scores: Dict[str, int] = {
            'melancholic': 0,
            'joyful':      0,
            'dramatic':    0,
            'neutral':     1,   # وزن افتراضي
        }

        _sad_signals: Dict[str, int] = {
            'died': 3, 'death': 2, 'grief': 3, 'mourning': 3, 'tears': 2,
            'sorrow': 3, 'tragic': 3, 'hopeless': 2, 'despair': 3, 'funeral': 3,
            'lost': 1, 'pain': 1, 'hurt': 1, 'sad': 2, 'melancholy': 3,
            'weeping': 2, 'grave': 2, 'burial': 3, 'widow': 2
        }
        _happy_signals: Dict[str, int] = {
            'joy': 2, 'happy': 2, 'celebration': 3, 'triumph': 2, 'success': 1,
            'smile': 1, 'laugh': 1, 'wonderful': 1, 'beautiful': 1,
            'wedding': 2, 'birthday': 2, 'victory': 2, 'delight': 2
        }
        _dramatic_signals: Dict[str, int] = {
            'conflict': 2, 'tension': 2, 'crisis': 2, 'climax': 2,
            'struggle': 2, 'fight': 1, 'war': 2, 'battle': 2, 'danger': 2,
            'escape': 1, 'chase': 1, 'explosion': 2, 'confrontation': 2
        }

        for word, weight in _sad_signals.items():
            if word in sample:
                tone_scores['melancholic'] += weight

        for word, weight in _happy_signals.items():
            if word in sample:
                tone_scores['joyful'] += weight

        for word, weight in _dramatic_signals.items():
            if word in sample:
                tone_scores['dramatic'] += weight

        # تصحيح سياقي: "love" في نص حزين لا يعني الفرح
        if 'love' in sample and tone_scores['melancholic'] > 3:
            tone_scores['joyful'] = max(0, tone_scores['joyful'] - 1)

        tone = max(tone_scores, key=tone_scores.get)

        return {
            "genre":        genre,
            "tone":         tone,
            "genre_scores": genre_scores,
            "tone_scores":  tone_scores,
        }

    # ════════════════════════════════════════════════════════════
    #  ① بناء systemInstruction كحقل مستقل
    # ════════════════════════════════════════════════════════════

    def _build_system_instruction(self, text_analysis: Dict[str, Any]) -> str:
        """
        يبني التعليمات الثابتة التي تُمرَّر كـ systemInstruction منفصل لـ Gemini.
        الحقل المنفصل يحظى بوزن انتباه ثابت عالٍ ولا يُطغى عليه المحتوى.

        يحتوي:
        • هوية المترجم وقاعدة واحدة (لا قائمة من 7 طلبات متوازية)
        • معلومات الكتاب
        • المصطلحات مرتبة حسب الأهمية (تردد عالٍ أولاً)
        • ② ملف معرفة الكتاب (شخصيات + أماكن)
        • ⑤ أمثلة Few-Shot بعد الفصلين الأولين
        • توجيه النوع والنبرة
        """
        parts: List[str] = []

        # ── الهوية والقاعدة الأساسية ──
        parts.append(
            "أنت خبير ترجمة أدبية محترف. مهمتك: ترجمة كاملة وأمينة من الإنجليزية "
            "إلى العربية الفصحى. قاعدة ذهبية واحدة: لا تترك أي كلمة أو جملة دون ترجمة. "
            "أخرج النص المترجم فقط بدون تعليقات أو إضافات."
        )

        # ── معلومات الكتاب ──
        if self.book_knowledge['title']:
            book_info = f"الكتاب: {self.book_knowledge['title']}"
            if self.book_knowledge['author']:
                book_info += f" — {self.book_knowledge['author']}"
            parts.append(book_info)

        # ── توجيه النوع الأدبي والنبرة ──
        genre_guidance = self._get_genre_specific_guidance(text_analysis)
        if genre_guidance.strip():
            parts.append(genre_guidance)

        # ── ⑥ المصطلحات مرتبة حسب الأهمية (الأكثر تكراراً أولاً) ──
        high_priority_terms = self._get_high_priority_terms(n=30)
        if high_priority_terms:
            term_lines = ["المصطلحات والأسماء الثابتة الإلزامية:"]
            for orig, trans in high_priority_terms:
                term_lines.append(f"  {orig} ← {trans}")
            parts.append('\n'.join(term_lines))

        # ── ② معرفة الكتاب (شخصيات + أماكن) ──
        knowledge_section = self._build_book_knowledge_section()
        if knowledge_section:
            parts.append(knowledge_section)

        # ── ⑤ Few-Shot Examples (بعد الفصلين الأولين) ──
        if self.book_knowledge['few_shot_examples']:
            ex_lines = ["أمثلة على أسلوب ترجمة هذا الكتاب تحديداً:"]
            for ex in self.book_knowledge['few_shot_examples'][:3]:
                ex_lines.append(f"  الأصل: {ex['original']}")
                ex_lines.append(f"  الترجمة: {ex['translation']}")
            parts.append('\n'.join(ex_lines))

        return '\n\n'.join(p for p in parts if p.strip())

    def _build_book_knowledge_section(self) -> str:
        """② قسم معرفة الكتاب: شخصيات وأماكن معتمدة"""
        sections: List[str] = []

        if self.book_knowledge['characters']:
            char_lines = ["الشخصيات المعتمدة:"]
            for eng, info in list(self.book_knowledge['characters'].items())[:20]:
                line = f"  {eng} ← {info['arabic']}"
                if info.get('description'):
                    line += f" ({info['description']})"
                char_lines.append(line)
            sections.append('\n'.join(char_lines))

        if self.book_knowledge['places']:
            place_lines = ["الأماكن المعتمدة:"]
            for eng, arabic in list(self.book_knowledge['places'].items())[:10]:
                place_lines.append(f"  {eng} ← {arabic}")
            sections.append('\n'.join(place_lines))

        return '\n\n'.join(sections)

    def _get_high_priority_terms(self, n: int = 30) -> List[Tuple[str, str]]:
        """⑥ المصطلحات مرتبة حسب التردد (الأكثر استخداماً = الأهم)"""
        sorted_terms = sorted(
            [(orig, trans) for orig, trans in self.terminology_database.items()],
            key=lambda x: self.term_frequency.get(x[0], 1),
            reverse=True
        )
        return sorted_terms[:n]

    # ════════════════════════════════════════════════════════════
    #  ① create_complete_translation_prompt → يُعيد tuple
    # ════════════════════════════════════════════════════════════

    def create_complete_translation_prompt(
        self,
        text: str,
        context: str = "",
        text_analysis: Dict[str, Any] = None
    ) -> Tuple[str, str]:
        """
        يُعيد (system_instruction, user_prompt) منفصلين.
        system_instruction → حقل systemInstruction في Gemini API (وزن ثابت عالٍ).
        user_prompt        → محتوى الرسالة فقط: السياق المباشر + النص.

        التغيير عن النسخة القديمة:
        • لا دمج للتعليمات مع المحتوى في حقل واحد
        • الـ prompt نظيف ومركّز: سياق + نص فقط
        • ⑥ طلب JSON المصطلحات مُضمَّن في نفس الطلب (لا API call منفصل)
        """
        if not text_analysis:
            text_analysis = self.detect_text_genre_and_tone(text)

        # system_instruction: كل ما هو ثابت
        system_instruction = self._build_system_instruction(text_analysis)

        # ── السياق المباشر (ملخصات الأحداث لا مقتطعات نصية) ──
        context_section = ""
        if self.book_knowledge['events']:
            recent = self.book_knowledge['events'][-3:]
            ev_lines = ["ملخص الفصول السابقة:"]
            for ev in recent:
                ev_lines.append(f"  - {ev['chapter']}: {ev['summary']}")
            context_section = '\n'.join(ev_lines)
        elif context:
            context_section = f"السياق: {context[:300]}"

        # ── ⑥ طلب JSON المصطلحات ضمن نفس الرد ──
        json_hint = (
            "\n\n---\n"
            "إذا وجدت أسماء أو مصطلحات مهمة جديدة في هذا النص، أضف في نهاية ردك:\n"
            'TERMS_JSON:{"terms":{"EnglishName":"الاسم العربي"}}\n'
            "إذا لم تجد مصطلحات جديدة، لا تُضف هذا القسم."
        )

        user_prompt = ""
        if context_section:
            user_prompt += context_section + "\n\n"

        user_prompt += f'النص المطلوب ترجمته:\n"""\n{text}\n"""{json_hint}'

        return system_instruction, user_prompt.strip()

    def _get_genre_specific_guidance(self, text_analysis: Dict[str, Any]) -> str:
        """توجيهات مختصرة حسب نوع النص ونبرته"""
        genre_guides = {
            "poetry": (
                "النوع: شعر — احتفظ بالإيقاع والصور الشعرية، "
                "لا تفقد أي بيت أو مقطع."
            ),
            "drama": (
                "النوع: حوار/مسرح — اجعل الحوار طبيعياً ومعبراً، "
                "اترجم كل إرشاد مسرحي."
            ),
            "narrative": (
                "النوع: سرد روائي — حافظ على تدفق الحكاية وتسلسل الأحداث، "
                "لا تفوت أي تفصيل."
            ),
            "prose": (
                "النوع: نثر — اجعل النثر متدفقاً وسليماً لغوياً، "
                "حافظ على ترابط الأفكار."
            ),
        }
        tone_guides = {
            "melancholic": "النبرة: حزينة/كآبة — انقل العمق العاطفي بدقة وحساسية.",
            "joyful":      "النبرة: مرحة/سعيدة — استخدم تعابير مشرقة ومفرحة.",
            "dramatic":    "النبرة: درامية — حافظ على التوتر والإثارة كما في الأصل.",
            "neutral":     "النبرة: محايدة — حافظ على التوازن والوضوح.",
        }
        g = genre_guides.get(text_analysis.get('genre', 'prose'), genre_guides['prose'])
        t = tone_guides.get(text_analysis.get('tone', 'neutral'), tone_guides['neutral'])
        return f"{g}\n{t}"

    # ════════════════════════════════════════════════════════════
    #  ② تحديث ملف معرفة الكتاب بعد كل فصل
    # ════════════════════════════════════════════════════════════

    def _update_book_knowledge(
        self, chapter_title: str, original: str, translated: str
    ):
        """
        يُحدّث ملف معرفة الكتاب بعد إتمام ترجمة كل فصل:
        • يستخرج الكيانات المسماة (spaCy) ويربطها بالترجمة المعتمدة
        • يُضيف ملخصاً للأحداث (3 جمل من بداية الترجمة)
        • يجمع ⑤ أمثلة Few-Shot من الفصلين الأولين
        """
        # ── استخراج كيانات spaCy وربطها بالمصطلحات ──
        entities = self.content_processor._extract_named_entities_spacy(original)
        for entity in entities:
            if (entity not in self.book_knowledge['characters']
                    and entity not in self.book_knowledge['places']):
                arabic_name = self.terminology_database.get(
                    entity, self.terminology_database.get(entity.lower(), entity)
                )
                # تصنيف بسيط: كلمة واحدة كبيرة → شخصية، متعددة → مكان
                if ' ' not in entity.strip():
                    self.book_knowledge['characters'][entity] = {
                        'arabic': arabic_name,
                        'description': ''
                    }
                else:
                    self.book_knowledge['places'][entity] = arabic_name

        # ── ملخص أحداث الفصل (من بداية الترجمة) ──
        summary_raw = translated[:300].strip()
        # اقتطع عند نهاية جملة
        for sep in ['؟', '!', '.', '،']:
            idx = summary_raw.rfind(sep)
            if idx > 80:
                summary_raw = summary_raw[:idx + 1]
                break
        self.book_knowledge['events'].append({
            'chapter': chapter_title,
            'summary': summary_raw
        })

        # ── ⑤ Few-Shot: جمع أمثلة من الفصلين الأولين ──
        if (self.chapters_completed <= 2
                and len(self.book_knowledge['few_shot_examples']) < 3):
            orig_sents = re.split(r'(?<=[.!?])\s+', original)
            trans_sents = re.split(r'(?<=[.!?؟!])\s+', translated)
            # نختار جملة من المنتصف (أوضح من البداية)
            pick_idx = min(len(orig_sents), len(trans_sents), 4) - 1
            if pick_idx >= 0:
                ex_orig = orig_sents[pick_idx].strip()
                ex_trans = trans_sents[pick_idx].strip()
                if 20 < len(ex_orig) < 180 and 20 < len(ex_trans) < 220:
                    self.book_knowledge['few_shot_examples'].append({
                        'original': ex_orig,
                        'translation': ex_trans
                    })

    # ════════════════════════════════════════════════════════════
    #  ⑥ استخراج المصطلحات من رد الترجمة (بدون API منفصل)
    # ════════════════════════════════════════════════════════════

    @staticmethod
    def _extract_first_json_object(text: str) -> str:
        """
        يستخرج أول كائن JSON متوازن الأقواس من النص.

        لماذا هذه الطريقة؟
        الـ regex البسيط يرفض أي أقواس داخلية،
        لذا يفشل مع JSON المتداخل مثل {"terms":{"key":"val"}}.
        هذه الدالة تتتبع عمق الأقواس بدقة مع دعم السلاسل النصية
        والأحرف المهرّبة لضمان الاستخراج الصحيح دائماً.
        """
        start = text.find('{')
        if start == -1:
            return ''
        depth = 0
        in_string = False
        escape_next = False
        for i, ch in enumerate(text[start:], start):
            if escape_next:
                escape_next = False
                continue
            if ch == '\\' and in_string:
                escape_next = True
                continue
            if ch == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    return text[start: i + 1]
        return ''

    @staticmethod
    def _safe_json_load(text: str) -> dict:
        """
        استخراج JSON آمن بمحاولتين:
        ① _extract_first_json_object (يتعامل مع JSON متداخل وعلى يمينه نص)
        ② json.loads مباشرة كـ fallback
        """
        candidate = CompleteTranslationEngine._extract_first_json_object(text)
        if candidate:
            try:
                return json.loads(candidate)
            except Exception:
                pass
        try:
            return json.loads(text.strip())
        except Exception:
            return {}

    def _parse_terms_from_response(self, response: str) -> Tuple[str, int]:
        """
        يفصل النص المترجم عن قسم JSON المصطلحات في نفس الرد.
        يُعيد (cleaned_translation, terms_count).

        الإصلاح:
        ─────────
        يستخدم _extract_first_json_object بدلاً من regex [^{}]*
        الذي كان يفشل مع أي JSON متداخل → لم يُستخرج أي مصطلح أبداً.

        مصطلحات مكررة (تردد عالٍ) → إلزامية في system instruction.
        مصطلحات لمرة واحدة → تُسجَّل بثقة منخفضة.
        """
        if not response or 'TERMS_JSON:' not in response:
            return response, 0

        parts = response.split('TERMS_JSON:', 1)
        translation_clean = parts[0].strip()
        terms_raw = parts[1].strip()

        terms_count = 0
        try:
            data = self._safe_json_load(terms_raw)
            terms = data.get('terms', {})
            for english, arabic in terms.items():
                english = str(english).strip()
                arabic  = str(arabic).strip()
                if english and arabic and len(english) > 2:
                    self.terminology_database[english] = arabic
                    self.term_frequency[english] = (
                        self.term_frequency.get(english, 0) + 1
                    )
                    terms_count += 1
                    # تحديث ملف معرفة الكتاب
                    if english and (english.istitle() or english[0].isupper()):
                        if ' ' not in english.strip():
                            if english not in self.book_knowledge['characters']:
                                self.book_knowledge['characters'][english] = {
                                    'arabic': arabic, 'description': ''
                                }
                        else:
                            if english not in self.book_knowledge['places']:
                                self.book_knowledge['places'][english] = arabic
        except Exception as e:
            logger.debug(f"[Terms] JSON parse warning: {e}")

        if terms_count:
            logger.info(
                f"[Terms] ⑥ Extracted {terms_count} terms inline "
                f"(no extra API call) | total_db={len(self.terminology_database)}"
            )

        return translation_clean, terms_count

    # ════════════════════════════════════════════════════════════
    #  ④ سلسلة المراجعة المستهدفة (محلي → API فقط عند الحاجة)
    # ════════════════════════════════════════════════════════════

    def _find_problematic_sentences(
        self, translation: str, original: str
    ) -> List[Dict[str, Any]]:
        """
        ④ المرحلة 2 (محلي - بدون API):
        يكشف الجمل التي تحتوي كلمات إنجليزية غير مقبولة باستخدام langdetect
        ويُرقّمها لإرسالها فقط إلى الـ API في المرحلة 3.
        يُقلّل حجم طلبات المراجعة بنسبة 80-90%.
        """
        # بناء قائمة بيضاء مؤقتة من النص الأصلي
        orig_whitelist = set(
            w for w in re.findall(r'\b[A-Za-z]{2,}\b', original)
        )

        problematic: List[Dict[str, Any]] = []
        sentences = re.split(r'(?<=[.!?؟!])\s+', translation)

        for i, sent in enumerate(sentences):
            eng_words = re.findall(r'\b[A-Za-z]{3,}\b', sent)
            bad_words = []
            for w in eng_words:
                if (w in orig_whitelist or w.lower() in orig_whitelist):
                    continue
                if self.content_processor._is_acceptable_english(w):
                    continue
                bad_words.append(w)
            if bad_words:
                problematic.append({
                    'index':    i,
                    'sentence': sent,
                    'issues':   bad_words
                })

        return problematic

    async def _targeted_correction(
        self,
        problematic: List[Dict[str, Any]],
        full_translation: str,
        system_instruction: str
    ) -> str:
        """
        ④ المرحلة 3 (API فقط عند الحاجة):
        يرسل الجمل المشكلة فقط (لا النص الكامل) للنموذج للتصحيح.
        المرحلة 4: يدمج الإصلاحات في النص الكامل.

        الإصلاحات:
        ──────────
        Bug #2: مثال الـ JSON الآن يستخدم الـ indices الحقيقية (لا 0,1,2 الوهمية)
                حتى يفهم النموذج أن المطلوب keys بالأرقام الفعلية للجمل.
        Bug #3: استخدام _safe_json_load بدلاً من regex جشع مع DOTALL
                الذي كان يفشل عند وجود أقواس عربية بعد الـ JSON.
        """
        if not problematic:
            return full_translation

        batch = problematic[:12]   # حد أقصى 12 جملة في طلب واحد

        # بناء مثال JSON ديناميكي من الـ indices الحقيقية (Bug #2 Fix)
        example_keys = {str(p['index']): "الجملة المصححة هنا" for p in batch[:2]}
        example_json  = json.dumps({"corrections": example_keys}, ensure_ascii=False)

        req_lines = [
            "صحح الجمل التالية: ترجم الكلمات الإنجليزية المُشار إليها إلى العربية.",
            f"أجب بـ JSON فقط بهذا الشكل: {example_json}",
            ""
        ]
        for item in batch:
            req_lines.append(
                f"جملة_{item['index']}: {item['sentence']}\n"
                f"  كلمات تحتاج ترجمة: {', '.join(item['issues'])}"
            )
        req_lines.append(
            f'\nالرد (JSON فقط بـ keys هي أرقام الجمل أعلاه):'
        )

        targeted_prompt = '\n'.join(req_lines)
        targeted_sys = (
            "أنت مصحح ترجمة دقيق. مهمتك الوحيدة: ترجمة الكلمات الإنجليزية المُحددة "
            "وإعادة الجملة كاملة بالعربية. أجب بـ JSON فقط، "
            "لا تكتب أي نص قبل الـ JSON أو بعده."
        )

        result = await self.api_manager.make_precision_request(
            targeted_prompt,
            system_instruction=targeted_sys,
            temperature=0.05,
            request_type="targeted_correction"
        )

        response_text, _, _ = result
        if not response_text:
            return full_translation

        # ── المرحلة 4: دمج التصحيحات (Bug #3 Fix: safe JSON extraction) ──
        try:
            parsed = self._safe_json_load(response_text)
            corrections: Dict[str, str] = parsed.get('corrections', {})

            if not corrections:
                logger.warning(
                    "[TargetedCorrection] No corrections found in response, "
                    "returning original"
                )
                return full_translation

            sentences = re.split(r'(?<=[.!?؟!])\s+', full_translation)
            applied = 0
            for item in batch:
                key = str(item['index'])
                if key in corrections and item['index'] < len(sentences):
                    corr = corrections[key].strip()
                    if corr:  # لا نستبدل بجملة فارغة
                        sentences[item['index']] = corr
                        applied += 1

            merged = ' '.join(sentences)
            logger.info(
                f"[TargetedCorrection] ④ Applied {applied}/{len(batch)} targeted fixes "
                f"(sent only {len(batch)} sentences vs full text)"
            )
            return merged

        except Exception as e:
            logger.warning(f"[TargetedCorrection] Merge failed: {e}")
            return full_translation

    # ════════════════════════════════════════════════════════════
    #  تقسيم النص إلى قطع آمنة
    # ════════════════════════════════════════════════════════════

    def _split_into_safe_chunks(self, text: str) -> List[str]:
        """
        تقسيم النص إلى قطع آمنة لا تتجاوز SAFE_CHUNK_WORDS كلمة.
        خوارزمية: فقرات أولاً، ثم جمل عند الضرورة، مع جسر سياقي بين القطع.
        """
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        if not paragraphs:
            return [text] if text.strip() else []

        chunks: List[str] = []
        current_parts: List[str] = []
        current_words = 0

        for para in paragraphs:
            para_words = len(para.split())

            if para_words > self.SAFE_CHUNK_WORDS:
                if current_parts:
                    chunks.append('\n\n'.join(current_parts))
                    current_parts = []
                    current_words = 0
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sub_parts: List[str] = []
                sub_words = 0
                for sent in sentences:
                    sw = len(sent.split())
                    if sub_words + sw > self.SAFE_CHUNK_WORDS and sub_parts:
                        chunks.append(' '.join(sub_parts))
                        sub_parts = [sent]
                        sub_words = sw
                    else:
                        sub_parts.append(sent)
                        sub_words += sw
                if sub_parts:
                    current_parts = [' '.join(sub_parts)]
                    current_words = sub_words
                continue

            if current_words + para_words > self.SAFE_CHUNK_WORDS and current_parts:
                chunks.append('\n\n'.join(current_parts))
                current_parts = [para]
                current_words = para_words
            else:
                current_parts.append(para)
                current_words += para_words

        if current_parts:
            chunks.append('\n\n'.join(current_parts))

        return chunks if chunks else [text]

    async def _translate_in_chunks(
        self, text: str, context: str, text_analysis: Dict[str, Any]
    ) -> Optional[str]:
        """
        ترجمة نص كبير قطعةً قطعة.
        • ① كل قطعة تُرسَل مع system_instruction منفصل
        • جسر سياقي: آخر 300 حرف من الترجمة السابقة
        • كشف ###TRUNCATED### مع إعادة محاولة بقطعة أصغر
        """
        chunks = self._split_into_safe_chunks(text)
        total = len(chunks)
        logger.info(
            f"[Chunking] Text split into {total} safe chunks "
            f"(max {self.SAFE_CHUNK_WORDS} words each)"
        )

        translated_chunks: List[str] = []
        running_context = context

        for idx, chunk in enumerate(chunks):
            chunk_words = len(chunk.split())
            logger.info(
                f"[Chunking] Translating chunk {idx+1}/{total} ({chunk_words} words)..."
            )

            # ① system_instruction + user_prompt منفصلان
            sys_inst, user_prompt = self.create_complete_translation_prompt(
                chunk, running_context, text_analysis
            )
            result = await self.api_manager.make_precision_request(
                user_prompt,
                system_instruction=sys_inst,
                temperature=0.1,
                request_type="chunk_translation"
            )

            chunk_translation, r_time, a_key = result if result else (None, 0.0, None)

            # تهيئة sub_results قبل الكتلة الشرطية لتجنب UnboundLocalError
            # عند الحالة العادية (بدون truncation) تبقى قائمة فارغة
            sub_results: List[str] = []

            # ── كشف القطع المُقطوعة ──
            if chunk_translation and "###TRUNCATED###" in chunk_translation:
                logger.warning(
                    f"[Chunking] Chunk {idx+1} truncated — retrying with half-size..."
                )
                chunk_translation = chunk_translation.replace("###TRUNCATED###", "").strip()

                mid_words = chunk.split()
                half = len(mid_words) // 2
                for sub_chunk in [' '.join(mid_words[:half]), ' '.join(mid_words[half:])]:
                    s_sys, s_prompt = self.create_complete_translation_prompt(
                        sub_chunk, running_context, text_analysis
                    )
                    sub_result = await self.api_manager.make_precision_request(
                        s_prompt,
                        system_instruction=s_sys,
                        temperature=0.1,
                        request_type="subchunk_translation"
                    )
                    sub_trans, _, _ = sub_result if sub_result else (None, 0.0, None)
                    if sub_trans:
                        # ⑥ استخراج المصطلحات هنا — لا نُكرر الاستخراج بعدها
                        sub_trans_clean, _ = self._parse_terms_from_response(
                            sub_trans.replace("###TRUNCATED###", "").strip()
                        )
                        sub_results.append(sub_trans_clean)
                        running_context = sub_trans_clean[-300:]

                if sub_results:
                    # المصطلحات استُخرجت بالفعل في الحلقة أعلاه
                    chunk_translation = '\n\n'.join(sub_results)

            if chunk_translation:
                # ⑥ استخراج المصطلحات من الرد (فقط إذا لم تُعالَج كـ sub_results أعلاه)
                # sub_results جاهزة ونظيفة بالفعل؛ هذا المسار للحالة العادية (بدون truncation)
                clean_chunk = chunk_translation.replace("###TRUNCATED###", "").strip()
                if sub_results:
                    # المصطلحات استُخرجت في حلقة sub_results — نتجنب الاستدعاء المكرر
                    clean = clean_chunk
                else:
                    clean, _ = self._parse_terms_from_response(clean_chunk)
                translated_chunks.append(clean)
                # تحديث الجسر السياقي: آخر 300 حرف
                running_context = clean[-300:]
                logger.info(f"[Chunking] Chunk {idx+1}/{total} done ✓")
            else:
                logger.error(f"[Chunking] ❌ Chunk {idx+1}/{total} FAILED — "
                             f"original text will be inserted as placeholder")
                # أدخل النص الأصلي كـ placeholder لتجنب فقدان البيانات
                translated_chunks.append(f"\n[ترجمة مفقودة - النص الأصلي:\n{chunk}\n]")

        return '\n\n'.join(translated_chunks)

    async def translate_with_completion_guarantee(
        self, text: str, context: str = ""
    ) -> Tuple[Optional[str], float, Optional[str]]:
        """
        ترجمة مع ضمان الاكتمال الشامل.
        سلسلة العمل المحسنة:
        ① ترجمة أولية (API) مع system_instruction منفصل
        ② كشف الجمل المشكلة محلياً (بدون API)
        ③ تصحيح مستهدف للجمل المشكلة فقط (API عند الحاجة فقط)
        ④ تحديث ملف معرفة الكتاب
        """
        logger.info(
            f"[CompletionGuarantee] Starting for text of {len(text)} chars"
        )

        text_analysis = self.detect_text_genre_and_tone(text)
        logger.info(
            f"Text analysis: Genre={text_analysis['genre']}, Tone={text_analysis['tone']}"
        )

        word_count = len(text.split())

        # ── المرحلة 1: ترجمة أولية ──
        response_time = 0.0
        api_key_used = None

        if word_count > self.SAFE_CHUNK_WORDS:
            logger.info(
                f"[Chunking] {word_count} words > {self.SAFE_CHUNK_WORDS} → chunk-based"
            )
            initial_translation = await self._translate_in_chunks(
                text, context, text_analysis
            )
        else:
            sys_inst, user_prompt = self.create_complete_translation_prompt(
                text, context, text_analysis
            )
            result = await self.api_manager.make_precision_request(
                user_prompt,
                system_instruction=sys_inst,
                temperature=0.1,
                request_type="complete_translation"
            )
            initial_translation, response_time, api_key_used = (
                result if result else (None, 0.0, None)
            )
            if initial_translation:
                initial_translation, _ = self._parse_terms_from_response(
                    initial_translation.replace("###TRUNCATED###", "").strip()
                )

        if not initial_translation:
            logger.error("[CompletionGuarantee] Initial translation failed")
            return None, 0.0, None

        logger.info("[CompletionGuarantee] Initial done — checking coverage...")

        # ── المرحلة 2: كشف الجمل المشكلة محلياً ──
        issues = self.content_processor.detect_incomplete_translation(
            text, initial_translation
        )
        problematic_sents = self._find_problematic_sentences(initial_translation, text)

        # ── المرحلة 3: مراجعة مستهدفة فقط عند وجود مشاكل ──
        if problematic_sents or issues['coverage_percentage'] < 65.0:
            quality_logger.warning(
                f"[CompletionGuarantee] Issues: {len(problematic_sents)} bad sentences, "
                f"coverage={issues['coverage_percentage']:.1f}%"
            )

            if problematic_sents:
                # تصحيح الجمل المشكلة فقط (80-90% أقل توكنز)
                sys_inst_short = self._build_system_instruction(text_analysis)
                fixed = await self._targeted_correction(
                    problematic_sents, initial_translation, sys_inst_short
                )
            else:
                fixed = initial_translation

            # إذا التغطية لا تزال منخفضة → مراجعة للاكتمال فقط
            final_check = self.content_processor.detect_incomplete_translation(
                text, fixed
            )
            if final_check['coverage_percentage'] < 65.0:
                quality_logger.warning(
                    "[CompletionGuarantee] Coverage still low → completion review..."
                )
                sys_review = (
                    "أنت مراجع ترجمة. مهمتك: ترجمة الأجزاء الناقصة فقط دون تغيير ما هو صحيح. "
                    "أخرج الترجمة الكاملة فقط."
                )
                review_prompt = (
                    f"الترجمة الحالية ناقصة (تغطية {final_check['coverage_percentage']:.0f}%).\n\n"
                    f"النص الأصلي:\n\"\"\"\n{text}\n\"\"\"\n\n"
                    f"الترجمة الحالية:\n\"\"\"\n{fixed}\n\"\"\"\n\n"
                    "قدم الترجمة الكاملة:"
                )
                review_result = await self.api_manager.make_precision_request(
                    review_prompt,
                    system_instruction=sys_review,
                    temperature=0.05,
                    request_type="completion_review"
                )
                reviewed, r_t, r_k = review_result
                if reviewed:
                    if r_t: response_time += r_t
                    if r_k: api_key_used = r_k
                    reviewed, _ = self._parse_terms_from_response(reviewed)
                    final_translation = reviewed if reviewed else fixed
                else:
                    final_translation = fixed
            else:
                final_translation = fixed
                logger.info(
                    f"[CompletionGuarantee] ✅ Coverage {final_check['coverage_percentage']:.1f}% after targeted fix"
                )
        else:
            logger.info(
                f"[CompletionGuarantee] ✅ Initial translation clean "
                f"(coverage={issues['coverage_percentage']:.1f}%)"
            )
            final_translation = initial_translation

        # ── تحويل الأرقام ──
        final_translation = self.content_processor.convert_numbers_to_arabic(
            final_translation
        )

        # ── المرحلة 4: تحديث ملف المعرفة ──
        if final_translation:
            self.context_history.append(final_translation[:500])
            if len(self.context_history) > 5:
                self.context_history.pop(0)
            self._update_book_knowledge(context or "فصل", text, final_translation)
            self.chapters_completed += 1

        return final_translation, response_time, api_key_used

    async def translate_with_comprehensive_review(
        self, text: str, context: str = ""
    ) -> Tuple[Optional[str], float, Optional[str]]:
        """
        ترجمة شاملة مع مراجعة مستهدفة لضمان عدم ترك أي محتوى أجنبي.

        سلسلة العمل المحسنة بالنقاط الست:
        ① الترجمة الأولية مع systemInstruction منفصل
        ② كشف المحتوى الأجنبي محلياً بدون API
        ③ تصحيح مستهدف للجمل المشكلة فقط (80-90% أقل توكنز)
        ④ تحديث ملف معرفة الكتاب (شخصيات + أحداث + أسلوب)
        ⑤ Few-Shot يُفعَّل تلقائياً بعد الفصلين الأولين
        ⑥ المصطلحات تُستخرج من نفس الرد (لا API call منفصل)
        """
        logger.info(
            f"[ComprehensiveReview] Starting for text of {len(text)} chars"
        )

        # المرحلة 1: تحليل النص ③
        text_analysis = self.detect_text_genre_and_tone(text)
        logger.info(
            f"Text analysis: Genre={text_analysis['genre']}, "
            f"Tone={text_analysis['tone']} | "
            f"genre_scores={text_analysis.get('genre_scores', {})}"
        )

        word_count = len(text.split())
        response_time = 0.0
        api_key_used = None

        # المرحلة 2: ترجمة أولية ①
        if word_count > self.SAFE_CHUNK_WORDS:
            logger.info(
                f"[Chunking] {word_count} words > {self.SAFE_CHUNK_WORDS} "
                "→ chunk-based translation"
            )
            initial_translation = await self._translate_in_chunks(
                text, context, text_analysis
            )
        else:
            sys_inst, user_prompt = self.create_complete_translation_prompt(
                text, context, text_analysis
            )
            result = await self.api_manager.make_precision_request(
                user_prompt,
                system_instruction=sys_inst,
                temperature=0.1,
                request_type="contextual_translation"
            )
            initial_translation, response_time, api_key_used = (
                result if result else (None, 0.0, None)
            )
            if initial_translation:
                # ⑥ استخراج المصطلحات من نفس الرد
                initial_translation, _terms_count = self._parse_terms_from_response(
                    initial_translation.replace("###TRUNCATED###", "").strip()
                )

        if not initial_translation:
            logger.error("[ComprehensiveReview] Initial translation failed")
            return None, 0.0, None

        logger.info("[ComprehensiveReview] Initial done — checking foreign content...")

        # المرحلة 3: ② كشف المحتوى الأجنبي محلياً
        has_foreign = self.content_processor.has_any_foreign_content(
            initial_translation, original_text=text
        )
        problematic_sents = (
            self._find_problematic_sentences(initial_translation, text)
            if has_foreign else []
        )

        if problematic_sents:
            quality_logger.warning(
                f"[ComprehensiveReview] {len(problematic_sents)} sentences with foreign content"
            )
            # المرحلة 3: ③ تصحيح مستهدف للجمل المشكلة فقط
            sys_inst_for_fix = self._build_system_instruction(text_analysis)
            fixed = await self._targeted_correction(
                problematic_sents, initial_translation, sys_inst_for_fix
            )

            # فحص بعد التصحيح
            still_has_foreign = self.content_processor.has_any_foreign_content(
                fixed, original_text=text
            )
            if still_has_foreign:
                quality_logger.warning(
                    "[ComprehensiveReview] Foreign content remains after targeted fix → final cleanup"
                )
                cleanup_sys = (
                    "أنت مصحح نهائي. احذف أو ترجم أي كلمة إنجليزية غير مقبولة. "
                    "أخرج النص العربي فقط."
                )
                cleanup_prompt = (
                    f"صحّح هذه الترجمة من أي كلمات إنجليزية متبقية:\n\n{fixed}"
                )
                cleanup_result = await self.api_manager.make_precision_request(
                    cleanup_prompt,
                    system_instruction=cleanup_sys,
                    temperature=0.02,
                    request_type="final_cleanup"
                )
                cleaned, r_t, r_k = cleanup_result
                if cleaned:
                    if r_t: response_time += r_t
                    if r_k: api_key_used = r_k
                    final_translation = cleaned
                else:
                    final_translation = fixed
            else:
                logger.info(
                    "[ComprehensiveReview] ✅ Clean after targeted correction"
                )
                final_translation = fixed
        else:
            logger.info(
                "[ComprehensiveReview] ✅ Initial translation free of foreign content"
            )
            final_translation = initial_translation

        # تحويل الأرقام
        final_translation = self.content_processor.convert_numbers_to_arabic(
            final_translation
        )

        # المرحلة 4: تحديث ملف المعرفة ② ⑤
        if final_translation:
            self.context_history.append(final_translation[:500])
            if len(self.context_history) > 5:
                self.context_history.pop(0)
            self._update_book_knowledge(context or "فصل", text, final_translation)
            self.chapters_completed += 1

        return final_translation, response_time, api_key_used

    async def extract_terminology(self, original: str, translation: str):
        """
        للتوافق مع الكود الخارجي.
        المصطلحات تُستخرج الآن تلقائياً ضمن رد الترجمة (⑥).
        هذه الدالة تبقى كواجهة عامة لكنها لا تُنشئ API call إضافياً.
        """
        # في حال استُدعيت مباشرة، نحاول استخراج المصطلحات محلياً
        # من خلال تحليل النصين بدون استدعاء API إضافي
        eng_words = re.findall(r'\b[A-Z][a-z]{2,}\b', original)
        for word in eng_words:
            if word in self.terminology_database:
                # زيادة التردد للمصطلح الموجود
                self.term_frequency[word] = self.term_frequency.get(word, 0) + 1
        logger.debug(
            "[extract_terminology] Called externally — terms already extracted inline ⑥"
        )


class ProfessionalDocumentProcessor:
    """
    معالج المستندات الاحترافي المطوّر - النسخة المطوّرة الشاملة

    التطويرات الجوهرية المُطبَّقة:
    ✅ الطبقة 1 : استخراج نصي بـ pymupdf مع البيانات الطباعية الكاملة
                 (حجم الخط، نوعه، موضعه، Bold/Italic)
    ✅ الطبقة 2 : OCR تلقائي بـ easyocr للصفحات المصوَّرة (< 50 حرفاً)
    ✅ الطبقة 3 : إشعار واضح ومفصَّل عند الملفات المحمية
    ✅ تراجع آمن لـ PyPDF2 عند غياب pymupdf تماماً
    ✅ قراءة TOC الداخلي للـ PDF مباشرةً (عند الوجود)
    ✅ كشف فصول ذكي يعتمد على حجم الخط والموضع والنمط الطباعي
    ✅ تقسيم نصي يحترم الوحدات السردية ونقاط القطع الطبيعية
    ✅ فحص سلامة الاستخراج مع إعادة محاولة تلقائية
    ✅ تنظيف نص محسَّن يُزيل الترويسات/التذييلات وأرقام الصفحات
    """

    # ─────────────────────────────────────────────────────────────────────
    #  ثوابت الكلاس
    # ─────────────────────────────────────────────────────────────────────
    _MIN_CHARS_PER_PAGE   : int   = 50      # أقل من هذا → صفحة مصوَّرة → OCR
    _MIN_CHAPTER_WORDS    : int   = 50      # فصل أقل من هذا غير منطقي
    _MAX_CHAPTER_WORDS    : int   = 20_000  # فصل أكثر من هذا مشبوه
    _FONT_SIZE_RATIO      : float = 1.25    # نسبة خط العنوان إلى خط النص العادي
    _SCENE_BREAK_PATTERNS : re.Pattern = re.compile(
        r'^(\*{3,}|\-{3,}|~{3,}|#{3,}|={3,}|\+{3,}|◆+|◇+|•{3,})$'
    )

    # ─────────────────────────────────────────────────────────────────────
    #  استيراد المكتبات الاختيارية
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _import_fitz() -> Optional[Any]:
        """استيراد pymupdf (fitz) مع رسالة إرشادية واضحة عند غيابه."""
        try:
            import fitz  # type: ignore
            return fitz
        except ImportError:
            logger.warning(
                "[ProfessionalDocumentProcessor] pymupdf غير مثبَّت. "
                "لتفعيل الاستخراج المتقدم نفِّذ: pip install pymupdf"
            )
            return None

    @staticmethod
    def _import_easyocr() -> Optional[Any]:
        """استيراد easyocr بصمت عند غيابه (مكتبة اختيارية)."""
        try:
            import easyocr  # type: ignore
            return easyocr
        except ImportError:
            return None

    # ─────────────────────────────────────────────────────────────────────
    #  الطبقة 1: الاستخراج المتقدم بـ pymupdf
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_layer1_fitz(file_path: str) -> Dict[str, Any]:
        """
        الطبقة 1: استخراج نصي كامل بـ pymupdf.
        يعيد قاموساً يتضمن:
          pages           – قائمة بيانات الصفحات (نص + بلوكات + عناوين مكتشفة)
          toc             – جدول محتويات داخلي للملف
          metadata        – بيانات الملف (العنوان، المؤلف …)
          total_pages     – عدد الصفحات الكلي
          is_protected    – هل الملف محمي بكلمة مرور؟
          needs_ocr_pages – أرقام الصفحات التي تحتاج OCR
        """
        fitz = ProfessionalDocumentProcessor._import_fitz()
        if fitz is None:
            return {"error": "pymupdf_missing"}

        result: Dict[str, Any] = {
            "pages":           [],
            "toc":             [],
            "metadata":        {},
            "total_pages":     0,
            "is_protected":    False,
            "needs_ocr_pages": [],
        }

        try:
            doc = fitz.open(file_path)
        except Exception as exc:
            logger.error(f"[Layer1] فشل فتح الملف: {exc}")
            return {"error": str(exc)}

        result["total_pages"] = len(doc)

        # ── الملف المحمي ──────────────────────────────────────────────
        if doc.needs_pass:
            result["is_protected"] = True
            doc.close()
            return result

        # ── الميتاداتا ────────────────────────────────────────────────
        meta = doc.metadata or {}
        result["metadata"] = {
            "title":    meta.get("title",    ""),
            "author":   meta.get("author",   ""),
            "subject":  meta.get("subject",  ""),
            "keywords": meta.get("keywords", ""),
            "creator":  meta.get("creator",  ""),
        }

        # ── TOC الداخلي ───────────────────────────────────────────────
        try:
            toc_raw = doc.get_toc(simple=False)  # [(level, title, page, dest), ...]
            result["toc"] = [
                {"level": item[0], "title": item[1], "page": item[2]}
                for item in toc_raw
                if item[1] and item[1].strip()
            ]
            if result["toc"]:
                logger.info(f"[Layer1] TOC داخلي: {len(result['toc'])} إدخال")
        except Exception:
            pass

        # ── استخراج الصفحات (حلقة واحدة تجمع الخطوط والبيانات معاً) ──
        # المرور الأول: جمع أحجام الخطوط لحساب المنوال
        all_font_sizes: List[float] = []
        # تخزين raw_dict لكل صفحة لتجنب القراءة المزدوجة
        pages_raw: List[Optional[Any]] = [None] * len(doc)

        try:
            for page_idx in range(len(doc)):
                try:
                    page_obj = doc[page_idx]
                    raw_dict = page_obj.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                    pages_raw[page_idx] = raw_dict
                    for blk in raw_dict.get("blocks", []):
                        if blk.get("type") != 0:
                            continue
                        for line in blk.get("lines", []):
                            for span in line.get("spans", []):
                                sz = span.get("size", 0)
                                if 6.0 <= sz <= 72.0:
                                    all_font_sizes.append(sz)
                except Exception:
                    pass

            body_font_size: float = 12.0
            if all_font_sizes:
                from collections import Counter
                size_counter = Counter(round(s, 1) for s in all_font_sizes)
                body_font_size = float(size_counter.most_common(1)[0][0])
            logger.info(f"[Layer1] حجم خط النص الأساسي: {body_font_size}pt")

            # المرور الثاني: بناء بيانات الصفحات من الـ raw_dict المحفوظة
            for page_idx, raw_dict in enumerate(pages_raw):
                if raw_dict is None:
                    continue
                try:
                    plain_text = ""
                    page_blocks: List[Dict[str, Any]] = []
                    potential_headings: List[str] = []

                    for blk in raw_dict.get("blocks", []):
                        if blk.get("type") != 0:
                            continue

                        block_text = ""
                        block_max_font: float = 0.0
                        is_bold = False

                        for ln in blk.get("lines", []):
                            line_text = ""
                            for span in ln.get("spans", []):
                                span_text  = span.get("text", "")
                                span_size  = span.get("size", 0.0)
                                span_flags = span.get("flags", 0)
                                line_text += span_text
                                if span_size > block_max_font:
                                    block_max_font = span_size
                                if span_flags & 16:   # bit 4 = Bold في pymupdf
                                    is_bold = True
                            block_text += line_text.strip() + "\n"

                        block_text = block_text.strip()
                        if not block_text:
                            continue

                        font_ratio = (block_max_font / body_font_size) if body_font_size > 0 else 1.0
                        is_heading = (
                            font_ratio >= ProfessionalDocumentProcessor._FONT_SIZE_RATIO
                            and len(block_text.split()) <= 20
                            and len(block_text) >= 3
                            and (is_bold or font_ratio >= 1.4)
                        )

                        if is_heading:
                            potential_headings.append(block_text.strip())

                        page_blocks.append({
                            "text":       block_text,
                            "font_size":  block_max_font,
                            "is_bold":    is_bold,
                            "is_heading": is_heading,
                            "y_pos":      blk.get("bbox", [0, 0, 0, 0])[1],
                        })
                        plain_text += block_text + "\n\n"

                    page_text_stripped = plain_text.strip()
                    needs_ocr = len(page_text_stripped) < ProfessionalDocumentProcessor._MIN_CHARS_PER_PAGE

                    result["pages"].append({
                        "page_number":        page_idx + 1,
                        "text":               page_text_stripped,
                        "blocks":             page_blocks,
                        "potential_headings": potential_headings,
                        "needs_ocr":          needs_ocr,
                    })
                    if needs_ocr:
                        result["needs_ocr_pages"].append(page_idx + 1)

                except Exception as exc:
                    logger.warning(f"[Layer1] خطأ في الصفحة {page_idx + 1}: {exc}")

        finally:
            doc.close()   # مضمون الإغلاق حتى عند أي exception غير متوقع

        return result

    # ─────────────────────────────────────────────────────────────────────
    #  الطبقة 2: OCR للصفحات المصوَّرة
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_layer2_ocr(
        file_path: str, page_numbers: List[int]
    ) -> Dict[int, str]:
        """
        الطبقة 2: OCR بـ easyocr للصفحات التي أعادت نصاً قصيراً جداً.
        يعيد {رقم_الصفحة: النص_المستخرج}.
        """
        if not page_numbers:
            return {}

        easyocr_mod = ProfessionalDocumentProcessor._import_easyocr()
        fitz        = ProfessionalDocumentProcessor._import_fitz()

        if easyocr_mod is None or fitz is None:
            logger.warning(
                "[Layer2] easyocr أو pymupdf غير متاح — تخطّي OCR. "
                "للتثبيت: pip install easyocr"
            )
            return {}

        ocr_results: Dict[int, str] = {}
        logger.info(f"[Layer2] تشغيل OCR على {len(page_numbers)} صفحة مصوَّرة …")

        try:
            reader = easyocr_mod.Reader(["ar", "en"], gpu=False, verbose=False)
        except Exception as exc:
            logger.warning(f"[Layer2] فشل تهيئة easyocr: {exc}")
            return {}

        try:
            doc = fitz.open(file_path)
            try:
                for pn in page_numbers:
                    try:
                        page      = doc[pn - 1]
                        mat       = fitz.Matrix(2, 2)   # دقة مضاعفة تُحسِّن OCR
                        pix       = page.get_pixmap(matrix=mat)
                        img_bytes = pix.tobytes("png")

                        import io
                        import numpy as np
                        from PIL import Image
                        img       = Image.open(io.BytesIO(img_bytes))
                        img_array = np.array(img)

                        ocr_lines = reader.readtext(img_array, detail=0, paragraph=True)
                        ocr_text  = "\n".join(ocr_lines)

                        if ocr_text.strip():
                            ocr_results[pn] = ocr_text
                            logger.info(f"[Layer2] OCR صفحة {pn}: {len(ocr_text)} حرف")
                        else:
                            logger.warning(f"[Layer2] OCR لم يجد نصاً في الصفحة {pn}")

                    except Exception as exc:
                        logger.warning(f"[Layer2] خطأ OCR صفحة {pn}: {exc}")
            finally:
                doc.close()   # مضمون الإغلاق حتى عند أي exception غير متوقع
        except Exception as exc:
            logger.error(f"[Layer2] خطأ عام OCR: {exc}")

        return ocr_results

    # ─────────────────────────────────────────────────────────────────────
    #  التراجع الآمن: PyPDF2
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_fallback_pypdf2(file_path: str) -> Dict[str, Any]:
        """
        تراجع آمن لـ PyPDF2 عند غياب pymupdf.
        يُرجع نفس هيكل _extract_layer1_fitz لضمان توافق باقي الكود.
        """
        logger.warning("[Fallback] استخدام PyPDF2 — النتائج قد تكون أقل دقة")
        result: Dict[str, Any] = {
            "pages":           [],
            "toc":             [],
            "metadata":        {},
            "total_pages":     0,
            "is_protected":    False,
            "needs_ocr_pages": [],
        }
        try:
            with open(file_path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                result["total_pages"] = len(reader.pages)

                if reader.metadata:
                    meta = dict(reader.metadata)
                    result["metadata"] = {
                        "title":  meta.get("/Title",  ""),
                        "author": meta.get("/Author", ""),
                    }

                for idx, pg in enumerate(reader.pages):
                    try:
                        txt      = pg.extract_text() or ""
                        needs_ocr = len(txt.strip()) < ProfessionalDocumentProcessor._MIN_CHARS_PER_PAGE
                        result["pages"].append({
                            "page_number":       idx + 1,
                            "text":              txt,
                            "blocks":            [],
                            "potential_headings": [],
                            "needs_ocr":         needs_ocr,
                        })
                        if needs_ocr:
                            result["needs_ocr_pages"].append(idx + 1)
                    except Exception as exc:
                        logger.warning(f"[Fallback] خطأ صفحة {idx + 1}: {exc}")
        except Exception as exc:
            logger.error(f"[Fallback] خطأ PyPDF2: {exc}")
            raise
        return result

    # ─────────────────────────────────────────────────────────────────────
    #  بناء الفصول من TOC الداخلي
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _detect_chapters_from_toc(
        toc: List[Dict[str, Any]], pages: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        بناء الفصول من TOC الداخلي للـ PDF (المسار الأكثر دقة).
        يُصفِّي المستويات العليا (level ≤ 2) ليحصل على الفصول الرئيسية.
        """
        if not toc:
            return []

        top_level = [item for item in toc if item.get("level", 1) <= 2] or toc
        total_pages = len(pages)
        chapters: List[Dict[str, Any]] = []

        for idx, toc_item in enumerate(top_level):
            start_pg = max(1, min(toc_item.get("page", 1), total_pages))
            end_pg   = (
                max(start_pg, min(top_level[idx + 1].get("page", total_pages) - 1, total_pages))
                if idx + 1 < len(top_level)
                else total_pages
            )

            chapter_text = "\n\n".join(
                p.get("text", "")
                for p in pages
                if start_pg <= p.get("page_number", 0) <= end_pg
                and p.get("text", "").strip()
            ).strip()

            if not chapter_text:
                continue

            chapters.append({
                "id":         f"chapter_{idx + 1:03d}",
                "title":      toc_item["title"].strip(),
                "content":    chapter_text,
                "start_page": start_pg,
                "end_page":   end_pg,
                "word_count": len(chapter_text.split()),
            })

        logger.info(f"[ChapterDetect/TOC] {len(chapters)} فصل من TOC الداخلي")
        return chapters

    # ─────────────────────────────────────────────────────────────────────
    #  بناء الفصول من البيانات الطباعية
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _detect_chapters_from_typography(
        pages: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        كشف الفصول من حجم الخط والموضع الطباعي (يُستخدم عند غياب TOC).
        يُفضِّل العناوين المكتشفة بالخط، ويتراجع للأنماط النصية الكلاسيكية.
        """
        chapters: List[Dict[str, Any]] = []
        current_chapter: Optional[Dict[str, Any]] = None
        chapter_counter = 1

        for page_data in pages:
            page_num  = page_data.get("page_number", 0)
            page_text = page_data.get("text", "").strip()
            blocks    = page_data.get("blocks", [])

            if not page_text:
                continue

            # ── كشف العنوان في هذه الصفحة ──────────────────────────
            detected_title: Optional[str] = None

            # المسار الأول: عنوان مكتشف بالخط
            for blk in blocks:
                if blk.get("is_heading"):
                    candidate = blk.get("text", "").strip()
                    if 2 <= len(candidate.split()) <= 15 and len(candidate) >= 3:
                        detected_title = candidate
                        break

            # المسار الثاني: أنماط كلاسيكية
            if not detected_title:
                for ln in page_text.split("\n"):
                    ln = ln.strip()
                    if ProfessionalDocumentProcessor._is_chapter_title_by_pattern(ln):
                        detected_title = ln
                        break

            # ── تجميع الفصول ────────────────────────────────────────
            if detected_title:
                if current_chapter and current_chapter.get("content", "").strip():
                    chapters.append(current_chapter)
                current_chapter = {
                    "id":         f"chapter_{chapter_counter:03d}",
                    "title":      detected_title,
                    "content":    page_text,
                    "start_page": page_num,
                    "end_page":   page_num,
                    "word_count": len(page_text.split()),
                }
                chapter_counter += 1
            elif current_chapter is not None:
                current_chapter["content"]   += "\n\n" + page_text
                current_chapter["end_page"]   = page_num
                current_chapter["word_count"] = len(current_chapter["content"].split())
            else:
                current_chapter = {
                    "id":         f"chapter_{chapter_counter:03d}",
                    "title":      f"الجزء {chapter_counter}",
                    "content":    page_text,
                    "start_page": page_num,
                    "end_page":   page_num,
                    "word_count": len(page_text.split()),
                }
                chapter_counter += 1

        if current_chapter and current_chapter.get("content", "").strip():
            chapters.append(current_chapter)

        logger.info(f"[ChapterDetect/Typography] {len(chapters)} فصل بالكشف الطباعي")
        return chapters

    @staticmethod
    def _is_chapter_title_by_pattern(line: str) -> bool:
        """فحص سريع: هل هذا السطر عنوان فصل بالأنماط الكلاسيكية؟"""
        if not line or len(line) < 2 or len(line) > 120:
            return False
        patterns = [
            r'^(Chapter|CHAPTER)\s+(\d+|[IVX]+)[\:\.\-\s]*',
            r'^(الفصل|فصل|القسم|الباب)\s+(\d+|[ا-ي]+)[\:\.\-\s]*',
            r'^\s*(\d+)[\.\-]\s+.{3,50}$',
            # الرقم الروماني يجب أن يتبعه نقطة أو شرطة فقط (لا مسافة)
            # يمنع مطابقة جمل إنجليزية تبدأ بـ "I" كـ "I hate hospitals."
            r'^\s*([IVX]{1,6})[.\-]\s+.{3,50}$',
        ]
        return any(re.match(p, line, re.IGNORECASE) for p in patterns)

    # ─────────────────────────────────────────────────────────────────────
    #  التقسيم الذكي المحترم للوحدات السردية
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def smart_text_division(
        text: str, target_chunk_size: int = 5000
    ) -> List[Dict[str, Any]]:
        """
        تقسيم ذكي للنص يحترم الوحدات السردية الطبيعية:
        ✅ يكتشف فواصل المشاهد (*** --- ~~~) ويقطع عندها
        ✅ يكمل الفقرة الحالية قبل القطع (لا يكسر الجملة)
        ✅ لا ينتج جزءاً بحجم صفر
        ✅ يحافظ على حد الكلمات كحد أقصى لا حد إجباري
        """
        paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        chunks: List[Dict[str, Any]] = []
        current_parts: List[str] = []
        current_word_count       = 0
        chunk_counter            = 1

        def _flush() -> None:
            nonlocal chunk_counter
            content = "\n\n".join(current_parts).strip()
            if not content:
                return
            chunks.append({
                "id":         f"chapter_{chunk_counter:03d}",
                "title":      f"الجزء {chunk_counter}",
                "content":    content,
                "word_count": len(content.split()),
                "start_page": chunk_counter,
                "end_page":   chunk_counter,
            })
            chunk_counter += 1
            current_parts.clear()

        for para in paragraphs:
            para_words = len(para.split())
            is_scene_break = bool(
                ProfessionalDocumentProcessor._SCENE_BREAK_PATTERNS.match(para)
            )

            # فاصل مشهد صريح + تجاوزنا نصف الحجم المستهدف → قطع هنا
            if is_scene_break and current_word_count >= target_chunk_size // 2:
                _flush()
                current_word_count = 0
                continue  # لا نُضيف رمز الفاصل للنص

            # تجاوز الحجم المستهدف → أنهِ الجزء الحالي ثم ابدأ جديداً
            if current_word_count + para_words > target_chunk_size and current_parts:
                _flush()
                current_word_count = 0

            current_parts.append(para)
            current_word_count += para_words

        _flush()   # الجزء الأخير
        return chunks

    # ─────────────────────────────────────────────────────────────────────
    #  فحص سلامة الاستخراج
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _verify_extraction_integrity(
        chapters: List[Dict[str, Any]], total_pages: int
    ) -> Tuple[bool, str]:
        """
        فحص منطقية نتائج الاستخراج قبل المتابعة.
        يُرجع (True, "OK") إذا كانت سليمة، أو (False, سبب_المشكلة) إذا لا.
        """
        if not chapters:
            return False, "لم يُستخرج أي فصل"

        empty = [ch for ch in chapters if ch.get("word_count", 0) == 0]
        if empty:
            return False, f"يوجد {len(empty)} فصل فارغ"

        avg_words = sum(ch.get("word_count", 0) for ch in chapters) / len(chapters)

        if avg_words < 30:
            return False, f"متوسط طول الفصل منخفض جداً: {avg_words:.0f} كلمة"

        if avg_words > ProfessionalDocumentProcessor._MAX_CHAPTER_WORDS:
            return (
                False,
                f"متوسط طول الفصل مرتفع جداً: {avg_words:.0f} كلمة "
                f"(يُرجَّح فشل اكتشاف الفصول)",
            )

        if total_pages > 0:
            pages_per_chapter = total_pages / len(chapters)
            if pages_per_chapter > 100:
                return (
                    False,
                    f"فصل واحد لكل {pages_per_chapter:.0f} صفحة — "
                    f"يبدو أن اكتشاف الفصول فشل",
                )

        return True, "OK"

    # ─────────────────────────────────────────────────────────────────────
    #  تنظيف النص المستخرج (محسَّن جوهرياً)
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        تنظيف النص المستخرج من PDF بمنطق محسَّن:
        ✅ إزالة أرقام الصفحات المنفردة
        ✅ الحفاظ على فواصل المشاهد (*** --- ~~~)
        ✅ دمج الأسطر بمنطق يُفرِّق بين العربية والإنجليزية
        ✅ إزالة الأسطر القصيرة عديمة المعنى
        ✅ تنظيف المسافات والأسطر الزائدة
        """
        if not text:
            return ""

        lines = text.split("\n")
        cleaned: List[str] = []

        for ln in lines:
            stripped = ln.strip()
            if not stripped:
                continue
            # الحفاظ على فواصل المشاهد
            if ProfessionalDocumentProcessor._SCENE_BREAK_PATTERNS.match(stripped):
                cleaned.append(stripped)
                continue
            # إزالة أرقام الصفحات المنفردة
            if re.match(r'^\d{1,4}$', stripped):
                continue
            # تجاهل الأسطر القصيرة جداً
            if len(stripped) < 3:
                continue
            cleaned.append(stripped)

        # ── دمج الأسطر بمنطق سياقي ──────────────────────────────────
        merged = ""
        for i, ln in enumerate(cleaned):
            if not merged:
                merged = ln
                continue

            prev = cleaned[i - 1]

            is_scene_break     = ProfessionalDocumentProcessor._SCENE_BREAK_PATTERNS.match(ln)
            prev_ends_sentence = prev.endswith((".", "؟", "!", ":", "،", "؛"))
            curr_starts_upper  = ln[0].isupper() if ln[0].isalpha() else False
            curr_starts_arabic = "\u0600" <= ln[0] <= "\u06FF"
            prev_is_short      = len(prev) < 40

            if is_scene_break:
                merged += "\n\n" + ln + "\n\n"
            elif prev_ends_sentence or prev_is_short or curr_starts_upper:
                merged += "\n\n" + ln
            elif curr_starts_arabic:
                merged += " " + ln
            else:
                merged += " " + ln

        # تنظيف نهائي
        merged = re.sub(r"\n{3,}", "\n\n", merged)
        merged = re.sub(r" {2,}",  " ",    merged)
        merged = re.sub(r"\n ",    "\n",   merged)
        return merged.strip()

    # ─────────────────────────────────────────────────────────────────────
    #  detect_chapter_titles — مُبقى للتوافق مع باقي الكود
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def detect_chapter_titles(text: str) -> List[str]:
        """
        كشف عناوين الفصول بالأنماط النصية الكلاسيكية.
        مُبقى للتوافق مع باقي الكود؛ الكشف الرئيسي يتم الآن
        بواسطة _detect_chapters_from_toc و _detect_chapters_from_typography.
        """
        lines  = text.split("\n")
        titles: List[str] = []
        patterns = [
            r'^(Chapter|CHAPTER)\s+(\d+|[IVX]+)[\:\.\-\s]*(.*)',
            r'^(الفصل|فصل|القسم|الباب)\s+(\d+|[ا-ي]+)[\:\.\-\s]*(.*)',
            r'^\s*(\d+)[\.\-\s](.{5,50})',
            # الرقم الروماني يجب أن يتبعه نقطة أو شرطة فقط (لا مسافة)
            # يمنع مطابقة جمل إنجليزية تبدأ بـ "I" كـ "I hate hospitals."
            r'^\s*([IVX]{1,6})[.\-]\s+(.{5,50})',
            r'^([A-Z][A-Z\s]{10,80})',
        ]
        for ln in lines:
            ln = ln.strip()
            if len(ln) < 3 or len(ln) > 100:
                continue
            for pat in patterns:
                if re.match(pat, ln, re.IGNORECASE):
                    titles.append(ln)
                    break
        return titles

    # ─────────────────────────────────────────────────────────────────────
    #  الدالة الرئيسية للاستخراج — الواجهة العامة
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def extract_pdf_with_precision(file_path: str) -> Dict[str, Any]:
        """
        استخراج دقيق للنص من PDF بنظام الطبقات الثلاث مع التحقق التلقائي.

        الترتيب:
          1. pymupdf (أو PyPDF2 تراجعاً)
          2. OCR بـ easyocr للصفحات المصوَّرة
          3. إشعار واضح عند الملفات المحمية

        ثم:
          • قراءة TOC الداخلي وبناء الفصول منه إن وُجد
          • كشف الفصول طباعياً (حجم الخط) إن لم يوجد TOC
          • تقسيم ذكي كمسار أخير
          • فحص سلامة مع إعادة محاولة تلقائية
        """
        logger.info(f"[ProfessionalDocumentProcessor] بدء معالجة: {file_path}")

        document_info: Dict[str, Any] = {
            "title":       "",
            "author":      "",
            "chapters":    [],
            "total_pages": 0,
            "metadata":    {},
        }

        fitz_available = ProfessionalDocumentProcessor._import_fitz() is not None

        # ── الطبقة 1 ─────────────────────────────────────────────────
        raw = (
            ProfessionalDocumentProcessor._extract_layer1_fitz(file_path)
            if fitz_available
            else ProfessionalDocumentProcessor._extract_fallback_pypdf2(file_path)
        )

        if raw.get("error") and raw["error"] != "pymupdf_missing":
            logger.error(f"[ProfessionalDocumentProcessor] خطأ في الاستخراج: {raw['error']}")
            raise RuntimeError(f"فشل استخراج PDF: {raw['error']}")

        if raw.get("is_protected"):
            logger.error("[ProfessionalDocumentProcessor] الملف محمي بكلمة مرور")
            raise RuntimeError(
                "الملف محمي بكلمة مرور. يرجى فكّ الحماية قبل المعالجة.\n"
                "مثال: qpdf --decrypt input.pdf output.pdf"
            )

        document_info["total_pages"] = raw.get("total_pages", 0)
        document_info["metadata"]    = raw.get("metadata", {})
        document_info["title"]       = raw["metadata"].get("title",  "")
        document_info["author"]      = raw["metadata"].get("author", "")

        pages: List[Dict[str, Any]] = raw.get("pages", [])
        toc:   List[Dict[str, Any]] = raw.get("toc",   [])

        # ── الطبقة 2: OCR ────────────────────────────────────────────
        needs_ocr_pages: List[int] = raw.get("needs_ocr_pages", [])
        if needs_ocr_pages:
            logger.info(
                f"[ProfessionalDocumentProcessor] {len(needs_ocr_pages)} صفحة تحتاج OCR"
            )
            ocr_map = ProfessionalDocumentProcessor._extract_layer2_ocr(
                file_path, needs_ocr_pages
            )
            for page_data in pages:
                pn = page_data["page_number"]
                if pn in ocr_map:
                    existing = page_data.get("text", "")
                    page_data["text"] = (existing + "\n\n" + ocr_map[pn]).strip()
                    page_data["needs_ocr"] = False

        # ── تنظيف نص الصفحات ─────────────────────────────────────────
        for page_data in pages:
            if page_data.get("text"):
                page_data["text"] = ProfessionalDocumentProcessor.clean_extracted_text(
                    page_data["text"]
                )

        # ── بناء الفصول (3 مسارات مُرتَّبة بالأولوية) ───────────────
        chapters: List[Dict[str, Any]] = []

        if toc:
            chapters = ProfessionalDocumentProcessor._detect_chapters_from_toc(toc, pages)

        if not chapters and fitz_available:
            chapters = ProfessionalDocumentProcessor._detect_chapters_from_typography(pages)

        if not chapters:
            logger.warning(
                "[ProfessionalDocumentProcessor] فشل اكتشاف الفصول — تفعيل التقسيم الذكي"
            )
            full_text = "\n\n".join(
                p.get("text", "") for p in pages if p.get("text")
            )
            chapters = ProfessionalDocumentProcessor.smart_text_division(full_text)

        document_info["chapters"] = chapters

        # ── فحص السلامة مع إعادة المحاولة ───────────────────────────
        ok, reason = ProfessionalDocumentProcessor._verify_extraction_integrity(
            chapters, document_info["total_pages"]
        )

        if not ok:
            logger.warning(
                f"[IntegrityCheck] نتائج غير منطقية ({reason}) — إعادة التقسيم الذكي"
            )
            full_text = "\n\n".join(
                p.get("text", "") for p in pages if p.get("text")
            )
            chapters = ProfessionalDocumentProcessor.smart_text_division(full_text)
            document_info["chapters"] = chapters

            ok2, reason2 = ProfessionalDocumentProcessor._verify_extraction_integrity(
                chapters, document_info["total_pages"]
            )
            if ok2:
                logger.info("[IntegrityCheck] الفحص الثاني: النتائج منطقية بعد إعادة التقسيم")
            else:
                logger.error(f"[IntegrityCheck] فشل الفحص الثاني: {reason2}")

        # ── إحصائيات ختامية ──────────────────────────────────────────
        total_words = sum(ch.get("word_count", 0) for ch in document_info["chapters"])
        logger.info(
            f"[ProfessionalDocumentProcessor] اكتمل: "
            f"{document_info['total_pages']} صفحة | "
            f"{len(document_info['chapters'])} فصل | "
            f"{total_words:,} كلمة | "
            f"OCR: {len(needs_ocr_pages)} صفحة | "
            f"TOC داخلي: {'✓' if toc else '✗'}"
        )

        return document_info


class EnhancedDocumentGenerator:
    """
    مولد المستندات المحسن للروايات.

    التحسينات المُطبَّقة على الكلاس:
      1. الفهرس من العناوين الحقيقية المستخرجة من PDF — لا تخمين.
      2. كاشف نوع المقطع (_detect_paragraph_type) لتطبيق أنماط متعددة.
      3. أنماط إضافية: DialogueLine, SceneBreak, EmbeddedPoem,
         EmbeddedDocument, SubHeading — مع استبقاء كل الأنماط الأصلية.
      4. إصلاح clean_novel_paragraph: حماية الحوارات والجمل الأدبية القصيرة.
      5. verify_document: فحص شامل بعد الحفظ مع تقرير جودة كامل.
      6. _setup_document_styles: استخراج إعداد الأنماط لدالة منفصلة.
    """

    # ── ثوابت أنواع المقاطع ──────────────────────────────────────────────
    PARA_NORMAL        = 'normal'
    PARA_DIALOGUE      = 'dialogue'
    PARA_SCENE_BREAK   = 'scene_break'
    PARA_EMBEDDED_POEM = 'embedded_poem'
    PARA_EMBEDDED_DOC  = 'embedded_document'
    PARA_SUBHEADING    = 'subheading'

    # خريطة النوع → اسم النمط في python-docx
    _PARA_STYLE_MAP = {
        'normal':           'NovelText',
        'dialogue':         'DialogueLine',
        'scene_break':      'SceneBreak',
        'embedded_poem':    'EmbeddedPoem',
        'embedded_document':'EmbeddedDocument',
        'subheading':       'SubHeading',
    }

    # بادئات تدل على حوار
    _DIALOGUE_STARTERS = ('"', '\u201c', '\u201d', '\u00ab', '\u00bb', '\u2014', '\u2013', '-')
    _DIALOGUE_VERBS    = (
        'قال', 'قالت', 'أجاب', 'أجابت', 'صاح', 'صاحت',
        'همس', 'همست', 'سأل', 'سألت', 'ردّ', 'ردّت', 'ردت',
    )

    # بادئات وثائق/رسائل مضمَّنة
    _DOC_STARTERS = (
        'عزيزي', 'عزيزتي', 'إلى:', 'من:', 'التاريخ:', 'الموضوع:',
        'Dear', 'To:', 'From:', 'Date:', 'Subject:',
    )

    # ─────────────────────────────────────────────────────────────────────
    #  مساعد: فحص عربية النص
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _is_arabic_text(text: str) -> bool:
        """يُعيد True إذا كان أكثر من 30% من الأحرف في النص أحرفاً عربية."""
        if not text:
            return False
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        return arabic_chars / max(len(text), 1) > 0.30

    # ─────────────────────────────────────────────────────────────────────
    #  كاشف نوع المقطع
    # ─────────────────────────────────────────────────────────────────────

    @classmethod
    def _detect_paragraph_type(cls, text: str) -> str:
        """
        يفحص نص الفقرة ويُعيد نوعها لاختيار النمط المناسب.

        الأنواع المُعادة:
          normal          — سرد عادي (الافتراضي الآمن)
          dialogue        — حوار (بدء بعلامة اقتباس صريحة أو فعل قول)
          scene_break     — فاصل مشهد (*** أو --- متجاورة بلا مسافات)
          embedded_poem   — شعر مضمَّن (علامة ^ أو # في بداية السطر)
          embedded_document — رسالة/وثيقة مضمَّنة
          subheading      — عنوان فرعي (يبدأ بـ ## أو علامة صريحة)
          empty           — فراغ

        ملاحظة التصميم:
          • SubHeading و EmbeddedPoem يتطلبان علامات صريحة فقط لمنع
            التصنيف الخاطئ للجمل السردية القصيرة.
          • أي نص لا تنطبق عليه شروط صريحة يُعامَل كـ normal.
        """
        stripped = text.strip()
        if not stripped:
            return 'empty'

        # ── فاصل مشهد: أحرف خاصة متجاورة فقط (بلا مسافات) ─────────
        # يُكتشف قبل وصوله هنا من خلال detect_and_preserve_scene_breaks
        if re.match(r'^[\*\-\~\_]{3,}$', stripped):
            return cls.PARA_SCENE_BREAK

        # ── وثيقة/رسالة مضمَّنة: بادئات صريحة ───────────────────────
        if any(stripped.startswith(s) for s in cls._DOC_STARTERS):
            return cls.PARA_EMBEDDED_DOC

        # ── حوار: يبدأ بعلامة اقتباس صريحة أو فعل قول ──────────────
        if any(stripped.startswith(s) for s in cls._DIALOGUE_STARTERS):
            return cls.PARA_DIALOGUE
        if any(stripped.startswith(v) for v in cls._DIALOGUE_VERBS):
            return cls.PARA_DIALOGUE

        # ── شعر مضمَّن: علامة ^ أو # في البداية (صريحة فقط) ─────────
        if stripped.startswith('^') or stripped.startswith('# '):
            return cls.PARA_EMBEDDED_POEM

        # ── عنوان فرعي: علامة ## أو [عنوان] صريحة ───────────────────
        if stripped.startswith('## ') or (stripped.startswith('[') and stripped.endswith(']')):
            return cls.PARA_SUBHEADING

        # ── سرد عادي: الافتراضي الآمن لجميع الحالات الأخرى ──────────
        return cls.PARA_NORMAL

    # ─────────────────────────────────────────────────────────────────────
    #  إعداد أنماط المستند (مستخرج لدالة مستقلة)
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def _setup_document_styles(doc: 'Document') -> None:
        """
        يُنشئ جميع الأنماط المخصصة في المستند إذا لم تكن موجودة بعد.

        الأنماط الأصلية (محتفظ بها كما هي):
          NovelTitle, ChapterTitle, NovelText, TOCTitle, TOCEntry

        الأنماط الجديدة:
          DialogueLine    — تمييز بصري للحوار
          SceneBreak      — فاصل مشهد مُوسَّط
          EmbeddedPoem    — شعر مضمَّن بخط مائل ومركزي
          EmbeddedDocument — رسالة/وثيقة بمسافة بادئة مزدوجة
          SubHeading      — عنوان فرعي داخل الفصل
        """
        styles = doc.styles

        def _get_or_create(name: str):
            """يُعيد النمط إذا وُجد، وإلا يُنشئه."""
            if name in styles:
                return styles[name]
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

        # ── النمط الأساسي ────────────────────────────────────────────
        base = styles['Normal']
        base.font.name  = 'Arial'
        base.font.rtl   = True
        base.font.size  = Pt(14)
        base.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        base.paragraph_format.line_spacing      = 1.3
        base.paragraph_format.space_after       = Pt(6)
        base.paragraph_format.space_before      = Pt(0)
        base.paragraph_format.first_line_indent = Inches(0.25)

        # ── عنوان الرواية (NovelTitle) ───────────────────────────────
        s = _get_or_create('NovelTitle')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(18)
        s.font.bold = True
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after  = Pt(10)
        s.paragraph_format.space_before = Pt(0)

        # ── عنوان الفصل (ChapterTitle) ───────────────────────────────
        s = _get_or_create('ChapterTitle')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(15)
        s.font.bold = True
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(8)

        # ── النص الأساسي (NovelText) ─────────────────────────────────
        s = _get_or_create('NovelText')
        s.base_style = styles['Normal']
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(14)
        s.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
        s.paragraph_format.line_spacing       = 1.25
        s.paragraph_format.space_after        = Pt(4)
        s.paragraph_format.space_before       = Pt(0)
        s.paragraph_format.first_line_indent  = Inches(0.2)
        s.paragraph_format.widow_control      = True
        s.paragraph_format.keep_together      = False

        # ── عنوان الفهرس (TOCTitle) ──────────────────────────────────
        s = _get_or_create('TOCTitle')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(16)
        s.font.bold = True
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after  = Pt(15)
        s.paragraph_format.space_before = Pt(0)

        # ── عناصر الفهرس (TOCEntry) ──────────────────────────────────
        s = _get_or_create('TOCEntry')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(13)
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.RIGHT
        s.paragraph_format.space_after  = Pt(6)
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.left_indent  = Inches(0.2)

        # ── حوار (DialogueLine) — جديد ───────────────────────────────
        s = _get_or_create('DialogueLine')
        s.base_style = styles['NovelText']
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(14)
        s.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.RIGHT
        s.paragraph_format.line_spacing      = 1.25
        s.paragraph_format.space_after       = Pt(3)
        s.paragraph_format.space_before      = Pt(3)
        s.paragraph_format.first_line_indent = Inches(0.3)
        s.paragraph_format.left_indent       = Inches(0.1)
        s.paragraph_format.widow_control     = True

        # ── فاصل مشهد (SceneBreak) — جديد ───────────────────────────
        s = _get_or_create('SceneBreak')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(12)
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after  = Pt(12)
        s.paragraph_format.space_before = Pt(12)

        # ── شعر مضمَّن (EmbeddedPoem) — جديد ────────────────────────
        s = _get_or_create('EmbeddedPoem')
        s.font.name   = 'Arial';  s.font.rtl = True;  s.font.size = Pt(13)
        s.font.italic = True
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.space_after  = Pt(6)
        s.paragraph_format.space_before = Pt(6)

        # ── وثيقة/رسالة مضمَّنة (EmbeddedDocument) — جديد ───────────
        s = _get_or_create('EmbeddedDocument')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(12)
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
        s.paragraph_format.line_spacing = 1.2
        s.paragraph_format.space_after  = Pt(4)
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.left_indent  = Inches(0.5)
        s.paragraph_format.right_indent = Inches(0.5)

        # ── عنوان فرعي (SubHeading) — جديد ──────────────────────────
        s = _get_or_create('SubHeading')
        s.font.name = 'Arial';  s.font.rtl = True;  s.font.size = Pt(14)
        s.font.bold = True
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after  = Pt(4)

    # ─────────────────────────────────────────────────────────────────────
    #  بناء الفهرس من العناوين الحقيقية
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    async def create_table_of_contents(chapters: List[Dict[str, Any]],
                                       api_manager: 'EnhancedGeminiAPI') -> List[Dict[str, str]]:
        """
        يبني الفهرس من العناوين الحقيقية المخزَّنة في بيانات الفصل
        (المستخرجة من PDF في مرحلة الاستخراج) — لا حاجة لتخمين من المحتوى.

        الأولوية لكل فصل:
          1. chapter['arabic_title']  — عنوان مترجم مُخزَّن مسبقاً (أفضل)
          2. chapter['title']         — العنوان الأصلي المستخرج:
               • عربي → يُستخدم مباشرةً
               • غير عربي وغير عام → يُترجم عبر API مرة واحدة فقط
          3. "الفصل N"                — احتياطي أخير إذا فشل كل ما سبق

        يضمن عدم تكرار أي عنوان في الفهرس.
        """
        logger.info("Building TOC from real extracted chapter titles (no content scanning)...")

        table_of_contents: List[Dict[str, str]] = []
        processed_original_titles: set = set()
        used_arabic_titles: set = set()
        fallback_counter = 1

        # نمط العناوين العامة الداخلية (chapter_003, الجزء 5, ...)
        _GENERIC_PATTERN = re.compile(
            r'^(chapter|الجزء|part|section|فصل)\s*[\d]+\s*$|^chapter_\d+$',
            re.IGNORECASE
        )

        for chapter in chapters:
            if not chapter.get('translated_content'):
                continue

            original_title: str = chapter.get('title', '').strip()

            # ── تجنب تكرار نفس العنوان الأصلي ───────────────────────
            if original_title in processed_original_titles:
                continue
            processed_original_titles.add(original_title)

            arabic_title: Optional[str] = None

            # ── المصدر 1: عنوان مترجم مُخزَّن مسبقاً ────────────────
            pre_translated = chapter.get('arabic_title', '').strip()
            if pre_translated and EnhancedDocumentGenerator._is_arabic_text(pre_translated):
                arabic_title = pre_translated[:60]

            # ── المصدر 2: العنوان الأصلي المستخرج ───────────────────
            if not arabic_title and original_title:
                is_generic = bool(_GENERIC_PATTERN.match(original_title))

                if not is_generic:
                    if EnhancedDocumentGenerator._is_arabic_text(original_title):
                        # العنوان عربي أصلاً → استخدمه مباشرةً
                        arabic_title = original_title[:60]
                    else:
                        # العنوان بلغة أخرى → ترجمه عبر API
                        prompt = (
                            f"اترجم عنوان الفصل التالي إلى العربية بشكل مختصر ومميز "
                            f"(٣-٨ كلمات فقط):\n\n{original_title}\n\n"
                            f"عنوان مترجم فقط (بدون شرح):"
                        )
                        try:
                            result = await api_manager.make_precision_request(
                                prompt,
                                temperature=0.2,
                                request_type="chapter_title_translation"
                            )
                            translated, _, _ = result if result else (None, 0.0, None)
                            if translated:
                                arabic_title = translated.strip()[:60]
                        except Exception as e:
                            logger.warning(
                                f"Title translation failed for '{original_title}': {e}"
                            )

            # ── الاحتياطي الأخير ─────────────────────────────────────
            if not arabic_title:
                arabic_title = f"الفصل {fallback_counter}"

            # ── ضمان فريد العنوان ────────────────────────────────────
            if arabic_title in used_arabic_titles:
                arabic_title = f"{arabic_title} ({fallback_counter})"
            used_arabic_titles.add(arabic_title)
            fallback_counter += 1

            table_of_contents.append({
                'original_title': original_title,
                'arabic_title':   arabic_title,
            })

        real_titles = sum(
            1 for e in table_of_contents
            if not e['arabic_title'].startswith('الفصل')
        )
        logger.info(
            f"TOC built: {len(table_of_contents)} entries "
            f"({real_titles} with real names, "
            f"{len(table_of_contents) - real_titles} fallback)"
        )
        return table_of_contents
    
    @staticmethod
    def create_novel_document(chapters: List[Dict[str, Any]],
                              output_path: str,
                              book_title: str = "الرواية المترجمة",
                              author: str = "مترجم بالذكاء الاصطناعي",
                              table_of_contents: List[Dict[str, str]] = None) -> str:
        """
        ينشئ مستند DOCX احترافي للرواية مع:
          • فهرس منفصل بعناوين حقيقية
          • تنسيق ذكي لكل نوع فقرة (سرد/حوار/شعر/وثيقة/فاصل مشهد/عنوان فرعي)
          • حماية المحتوى الأدبي القصير (الحوارات الدرامية)
          • فحص شامل للمستند بعد الحفظ مع تقرير جودة في السجل
        """
        logger.info(f"Creating novel document with smart multi-style formatting: {output_path}")

        try:
            # ── إعداد مجلد الإخراج ──────────────────────────────────
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            # ── إعداد الصفحة ─────────────────────────────────────────
            section = doc.sections[0]
            section.page_width    = Inches(6)
            section.page_height   = Inches(9)
            section.left_margin   = Inches(0.7)
            section.right_margin  = Inches(0.9)
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)

            # ── إنشاء جميع الأنماط (الأصلية + الجديدة) ───────────────
            EnhancedDocumentGenerator._setup_document_styles(doc)

            # ── صفحة العنوان ─────────────────────────────────────────
            doc.add_paragraph(book_title, style='NovelTitle')

            if author and author != "مترجم بالذكاء الاصطناعي":
                author_para = doc.add_paragraph(author)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if author_para.runs:
                    author_para.runs[0].font.size = Pt(13)
                    author_para.runs[0].font.rtl  = True
                author_para.paragraph_format.space_after = Pt(5)

            doc.add_page_break()

            # ── صفحة الفهرس المنفصلة ─────────────────────────────────
            if table_of_contents:
                doc.add_paragraph("فهرس المحتويات", style='TOCTitle')
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(8)

                for i, toc_entry in enumerate(table_of_contents, 1):
                    chapter_number_text = ComprehensiveContentProcessor.number_to_arabic_text(i)
                    toc_line = f"الفصل {chapter_number_text}: {toc_entry['arabic_title']}"
                    toc_para = doc.add_paragraph(toc_line, style='TOCEntry')
                    toc_para.paragraph_format.space_after = Pt(6)

            doc.add_page_break()

            # ── بناء خريطة: original_title → arabic_title ────────────
            toc_index_map: Dict[str, str] = {}
            if table_of_contents:
                for entry in table_of_contents:
                    toc_index_map[entry['original_title']] = entry['arabic_title']

            # ── محتوى الرواية ────────────────────────────────────────
            used_chapter_titles: set = set()

            for i, chapter in enumerate(chapters):
                if not chapter.get('translated_content'):
                    logger.warning(f"Skipping untranslated chapter: {chapter.get('title', '?')}")
                    continue

                # ── تحديد عنوان الفصل ────────────────────────────────
                orig = chapter.get('title', '')
                if orig in toc_index_map:
                    candidate = toc_index_map[orig]
                elif table_of_contents and i < len(table_of_contents):
                    candidate = table_of_contents[i]['arabic_title']
                else:
                    candidate = f"الفصل {i + 1}"

                # السماح بفصل بلا عنوان إذا كان مكرراً (بدلاً من إسقاطه)
                chapter_title: Optional[str] = None
                if candidate not in used_chapter_titles:
                    chapter_title = candidate
                    used_chapter_titles.add(chapter_title)

                if chapter_title and not chapter_title.startswith('الجزء'):
                    doc.add_paragraph(chapter_title, style='ChapterTitle')

                # ── فقرات الفصل بالتنسيق الذكي ───────────────────────
                content    = chapter['translated_content']
                paragraphs = content.split('\n\n')

                for raw_para in paragraphs:
                    raw_para = raw_para.strip()
                    if not raw_para:
                        continue

                    # إزالة العنوان المكرر من بداية الفقرة
                    if chapter_title and chapter_title in raw_para:
                        raw_para = raw_para.replace(chapter_title, '').strip()
                    if not raw_para:
                        continue

                    # ── تنظيف مع حماية الأدب القصير ──────────────────
                    clean_text = EnhancedDocumentGenerator.clean_novel_paragraph(raw_para)
                    if not clean_text:
                        continue

                    # ── كشف النوع واختيار النمط المناسب ──────────────
                    para_type  = EnhancedDocumentGenerator._detect_paragraph_type(clean_text)
                    if para_type == 'empty':
                        continue

                    style_name = EnhancedDocumentGenerator._PARA_STYLE_MAP.get(
                        para_type, 'NovelText'
                    )

                    doc.add_paragraph(clean_text, style=style_name)

            # ── حفظ المستند ──────────────────────────────────────────
            doc.save(output_path)
            logger.info(f"Novel document saved: {output_path}")
            logger.info("Font sizes: Body 14pt, Titles 15pt | TOC: Arabic numerals")

            # ── فحص ما بعد الإنشاء ───────────────────────────────────
            # عدد الفصول المتوقع في المستند = فصول مترجمة بعناوين فريدة
            # لا تبدأ بـ "الجزء" (وهي المستثناة من الكتابة في المستند)
            seen_titles: set = set()
            expected_chapters_in_doc = 0
            for ch in chapters:
                if not ch.get('translated_content'):
                    continue
                orig = ch.get('title', '')
                toc_title = toc_index_map.get(orig, '')
                if not toc_title and table_of_contents:
                    idx = next(
                        (i for i, c in enumerate(chapters) if c is ch), -1
                    )
                    toc_title = (table_of_contents[idx]['arabic_title']
                                 if 0 <= idx < len(table_of_contents) else f"الفصل {idx+1}")
                if toc_title and toc_title not in seen_titles and not toc_title.startswith('الجزء'):
                    seen_titles.add(toc_title)
                    expected_chapters_in_doc += 1
            expected_toc   = len(table_of_contents) if table_of_contents else 0

            verification = EnhancedDocumentGenerator.verify_document(
                output_path, expected_chapters_in_doc, expected_toc
            )

            if verification['passed']:
                logger.info(
                    f"✅ Post-creation verification PASSED | "
                    f"Paragraphs: {verification['stats'].get('non_empty_paragraphs', 0)} | "
                    f"Words: {verification['stats'].get('total_words', 0):,} | "
                    f"Size: {verification['stats'].get('file_size_bytes', 0):,} bytes"
                )
            else:
                logger.warning("⚠️ Post-creation verification found issues:")
                for w in verification.get('warnings', []):
                    logger.warning(f"   • {w}")

            quality_logger.info(
                "Document verification report",
                passed   = verification['passed'],
                checks   = verification['checks'],
                stats    = verification['stats'],
                warnings = verification.get('warnings', []),
            )

            return output_path

        except Exception as e:
            logger.error(f"Error creating novel document: {str(e)}")
            logger.error(traceback.format_exc())
            raise
    
    @staticmethod
    def clean_novel_paragraph(text: str) -> str:
        """
        تنظيف الفقرة مع حماية المحتوى الأدبي القصير ذي الثقل الدرامي.

        محمي (لا يُحذف):
          • فواصل المشاهد (*** أو --- أو ~~~) — تُُحفظ وتُوحَّد كـ "***"
          • جمل الحوار القصيرة ("لا.", "نعم.", "قالت: كلا!", "صمت طويل.")
          • أي نص قصير يحتوي أحرفاً عربية أو علامات اقتباس/حوار
          • ALL CAPS عربي (يدل على صراخ أو تأكيد في الحوار)

        محذوف:
          • أرقام الصفحات المنفردة
          • رموز بادئة بلا محتوى أدبي (•, *, :, ;) — مع استثناء فواصل المشاهد
          • تكرار مباشر للكلمات المتجاورة
          • نص لا يحتوي إلا على أرقام ورموز وفراغات
          • ALL CAPS إنجليزي قصير (قد يكون عنواناً مكرراً)
        """
        if not text:
            return ""

        stripped = text.strip()

        # ── حماية أولى: فاصل مشهد — يُحوَّل مباشرةً ولا يمر بالفلاتر ──
        # يتعامل مع: *** أو --- أو ~~~ أو * * * (بمسافات)
        if re.match(r'^[\*\-\~\_\s]{3,}$', stripped) and re.search(r'[\*\-\~\_]', stripped):
            return '***'

        text = stripped

        # ── إزالة أرقام الصفحات المنفردة ────────────────────────────
        text = re.sub(r'^\s*\d{1,4}\s*$', '', text, flags=re.MULTILINE)

        # ── إزالة رموز زخرفية بادئة بلا محتوى (لا تمس الحوار) ───────
        text = re.sub(r'^[•\*\.\:\;]\s+(?=\S)', '', text.strip())

        # ── إزالة أرقام البداية المتصلة بلا نص إضافي ────────────────
        text = re.sub(r'^\d+[\.\-\s]+(?=[^\d])', '', text)

        # ── إزالة التكرار المباشر للكلمات المتجاورة ─────────────────
        words = text.split()
        if len(words) > 1:
            clean_words = [words[0]]
            for w in words[1:]:
                if w != clean_words[-1]:
                    clean_words.append(w)
            text = ' '.join(clean_words)

        # ── تنظيف المسافات الزائدة ───────────────────────────────────
        text = re.sub(r'\s+', ' ', text).strip()

        if not text:
            return ""

        # ── فحص: نص رموز وأرقام فقط ← يُحذف ────────────────────────
        if re.match(r'^[\d\s\-\*•\.\:\;،]+$', text):
            return ""

        # ── حماية المحتوى الأدبي القصير (< 20 حرف) ──────────────────
        if len(text) < 20:
            has_arabic        = any('\u0600' <= c <= '\u06FF' for c in text)
            has_dialogue_mark = any(c in ('"', '\u201c', '\u201d', '\u00ab', '\u00bb',
                                          '\u2014', '\u2013') for c in text)
            # أبقِه إذا كان يحمل معنى أدبياً (عربي أو علامة حوار)
            if has_arabic or has_dialogue_mark:
                return text
            return ""

        # ── ALL CAPS: فرّق بين العربي (صراخ/تأكيد) والإنجليزي ───────
        if text.isupper() and len(text) < 50:
            has_arabic = any('\u0600' <= c <= '\u06FF' for c in text)
            if has_arabic:
                return text   # صراخ أو تأكيد في حوار → محفوظ
            return ""         # أحرف إنجليزية كبيرة فقط → عنوان مكرر → يُحذف

        return text

    # ─────────────────────────────────────────────────────────────────────
    #  فحص المستند بعد الإنشاء
    # ─────────────────────────────────────────────────────────────────────

    @staticmethod
    def verify_document(output_path: str,
                        expected_chapters: int,
                        expected_toc_entries: int) -> Dict[str, Any]:
        """
        فحص شامل للمستند بعد حفظه وإعادة تقرير جودة كامل.

        الفحوصات:
          ✓ الملف موجود
          ✓ حجم الملف > 10KB
          ✓ الملف قابل للفتح بـ python-docx
          ✓ يحتوي على فقرات غير فارغة
          ✓ عدد عناوين الفصول يطابق المتوقع
          ✓ عدد عناصر الفهرس يطابق المتوقع
          ✓ لا يوجد فصلان بعناوين متتالية بلا محتوى بينهما

        يُعيد:
          {
            'passed': bool,          ← True إذا نجحت الفحوصات الأساسية
            'checks': {key: bool},   ← نتيجة كل فحص
            'warnings': [str],       ← تحذيرات تفصيلية
            'stats': {key: value},   ← إحصائيات المستند
          }
        """
        report: Dict[str, Any] = {
            'passed':   False,
            'checks':   {},
            'warnings': [],
            'stats':    {},
        }

        file_path = Path(output_path)

        # ── فحص 1: الملف موجود ──────────────────────────────────────
        report['checks']['file_exists'] = file_path.exists()
        if not file_path.exists():
            report['warnings'].append(f"الملف غير موجود: {output_path}")
            return report

        # ── فحص 2: الحجم معقول (> 10KB) ────────────────────────────
        file_size = file_path.stat().st_size
        report['stats']['file_size_bytes'] = file_size
        size_ok = file_size > 10_240
        report['checks']['file_size_ok'] = size_ok
        if not size_ok:
            report['warnings'].append(
                f"حجم الملف صغير جداً: {file_size:,} بايت (الحد الأدنى 10KB)"
            )

        # ── فحص 3: الملف قابل للفتح ─────────────────────────────────
        try:
            verify_doc = Document(output_path)
            report['checks']['file_readable'] = True

            all_paragraphs  = verify_doc.paragraphs
            non_empty_paras = [p for p in all_paragraphs if p.text.strip()]
            total_words     = sum(len(p.text.split()) for p in non_empty_paras)

            report['stats']['total_paragraphs']     = len(all_paragraphs)
            report['stats']['non_empty_paragraphs'] = len(non_empty_paras)
            report['stats']['total_words']          = total_words

            # ── فحص 4: يحتوي فقرات ──────────────────────────────────
            paras_ok = len(non_empty_paras) > 0
            report['checks']['paragraphs_ok'] = paras_ok
            if not paras_ok:
                report['warnings'].append("المستند لا يحتوي على أي فقرات غير فارغة!")

            # ── فحص 5: عدد عناوين الفصول ────────────────────────────
            chapter_titles = [
                p for p in all_paragraphs
                if p.style.name == 'ChapterTitle' and p.text.strip()
            ]
            actual_chapters = len(chapter_titles)
            report['stats']['actual_chapters_in_doc'] = actual_chapters
            ch_ok = (actual_chapters == expected_chapters)
            report['checks']['chapter_count_ok'] = ch_ok
            if not ch_ok:
                report['warnings'].append(
                    f"عناوين الفصول في المستند ({actual_chapters}) "
                    f"≠ المتوقع ({expected_chapters})"
                )

            # ── فحص 6: عدد عناصر الفهرس ─────────────────────────────
            toc_entries = [
                p for p in all_paragraphs
                if p.style.name == 'TOCEntry' and p.text.strip()
            ]
            actual_toc = len(toc_entries)
            report['stats']['actual_toc_entries'] = actual_toc
            toc_ok = (actual_toc == expected_toc_entries)
            report['checks']['toc_count_ok'] = toc_ok
            if not toc_ok:
                report['warnings'].append(
                    f"عناصر الفهرس في المستند ({actual_toc}) "
                    f"≠ المتوقع ({expected_toc_entries})"
                )

            # ── فحص 7: لا فصول بلا محتوى (عناوين متتالية) ───────────
            empty_chapters = sum(
                1 for idx in range(len(all_paragraphs) - 1)
                if (all_paragraphs[idx].style.name == 'ChapterTitle' and
                    all_paragraphs[idx + 1].style.name == 'ChapterTitle')
            )
            report['stats']['chapters_without_content'] = empty_chapters
            report['checks']['no_empty_chapters'] = (empty_chapters == 0)
            if empty_chapters > 0:
                report['warnings'].append(
                    f"{empty_chapters} فصل بلا محتوى (عناوين ChapterTitle متتالية)"
                )

            # ── فحص 8: لا عناوين فهرس مكررة ─────────────────────────
            toc_texts  = [p.text.strip() for p in toc_entries]
            duplicates = len(toc_texts) - len(set(toc_texts))
            report['stats']['duplicate_toc_entries'] = duplicates
            report['checks']['no_duplicate_toc'] = (duplicates == 0)
            if duplicates > 0:
                report['warnings'].append(
                    f"{duplicates} عنوان مكرر في الفهرس"
                )

        except Exception as e:
            report['checks']['file_readable'] = False
            report['warnings'].append(f"خطأ في فتح المستند: {str(e)}")
            return report

        # ── الحكم النهائي (الفحوصات الأساسية فقط) ───────────────────
        critical = ['file_exists', 'file_size_ok', 'file_readable', 'paragraphs_ok']
        report['passed'] = all(report['checks'].get(c, False) for c in critical)

        return report


class MasterTranslationSystem:
    """النظام الرئيسي الشامل للترجمة عالية الجودة - المحسن"""

    def __init__(self, api_keys: List[str], target_language: str = "Arabic"):
        self.api_manager = EnhancedGeminiAPI(api_keys)
        self.translation_engine = CompleteTranslationEngine(self.api_manager, target_language)
        self.document_processor = ProfessionalDocumentProcessor()
        self.document_generator = EnhancedDocumentGenerator()

        # إعداد قاعدة البيانات مع فحص السلامة
        self.db_path = "master_translation_enhanced.db"
        self.init_advanced_database()

        # قفل asyncio للحماية الدقيقة عند الحاجة لقسم حرج متعدد الخطوات عبر await.
        # ملاحظة: asyncio أحادي الخيط (cooperative multitasking) — عمليات القاموس
        # الأتومية (تعيين/قراءة مفتاح واحد) آمنة بين نقاط await دون قفل.
        # يُستخدم هذا القفل فقط إذا احتجنا مستقبلاً لقسم read-modify-write
        # يمتد عبر await (مثلاً: async with _terminology_lock: ... await ... write).
        self._terminology_lock = asyncio.Lock()

        # إحصائيات مفصلة — كل عملية كتابة محمية بـ GIL في asyncio
        self.translation_stats = {
            'total_chapters': 0,
            'completed_chapters': 0,
            'skipped_chapters': 0,
            'total_words': 0,
            'translated_words': 0,
            'total_characters': 0,
            'translation_start_time': None,
            'estimated_completion_time': None,
            'average_quality_score': 0.0,
            'quality_scores': [],          # لحساب متوسط حقيقي
            'foreign_content_corrections': 0,
            'contextual_adaptations': 0,
            'smart_retries': 0,            # عدد مرات تفعيل الإعادة الذكية
            'split_recoveries': 0,         # نجاح إستراتيجية التقسيم
            'rephrasing_recoveries': 0,    # نجاح إستراتيجية إعادة الصياغة
        }

        logger.info("Enhanced main system for high-quality translation initialized")
    
    def init_advanced_database(self):
        """إنشاء قاعدة بيانات متقدمة مع فحص سلامة عند بدء التشغيل"""

        # ── فحص السلامة أولاً إذا كانت قاعدة البيانات موجودة مسبقاً ─────
        db_file = Path(self.db_path)
        if db_file.exists():
            integrity_ok, integrity_msg = self._verify_db_integrity()
            if not integrity_ok:
                logger.warning(
                    f"[DB] Database integrity check FAILED: {integrity_msg}. "
                    f"Renaming corrupted DB and starting fresh."
                )
                # إعادة تسمية قاعدة البيانات التالفة بدلاً من حذفها
                backup_name = f"{self.db_path}.corrupted_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                db_file.rename(backup_name)
                logger.warning(f"[DB] Corrupted DB backed up as: {backup_name}")
            else:
                logger.info(f"[DB] Integrity check passed: {integrity_msg}")

        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            cursor = conn.cursor()

            # جدول الفصول مع معلومات مفصلة
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS chapters (
                    id TEXT PRIMARY KEY,
                    title TEXT,
                    original_content TEXT,
                    translated_content TEXT,
                    word_count INTEGER,
                    character_count INTEGER,
                    genre TEXT DEFAULT 'prose',
                    tone TEXT DEFAULT 'neutral',
                    status TEXT DEFAULT 'pending',
                    translation_attempts INTEGER DEFAULT 0,
                    quality_score REAL DEFAULT 0.0,
                    translation_time REAL DEFAULT 0.0,
                    foreign_content_detected BOOLEAN DEFAULT 0,
                    corrections_applied INTEGER DEFAULT 0,
                    retry_strategy TEXT DEFAULT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            # إضافة عمود retry_strategy إذا لم يكن موجوداً (للتوافق مع قواعد بيانات قديمة)
            try:
                cursor.execute("ALTER TABLE chapters ADD COLUMN retry_strategy TEXT DEFAULT NULL")
                conn.commit()
            except sqlite3.OperationalError:
                pass  # العمود موجود بالفعل

            # جدول المصطلحات المتقدم
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS terminology (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    original_term TEXT UNIQUE,
                    translated_term TEXT,
                    category TEXT,
                    frequency INTEGER DEFAULT 1,
                    confidence_score REAL DEFAULT 1.0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            # جدول السجلات المفصل
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS translation_logs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chapter_id TEXT,
                    operation TEXT,
                    status TEXT,
                    message TEXT,
                    duration REAL,
                    api_key_used TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            # جدول الأحداث الاستخباراتية (الجديد)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS intelligent_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    event_type TEXT,
                    api_key TEXT,
                    duration REAL,
                    genre TEXT,
                    tone TEXT,
                    word_count INTEGER,
                    status TEXT,
                    error_type TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            conn.commit()

        logger.info("Advanced database created and extended with intelligence tables")

    def _verify_db_integrity(self) -> Tuple[bool, str]:
        """
        فحص سلامة قاعدة البيانات باستخدام PRAGMA integrity_check.
        يُعيد (True, 'ok') إذا كانت سليمة، أو (False, رسالة_الخطأ) إذا كانت تالفة.
        """
        try:
            with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
                conn.execute("PRAGMA journal_mode=WAL")   # وضع Write-Ahead Logging لأفضل موثوقية
                result = conn.execute("PRAGMA integrity_check").fetchone()
                if result and result[0] == "ok":
                    return True, "ok"
                else:
                    return False, str(result[0]) if result else "unknown error"
        except Exception as e:
            return False, str(e)
    
    def save_chapter_advanced(self, chapter_data: Dict[str, Any]):
        """حفظ متقدم للفصل مع جميع البيانات بما فيها إستراتيجية الإعادة"""

        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            cursor = conn.cursor()

            cursor.execute('''
                INSERT OR REPLACE INTO chapters
                (id, title, original_content, translated_content, word_count, character_count,
                 genre, tone, status, translation_attempts, quality_score, translation_time,
                 foreign_content_detected, corrections_applied, retry_strategy, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                chapter_data['id'],
                chapter_data['title'],
                chapter_data.get('original_content', ''),
                chapter_data.get('translated_content', ''),
                chapter_data.get('word_count', 0),
                chapter_data.get('character_count', 0),
                chapter_data.get('genre', 'prose'),
                chapter_data.get('tone', 'neutral'),
                chapter_data.get('status', 'pending'),
                chapter_data.get('translation_attempts', 0),
                chapter_data.get('quality_score', 0.0),
                chapter_data.get('translation_time', 0.0),
                chapter_data.get('foreign_content_detected', False),
                chapter_data.get('corrections_applied', 0),
                chapter_data.get('retry_strategy', None),
                datetime.now().isoformat()
            ))

            conn.commit()
    
    def log_operation(self, chapter_id: str, operation: str, status: str, 
                     message: str, duration: float = 0.0, api_key: str = ""):
        """تسجيل العمليات في السجل"""
        
        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            cursor = conn.cursor()

            cursor.execute('''
                INSERT INTO translation_logs
                (chapter_id, operation, status, message, duration, api_key_used)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (chapter_id, operation, status, message, duration, api_key[:15] if api_key else ""))

            conn.commit()

    def log_intelligent_event(self, event_type: str, api_key: str = None,
                              duration: float = None, genre: str = None, tone: str = None,
                              word_count: int = None, status: str = None, error_type: str = None):
        """تسجيل حدث استخباراتي مهيكل في قاعدة البيانات لتحليلات الأداء والمشاكل"""
        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO intelligent_events
                (event_type, api_key, duration, genre, tone, word_count, status, error_type)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (event_type, api_key[:15] if api_key else None, duration, genre, tone, word_count, status, error_type))
            conn.commit()

    def analyze_and_display_intelligence(self):
        """تحليل البيانات الاستخباراتية من قاعدة البيانات وعرضها باستخدام Rich"""
        # Table و Panel مستوردان من المستوى الأعلى

        console.print("\n[bold cyan]🔍 Extracting intelligence analytics from logs...[/bold cyan]")

        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            cursor = conn.cursor()

            try:
                # 1. تحليل متوسط وقت الترجمة لكل نوع أدبي
                cursor.execute('''
                    SELECT genre, AVG(duration) as avg_duration, COUNT(*) as count
                    FROM intelligent_events
                    WHERE event_type = 'translation_complete' AND duration IS NOT NULL
                    GROUP BY genre
                ''')
                genre_stats = cursor.fetchall()

                if genre_stats:
                    table_genre = Table(title="[bold]Average Translation Time by Genre[/bold]", show_header=True, header_style="bold magenta")
                    table_genre.add_column("Genre", justify="right")
                    table_genre.add_column("Avg Time (s)", justify="center")
                    table_genre.add_column("Chapters", justify="center")

                    for genre, avg_duration, count in genre_stats:
                        table_genre.add_row(str(genre), f"{avg_duration:.2f}", str(count))

                    console.print(table_genre)
                else:
                    console.print("[dim]Not enough data to analyze literary genres.[/dim]")

                # 2. تحليل الساعات الأكثر استجابة (متوسط وقت الاستجابة لكل ساعة)
                cursor.execute('''
                    SELECT strftime('%H', timestamp) as hour, AVG(duration) as avg_duration, COUNT(*) as count
                    FROM intelligent_events
                    WHERE event_type = 'translation_complete' AND duration IS NOT NULL
                    GROUP BY hour
                    ORDER BY avg_duration ASC
                ''')
                hour_stats = cursor.fetchall()

                if hour_stats:
                    table_hour = Table(title="[bold]Performance Patterns: Fastest Response Hours[/bold]", show_header=True, header_style="bold green")
                    table_hour.add_column("Hour (Server Time)", justify="center")
                    table_hour.add_column("Avg Response (s)", justify="center")
                    table_hour.add_column("Operations", justify="center")

                    for hour, avg_duration, count in hour_stats:
                        table_hour.add_row(f"{hour}:00", f"{avg_duration:.2f}", str(count))

                    console.print(table_hour)

                    # استنتاج ذكي
                    best_hour = hour_stats[0][0]
                    console.print(Panel(f"[bold green]💡 Smart Insight:[/bold green] The best time to use the system is around [bold]{best_hour}:00[/bold] as the API response is at its fastest.", title="Performance Tip"))
                else:
                    console.print("[dim]Not enough data to analyze the best response hours.[/dim]")

            except Exception as e:
                console.print(f"[bold red]Error extracting intelligence analytics: {str(e)}[/bold red]")

    def _load_completed_chapters_from_db(self) -> Dict[str, Any]:
        """تحميل الفصول المكتملة مسبقاً من قاعدة البيانات"""
        logger.info("Checking database for previously translated chapters...")

        with contextlib.closing(sqlite3.connect(self.db_path)) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()

            cursor.execute("SELECT * FROM chapters WHERE status = 'completed'")
            completed_chapters = {}
            rows = cursor.fetchall()
            for row in rows:
                chapter_data = dict(row)
                completed_chapters[chapter_data['id']] = chapter_data
            
            if completed_chapters:
                logger.info(f"Found {len(completed_chapters)} previously translated chapters. Skipping translation for them.")
            else:
                logger.info("No previously translated chapters found.")
            
            return completed_chapters

    # ================================================================== #
    #  اختبار مفاتيح API                                                  #
    # ================================================================== #

    async def test_all_api_keys(self) -> Dict[str, bool]:
        """
        اختبار جميع مفاتيح API قبل بدء الترجمة وعرض تقرير Rich.
        يُرسل طلباً تجريبياً بسيطاً لكل مفتاح بالتوازي،
        ويعرض جدولاً بالحالة والمعدل التاريخي لكل مفتاح.

        يُطلق RuntimeError إذا لم يكن هناك أي مفتاح صالح.
        """
        # Table و Panel مستوردان من المستوى الأعلى

        console.print("\n[bold cyan]🔑 Testing all API keys before starting translation...[/bold cyan]")

        test_prompt = "أجب بكلمة 'جاهز' فقط."
        results: Dict[str, bool] = {}

        async def _test_single_key(key: str) -> Tuple[str, bool, str]:
            """اختبار مفتاح واحد وإرجاع (key, valid, status_msg)"""
            try:
                await self.api_manager._ensure_session()
                test_body = {
                    "contents": [{"parts": [{"text": test_prompt}], "role": "user"}],
                    "generationConfig": {"maxOutputTokens": 10, "temperature": 0.1}
                }
                url = f"{self.api_manager.base_url}?key={key}"
                async with self.api_manager.session.post(
                    url,
                    json=test_body,
                    timeout=aiohttp.ClientTimeout(total=30)
                ) as response:
                    if response.status == 200:
                        return key, True, "[green]✅ Active[/green]"
                    elif response.status == 429:
                        # مفتاح صالح لكن به rate limit مؤقت
                        return key, True, "[yellow]⚡ Rate limited (valid)[/yellow]"
                    elif response.status in (401, 403):
                        return key, False, f"[red]❌ Invalid/Expired ({response.status})[/red]"
                    else:
                        return key, False, f"[yellow]⚠️  HTTP {response.status}[/yellow]"
            except asyncio.TimeoutError:
                return key, False, "[red]⏱️  Timeout[/red]"
            except Exception as e:
                return key, False, f"[red]🔴 Error: {str(e)[:40]}[/red]"

        # اختبار جميع المفاتيح بالتوازي
        tasks = [_test_single_key(key) for key in self.api_manager.api_keys]
        test_results = await asyncio.gather(*tasks)

        # جدول Rich للعرض
        table = Table(
            title="[bold]API Keys Status Report[/bold]",
            show_header=True,
            header_style="bold blue"
        )
        table.add_column("#",                    justify="center", width=4)
        table.add_column("Key (masked)",         justify="left",   width=26)
        table.add_column("Status",               justify="center", width=30)
        table.add_column("Historical Success %", justify="center", width=22)

        valid_count = 0
        for i, (key, is_valid, status_msg) in enumerate(test_results, 1):
            results[key] = is_valid
            if is_valid:
                valid_count += 1

            masked_key = f"{key[:8]}...{key[-4:]}" if len(key) > 12 else key

            # الحصول على معدل النجاح التاريخي من KeyStatistics
            hist_rate = self.api_manager.key_stats[key].get_success_rate()
            hist_str = f"{hist_rate:.1f}%" if hist_rate >= 0 else "N/A"
            hist_colored = f"[green]{hist_str}[/green]" if hist_rate >= 80 else (
                f"[yellow]{hist_str}[/yellow]" if hist_rate >= 50 else f"[red]{hist_str}[/red]"
            )

            table.add_row(str(i), masked_key, status_msg, hist_colored)

        console.print(table)

        # ملخص نهائي
        ready_color = "green" if valid_count > 0 else "red"
        console.print(Panel(
            f"[{ready_color}]Valid keys: {valid_count}/{len(self.api_manager.api_keys)}[/{ready_color}]  |  "
            f"Ready to translate: {'[green]YES ✅[/green]' if valid_count > 0 else '[red]NO ❌[/red]'}",
            title="API Keys Summary",
            border_style=ready_color
        ))

        if valid_count == 0:
            raise RuntimeError(
                "No valid API keys available. "
                "Please check your keys in EnhancedGeminiAPI or via --api-keys argument."
            )

        logger.info(f"API key test completed: {valid_count}/{len(self.api_manager.api_keys)} valid keys")
        return results

    # ================================================================== #
    #  حساب نقطة الجودة الحقيقية                                         #
    # ================================================================== #

    @staticmethod
    def _calculate_real_quality_score(
        original_content: str,
        translated_content: str,
        foreign_content_detected: bool,
        translation_time: float,
        retry_count: int = 0
    ) -> float:
        """
        حساب نقطة جودة حقيقية (0–10) بدلاً من القيمة الثابتة 8.5.

        المعايير المُستخدمة:
          1. نسبة الاكتمال  — هل الترجمة وافية بالنسبة للأصل؟         (3 نقاط)
          2. غياب المحتوى الأجنبي                                      (2.5 نقطة)
          3. كثافة الأحرف العربية في الناتج                            (2.5 نقطة)
          4. وقت الترجمة   — استجابة سريعة جداً قد تعني ناتجاً ناقصاً (1 نقطة)
          5. عقوبة إعادة المحاولة                                       (1 نقطة)
        """
        if not translated_content or not original_content:
            return 0.0

        score = 10.0

        # ── 1. نسبة الاكتمال ──────────────────────────────────────────
        orig_words  = max(1, len(original_content.split()))
        trans_words = len(translated_content.split())
        completeness = trans_words / orig_words

        # النص العربي أقصر من الإنجليزي عادةً (40-85% طبيعي)
        if completeness < 0.25:
            score -= 3.0       # ترجمة ناقصة جداً
        elif completeness < 0.40:
            score -= 1.5       # ناقصة نسبياً
        elif completeness > 2.8:
            score -= 1.0       # مطولة بشكل مبالغ (احتمال تكرار)

        # ── 2. المحتوى الأجنبي ───────────────────────────────────────
        if foreign_content_detected:
            score -= 2.5

        # ── 3. كثافة الأحرف العربية ──────────────────────────────────
        printable_chars = [c for c in translated_content if c.strip()]
        if printable_chars:
            arabic_chars  = sum(1 for c in printable_chars if '\u0600' <= c <= '\u06FF')
            arabic_ratio  = arabic_chars / len(printable_chars)
            if arabic_ratio < 0.45:
                score -= 2.0   # نسبة عربية منخفضة جداً
            elif arabic_ratio < 0.65:
                score -= 0.8

        # ── 4. سرعة الترجمة ───────────────────────────────────────────
        if translation_time > 0:
            words_per_sec = orig_words / translation_time
            if words_per_sec > 600:
                # استجابة فورية قد تعني رداً مقتضباً غير مكتمل
                score -= 0.5

        # ── 5. عقوبة إعادة المحاولة ──────────────────────────────────
        if retry_count > 0:
            score -= min(1.0, retry_count * 0.35)

        return max(0.0, min(10.0, round(score, 2)))

    # ================================================================== #
    #  إعادة المحاولة الذكية للفصول الفاشلة                              #
    # ================================================================== #

    async def _translate_with_smart_retry(self, chapter: Dict[str, Any]) -> Dict[str, Any]:
        """
        يُجرب ترجمة الفصل؛ فإن فشل يُطبّق إستراتيجيتين بالتتالي:

        إستراتيجية 1 — التقسيم إلى نصفين (للفصول الطويلة > 1500 كلمة):
          • يُقسّم المحتوى إلى نصفين، يُترجم كلاً على حدة بـ _count_stats=False
            لمنع الإحصاء المضاعف، ثم يدمجهما ويُحدّث الإحصائيات الكلية مرة واحدة.

        إستراتيجية 2 — إعادة الصياغة المحايدة (لأخطاء المحتوى الحساس):
          • يُضيف مقدمة سياقية محايدة تُقلل فرص الرفض من نماذج Safety.
          • يُعيد ضبط original_content و id إلى قيمهما الأصليتين قبل الحفظ.
        """
        # ── المحاولة الأولى العادية ────────────────────────────────────
        result = await self.translate_chapter_comprehensively(chapter)
        if result['status'] == 'completed':
            return result

        self.translation_stats['smart_retries'] += 1
        content     = chapter.get('content', '')
        word_count  = chapter.get('word_count', len(content.split()))

        # ── إستراتيجية 1: التقسيم ─────────────────────────────────────
        if word_count > 1500 and content:
            logger.info(
                f"[SmartRetry] '{chapter['title']}' failed (words={word_count}) "
                f"→ splitting into halves"
            )
            words = content.split()
            mid   = len(words) // 2

            half1 = dict(chapter)
            half2 = dict(chapter)

            half1.update({
                'content':              ' '.join(words[:mid]),
                'id':                   f"{chapter['id']}_part1",
                'title':                f"{chapter['title']} (الجزء الأول)",
                'translation_attempts': 0,
            })
            half2.update({
                'content':              ' '.join(words[mid:]),
                'id':                   f"{chapter['id']}_part2",
                'title':                f"{chapter['title']} (الجزء الثاني)",
                'translation_attempts': 0,
            })

            # _count_stats=False → يمنع الإحصاء المضاعف للأجزاء الفرعية
            r1 = await self.translate_chapter_comprehensively(half1, _count_stats=False)
            r2 = await self.translate_chapter_comprehensively(half2, _count_stats=False)

            if r1['status'] == 'completed' and r2['status'] == 'completed':
                merged_translation = (
                    r1.get('translated_content', '') + '\n\n' +
                    r2.get('translated_content', '')
                )
                foreign_merged = (
                    r1.get('foreign_content_detected', False) or
                    r2.get('foreign_content_detected', False)
                )
                time_merged = (
                    r1.get('translation_time', 0) + r2.get('translation_time', 0)
                )
                quality = self._calculate_real_quality_score(
                    content, merged_translation,
                    foreign_merged, time_merged, retry_count=1
                )
                merged = dict(chapter)
                merged.update({
                    'translated_content':      merged_translation,
                    'status':                  'completed',
                    'word_count':              word_count,
                    'quality_score':           quality,
                    'original_content':        content,  # ← الأصل الحقيقي دون تعديل
                    'translation_time':        time_merged,
                    'corrections_applied':     r1.get('corrections_applied', 0) + r2.get('corrections_applied', 0),
                    'foreign_content_detected': foreign_merged,
                    'genre':                   r1.get('genre', chapter.get('genre', 'prose')),
                    'tone':                    r1.get('tone',  chapter.get('tone',  'neutral')),
                    'retry_strategy':          'split_halves',
                })
                self.save_chapter_advanced(merged)

                # ── حذف إدخالات الأجزاء الفرعية من DB (لا لزوم لها بعد الدمج) ──
                try:
                    with contextlib.closing(sqlite3.connect(self.db_path)) as _conn:
                        _conn.execute(
                            "DELETE FROM chapters WHERE id IN (?, ?)",
                            (f"{chapter['id']}_part1", f"{chapter['id']}_part2")
                        )
                        _conn.commit()
                except Exception:
                    pass  # الحذف اختياري — لا يُوقف العملية عند الفشل

                # ── تحديث الإحصائيات الكلية (مرة واحدة فقط للفصل الأصلي) ──
                self.translation_stats['quality_scores'].append(quality)
                if self.translation_stats['quality_scores']:
                    self.translation_stats['average_quality_score'] = (
                        sum(self.translation_stats['quality_scores']) /
                        len(self.translation_stats['quality_scores'])
                    )
                self.translation_stats['completed_chapters']     += 1
                self.translation_stats['translated_words']       += word_count
                self.translation_stats['contextual_adaptations'] += 1
                self.translation_stats['split_recoveries']       += 1
                if foreign_merged:
                    self.translation_stats['foreign_content_corrections'] += 1

                logger.info(f"[SmartRetry] '{chapter['title']}' recovered via split_halves ✅")
                return merged

        # ── إستراتيجية 2: إعادة الصياغة المحايدة ─────────────────────
        logger.info(
            f"[SmartRetry] '{chapter['title']}' → retrying with neutral rephrasing"
        )
        neutral_preamble = (
            "النص التالي مقتطف أدبي من رواية. "
            "يُرجى ترجمته إلى العربية الفصحى بأسلوب أدبي رفيع:\n\n"
        )
        rephrased = dict(chapter)
        rephrased.update({
            'content':              neutral_preamble + content,
            'translation_attempts': 0,
        })

        # _count_stats=True هنا لأن هذا هو الفصل الكامل (لا جزء فرعي)
        result_r = await self.translate_chapter_comprehensively(rephrased)
        if result_r['status'] == 'completed':
            # ── استعادة القيم الأصلية قبل الحفظ النهائي ──────────────
            result_r['id']               = chapter['id']
            result_r['title']            = chapter.get('title', result_r.get('title', ''))
            result_r['original_content'] = content          # ← بدون المقدمة المحايدة
            result_r['retry_strategy']   = 'neutral_rephrasing'
            self.save_chapter_advanced(result_r)
            self.translation_stats['rephrasing_recoveries'] += 1
            logger.info(f"[SmartRetry] '{chapter['title']}' recovered via neutral_rephrasing ✅")
            return result_r

        # ── فشل كل الاستراتيجيات ─────────────────────────────────────
        logger.error(f"[SmartRetry] All retry strategies exhausted for '{chapter['title']}'")
        return result

    async def translate_chapter_comprehensively(
        self,
        chapter: Dict[str, Any],
        _count_stats: bool = True
    ) -> Dict[str, Any]:
        """
        ترجمة شاملة للفصل مع ضمانات الجودة المحسنة ونقطة جودة حقيقية.

        Args:
            chapter:       بيانات الفصل المراد ترجمته.
            _count_stats:  إذا كان False، لا تُحدَّث الإحصائيات الكلية
                           (يُستخدم داخلياً فقط من _translate_with_smart_retry
                           للنصفين الفرعيين حتى لا يُحسَب الفصل الواحد مرتين).
        """

        start_time = time.time()
        chapter_id = chapter['id']

        logger.info(f"Starting comprehensive translation for chapter: {chapter['title']}")
        quality_logger.info(f"Chapter {chapter_id}: Enhanced processing started")

        try:
            # إعداد البيانات الأولية
            content = chapter.get('content', '')
            chapter['original_content'] = content
            chapter['word_count'] = len(content.split())
            chapter['character_count'] = len(content)
            chapter['status'] = 'translating'
            chapter['translation_attempts'] = chapter.get('translation_attempts', 0) + 1

            # اكتشاف نوع النص ونبرته
            text_analysis = self.translation_engine.detect_text_genre_and_tone(content)
            chapter['genre'] = text_analysis['genre']
            chapter['tone']  = text_analysis['tone']

            self.save_chapter_advanced(chapter)
            self.log_operation(
                chapter_id, "translation_start", "info",
                f"Started translating chapter of {chapter['word_count']} words "
                f"- Genre: {text_analysis['genre']}, Tone: {text_analysis['tone']}"
            )

            # الترجمة الشاملة مع المراجعة
            translation_context = f"هذا الفصل بعنوان '{chapter['title']}' من رواية أدبية"

            # ── ملاحظة: asyncio أحادي الخيط (cooperative)، لا تتقاطع coroutines بين
            #    نقطتَي await. عمليات قاموس terminology_database آمنة دون قفل.
            #    _terminology_lock محجوز للحالات التي تحتاج قسماً حرجاً متعدد الخطوط
            #    عبر await (مثل read-modify-write مع تأخير خارجي). هنا غير مطلوب.
            translated_content, response_time, api_key_used = \
                await self.translation_engine.translate_with_comprehensive_review(
                    content, translation_context
                )

            if translated_content:
                translation_time = time.time() - start_time

                # فحص المحتوى الأجنبي النهائي على الناتج الكامل المُعاد من translate_with_comprehensive_review
                # (قد يختلف عن الفحص الداخلي داخل translate_with_comprehensive_review لأن
                #  has_any_foreign_content هنا تفحص النص بالكامل دون سياق الفقرات)
                foreign_content_detected = \
                    self.translation_engine.content_processor.has_any_foreign_content(translated_content)

                # ── حساب نقطة جودة حقيقية ─────────────────────────────
                retry_count = max(0, chapter.get('translation_attempts', 1) - 1)
                real_quality_score = self._calculate_real_quality_score(
                    original_content=content,
                    translated_content=translated_content,
                    foreign_content_detected=foreign_content_detected,
                    translation_time=translation_time,
                    retry_count=retry_count
                )

                # حساب عدد التصحيحات المطبقة بدقة
                # 0 = لا تصحيح | 1 = تصحيح محتوى أجنبي فقط | 2 = تصحيح + تعديل جودة
                corrections_count = (2 if real_quality_score < 7.0 else 1) if foreign_content_detected else 0

                # تحديث بيانات الفصل
                chapter.update({
                    'translated_content':       translated_content,
                    'status':                   'completed',
                    'translation_time':         translation_time,
                    'quality_score':            real_quality_score,
                    'foreign_content_detected': foreign_content_detected,
                    'corrections_applied':      corrections_count,
                })

                # حفظ النتائج
                self.save_chapter_advanced(chapter)

                # ── تحديث الإحصائيات والجودة (فقط للفصول الرئيسية، لا الأجزاء الداخلية) ──
                if _count_stats:
                    self.translation_stats['quality_scores'].append(real_quality_score)
                    if self.translation_stats['quality_scores']:
                        self.translation_stats['average_quality_score'] = (
                            sum(self.translation_stats['quality_scores']) /
                            len(self.translation_stats['quality_scores'])
                        )

                # تسجيل الحدث الاستخباراتي
                self.log_intelligent_event(
                    event_type="translation_complete",
                    api_key=api_key_used,
                    duration=response_time,
                    genre=text_analysis['genre'],
                    tone=text_analysis['tone'],
                    word_count=chapter['word_count'],
                    status="success"
                )

                # تسجيل النجاح في السجلات العادية
                self.log_operation(
                    chapter_id, "translation_complete", "success",
                    f"Translation completed in {translation_time:.2f}s, "
                    f"Genre: {text_analysis['genre']}, "
                    f"Quality: {real_quality_score}/10, "
                    f"Corrections: {corrections_count}",
                    translation_time,
                    api_key_used if api_key_used else ""
                )

                # تحديث عدادات الإحصاء (فقط للفصول الرئيسية)
                if _count_stats:
                    self.translation_stats['completed_chapters'] += 1
                    self.translation_stats['translated_words']   += chapter['word_count']
                    self.translation_stats['contextual_adaptations'] += 1

                    if foreign_content_detected:
                        self.translation_stats['foreign_content_corrections'] += 1
                        quality_logger.warning(
                            f"Chapter {chapter_id}: Applied corrections for foreign content "
                            f"(quality={real_quality_score})"
                        )
                    else:
                        quality_logger.info(
                            f"Chapter {chapter_id}: Free of foreign content "
                            f"(quality={real_quality_score})"
                        )
                else:
                    # وضع داخلي (جزء فرعي) — نسجّل فقط دون تحديث عدادات الكلية
                    quality_logger.info(
                        f"[SubChunk] {chapter_id}: quality={real_quality_score}, "
                        f"foreign={foreign_content_detected}"
                    )

                logger.info(
                    f"Translation finished for chapter '{chapter['title']}' — "
                    f"Time: {translation_time:.2f}s, "
                    f"Genre: {text_analysis['genre']}, "
                    f"Quality: {real_quality_score}/10"
                )
                return chapter

            else:
                # فشل في الترجمة
                chapter['status'] = 'failed'
                self.save_chapter_advanced(chapter)
                self.log_operation(chapter_id, "translation_failed", "error",
                                   "Failed to get translation from API")
                logger.error(f"Translation failed for chapter: {chapter['title']}")
                return chapter

        except Exception as e:
            chapter['status'] = 'error'
            error_message = str(e)
            self.save_chapter_advanced(chapter)
            self.log_operation(chapter_id, "translation_error", "error", error_message)
            logger.error(f"Error translating chapter {chapter['title']}: {error_message}")
            logger.error(traceback.format_exc())
            return chapter
    
    async def process_complete_book(self, pdf_path: str, output_dir: str,
                                    book_title: str = None, author: str = None) -> str:
        """
        معالجة كاملة للكتاب من PDF إلى رواية جاهزة للقراءة مع فهرس منفصل.

        التحسينات الجوهرية في هذه النسخة (مع الحفاظ على التسلسل الأصلي):
          ✅ فحص سلامة DB عند بدء التشغيل
          ✅ نقطة جودة حقيقية محسوبة لكل فصل
          ✅ إعادة محاولة ذكية للفصول الفاشلة (تقسيم / إعادة صياغة)
          ✅ تقرير HTML يُولَّد تلقائياً في مجلد الإخراج
          ✅ ترجمة تسلسلية مضمونة: فصل ثم فصل بالترتيب
        """
        # إنشاء مجلد الإخراج
        output_path_obj = Path(output_dir)
        output_path_obj.mkdir(parents=True, exist_ok=True)

        # تحديد اسم الملف المخرج
        pdf_name    = Path(pdf_path).stem
        output_file = output_path_obj / f"{pdf_name}_رواية_مترجمة.docx"

        logger.info("=" * 100)
        logger.info("Starting enhanced comprehensive processing of the novel with a separate TOC")
        logger.info(f"Source file: {pdf_path}")
        logger.info(f"Target file: {output_file}")
        logger.info("=" * 100)

        self.translation_stats['translation_start_time'] = time.time()

        try:
            # المرحلة 0: عرض التحليلات الاستخباراتية السابقة
            self.analyze_and_display_intelligence()

            # المرحلة 1: استخراج وتحليل المستند
            logger.info("📖 Phase 1: Extracting and analyzing the document...")
            document_structure = self.document_processor.extract_pdf_with_precision(pdf_path)

            # تحميل الفصول المكتملة مسبقاً
            previously_completed = self._load_completed_chapters_from_db()

            chapters = document_structure['chapters']
            self.translation_stats['total_chapters']    = len(chapters)
            self.translation_stats['total_words']       = sum(ch['word_count'] for ch in chapters)
            self.translation_stats['total_characters']  = sum(len(ch.get('content', '')) for ch in chapters)

            if not book_title:
                book_title = document_structure.get('title', 'Translated Novel') or 'Translated Novel'
            if not author:
                author = document_structure.get('author', 'Unknown Author') or 'Unknown Author'

            logger.info(f"📊 Extracted {len(chapters)} chapters")
            logger.info(f"📊 Total words: {self.translation_stats['total_words']:,}")
            logger.info(f"📊 Total characters: {self.translation_stats['total_characters']:,}")
            logger.info(f"📚 Book Title: {book_title}")
            logger.info(f"✍️ Author: {author}")

            # بناء القائمة البيضاء الديناميكية
            logger.info("🧠 Building dynamic whitelist from full book text (character names, places)...")
            full_book_text = " ".join(ch.get('content', '') for ch in chapters)
            self.translation_engine.content_processor.build_book_whitelist(full_book_text)

            # ================================================================ #
            # المرحلة 2: الترجمة التسلسلية — فصل ثم فصل بالترتيب             #
            # ================================================================ #
            logger.info("🔄 Phase 2: Starting sequential translation (chapter by chapter)...")

            all_processed_chapters = []

            with Progress(
                TextColumn("[progress.description]{task.description}"),
                BarColumn(complete_style="green", finished_style="bold green"),
                TaskProgressColumn(),
                TimeRemainingColumn(),
                console=console
            ) as progress:
                translation_task = progress.add_task(
                    "[cyan]Translating chapters...", total=len(chapters)
                )

                for i, chapter in enumerate(chapters):
                    logger.info("-" * 50)

                    # ── فصل مُترجم مسبقاً (skip) ────────────────────────
                    if chapter['id'] in previously_completed:
                        cached = previously_completed[chapter['id']]
                        all_processed_chapters.append(cached)

                        self.translation_stats['skipped_chapters']   += 1
                        self.translation_stats['completed_chapters'] += 1
                        self.translation_stats['translated_words']   += cached.get('word_count', 0)

                        # إضافة نقطة الجودة المحفوظة للمتوسط
                        cached_quality = cached.get('quality_score', 0.0)
                        if cached_quality > 0:
                            self.translation_stats['quality_scores'].append(cached_quality)
                            if self.translation_stats['quality_scores']:
                                self.translation_stats['average_quality_score'] = (
                                    sum(self.translation_stats['quality_scores']) /
                                    len(self.translation_stats['quality_scores'])
                                )

                        logger.info(
                            f"⏭️ Skipping chapter {i+1}/{len(chapters)}: "
                            f"'{chapter['title']}' (previously translated, "
                            f"quality={cached_quality:.2f})"
                        )
                        progress.update(
                            translation_task, advance=1,
                            description=f"[green]Skipped (Cached): {chapter['title']}"
                        )
                        continue

                    # ── ترجمة فعلية مع إعادة المحاولة الذكية ────────────
                    logger.info(f"📝 Translating chapter {i+1}/{len(chapters)}: {chapter['title']}")
                    progress.update(
                        translation_task,
                        description=f"[yellow]Translating: {chapter['title']}"
                    )

                    result = await self._translate_with_smart_retry(chapter)
                    all_processed_chapters.append(result)

                    progress.update(
                        translation_task, advance=1,
                        description=f"[green]Completed: {chapter['title']}"
                    )

                    # تسجيل إحصائيات التقدم كل 5 فصول
                    elapsed_time  = time.time() - self.translation_stats['translation_start_time']
                    chapters_done = i + 1
                    successful_so_far = sum(
                        1 for ch in all_processed_chapters if ch.get('status') == 'completed'
                    )
                    if successful_so_far > 0 and chapters_done % 5 == 0:
                        avg_time = elapsed_time / chapters_done
                        remaining = avg_time * (len(chapters) - chapters_done)
                        logger.info(
                            f"📊 Progress: {successful_so_far}/{len(chapters)} completed | "
                            f"Avg quality: {self.translation_stats['average_quality_score']:.2f}/10 | "
                            f"Est. remaining: {remaining/60:.1f} min"
                        )

            # ================================================================ #
            # المرحلة 3: التحقق النهائي من الجودة                              #
            # ================================================================ #
            logger.info("🔍 Phase 3: Final Quality Check...")

            successful_chapters = [
                ch for ch in all_processed_chapters if ch.get('status') == 'completed'
            ]
            failed_chapters = [
                ch for ch in all_processed_chapters if ch.get('status') in ['failed', 'error']
            ]

            if failed_chapters:
                logger.warning(f"⚠️  {len(failed_chapters)} chapters failed to translate:")
                for ch in failed_chapters:
                    logger.warning(f"   - {ch['title']}")

            chapters_with_foreign = [
                ch for ch in successful_chapters if ch.get('foreign_content_detected', False)
            ]
            if chapters_with_foreign:
                quality_logger.warning(
                    f"Foreign content corrections applied in {len(chapters_with_foreign)} chapters"
                )
            else:
                quality_logger.info("All chapters are completely free of foreign content")

            # ================================================================ #
            # المرحلة 4: إنشاء فهرس منظم بدون تكرار                           #
            # ================================================================ #
            logger.info("📋 Phase 4: Creating organized table of contents without duplicates...")

            table_of_contents = await self.document_generator.create_table_of_contents(
                successful_chapters, self.api_manager
            )
            logger.info(
                f"Professional TOC created with {len(table_of_contents)} unique titles (no page numbers)"
            )

            # ================================================================ #
            # المرحلة 5: إنشاء مستند الرواية النهائي                           #
            # ================================================================ #
            logger.info("📝 Phase 5: Generating final novel document with professional formatting...")
            logger.info("🎯 Font sizes: Body 14pt, Titles 15pt")
            logger.info("📄 TOC: Professional with written Arabic numerals")
            logger.info("📐 Layout: Optimized spacing mirroring printed novels")

            final_document_path = self.document_generator.create_novel_document(
                successful_chapters, str(output_file), book_title, author, table_of_contents
            )

            # ================================================================ #
            # المرحلة 6: تجميع الإحصائيات النهائية وتوليد التقرير HTML         #
            # ================================================================ #
            total_time       = time.time() - self.translation_stats['translation_start_time']
            total_successful = len(successful_chapters)
            total_failed     = len(failed_chapters)

            translated_words = sum(ch.get('word_count', 0) for ch in successful_chapters)
            words_per_minute = translated_words / (total_time / 60) if total_time > 0 else 0
            avg_quality      = self.translation_stats['average_quality_score']

            # ── تقرير نهائي شامل ────────────────────────────────────────────
            logger.info("=" * 100)
            logger.info("🎉 Novel processed successfully with a separate TOC!")
            logger.info("=" * 100)
            logger.info(f"📖 Novel Title: {book_title}")
            logger.info(f"✍️  Author: {author}")
            logger.info(f"📄 Total Chapters: {len(chapters)}")
            logger.info(f"✅ Successfully Translated: {total_successful}")
            logger.info(f"⏭️ Skipped (Previously Translated): {self.translation_stats['skipped_chapters']}")
            logger.info(f"❌ Failed Chapters: {total_failed}")
            logger.info(f"📊 Total Words Translated: {translated_words:,}")
            logger.info(f"⏱️  Total Time: {total_time/60:.1f} minutes")
            logger.info(f"🚀 Translation Rate: {words_per_minute:.0f} words/minute")
            logger.info(f"⭐ Average Quality Score: {avg_quality:.2f}/10")

            logger.info("=" * 50)
            logger.info("🔧 Applied Enhancements Statistics:")
            logger.info(f"   🌍 Foreign Content Corrections: {self.translation_stats['foreign_content_corrections']}")
            logger.info(f"   📖 Contextual Adaptations (Genre/Tone): {self.translation_stats['contextual_adaptations']}")
            logger.info(f"   🔑 Multiple API Keys Used: {len(self.api_manager.api_keys)} keys")
            logger.info(f"   📚 Saved Terms: {len(self.translation_engine.terminology_database)} terms")
            logger.info(f"   📋 Professional TOC: {len(table_of_contents)} chapters with Arabic numerals")
            logger.info(f"   🎯 Uniform Font Sizes: Body 14pt, Titles 15pt")
            logger.info(f"   📐 Enhanced Formatting: Optimized space, calculated spacing")
            logger.info(f"   🔁 Smart Retries Triggered: {self.translation_stats['smart_retries']}")
            logger.info(f"      ↳ Split recoveries:     {self.translation_stats['split_recoveries']}")
            logger.info(f"      ↳ Rephrasing recoveries:{self.translation_stats['rephrasing_recoveries']}")

            logger.info(f"📁 Final professionally formatted novel: {final_document_path}")
            logger.info("=" * 100)

            quality_logger.info("Final Quality Report:")
            quality_logger.info(
                f"Total corrections applied: "
                f"{sum(ch.get('corrections_applied', 0) for ch in successful_chapters)}"
            )

            # تصنيف الفصول حسب النوع
            genre_counts: Dict[str, int] = {}
            tone_counts:  Dict[str, int] = {}
            for ch in successful_chapters:
                genre = ch.get('genre', 'unknown')
                tone  = ch.get('tone',  'unknown')
                genre_counts[genre] = genre_counts.get(genre, 0) + 1
                tone_counts[tone]   = tone_counts.get(tone,  0) + 1

            quality_logger.info("Literary Genre Distribution:")
            for genre, count in genre_counts.items():
                quality_logger.info(f"  {genre}: {count} chapters")

            quality_logger.info("Emotional Tone Distribution:")
            for tone, count in tone_counts.items():
                quality_logger.info(f"  {tone}: {count} chapters")

            # ── توليد تقرير HTML تلقائي ─────────────────────────────────────
            html_path = self.generate_html_report(
                book_title=book_title,
                author=author,
                successful_chapters=successful_chapters,
                failed_chapters=failed_chapters,
                total_time=total_time,
                words_per_minute=words_per_minute,
                avg_quality=avg_quality,
                genre_counts=genre_counts,
                tone_counts=tone_counts,
                output_dir=output_dir,
            )
            logger.info(f"📊 HTML Report generated: {html_path}")

            return final_document_path

        except Exception as e:
            logger.error(f"Fatal error during novel processing: {str(e)}")
            logger.error(traceback.format_exc())
            raise

    # ================================================================== #
    #  توليد تقرير HTML                                                   #
    # ================================================================== #

    def generate_html_report(
        self,
        book_title: str,
        author: str,
        successful_chapters: List[Dict],
        failed_chapters: List[Dict],
        total_time: float,
        words_per_minute: float,
        avg_quality: float,
        genre_counts: Dict[str, int],
        tone_counts: Dict[str, int],
        output_dir: str,
    ) -> str:
        """
        يُولِّد تقرير HTML احترافي بعد اكتمال الترجمة.
        يشمل: إحصائيات عامة، توزيع الجودة، الأنواع الأدبية والنبرات،
        تفاصيل كل فصل، وإحصائيات الإعادة الذكية.
        يُحفظ تلقائياً في مجلد الإخراج.
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', book_title)[:40].strip().replace(' ', '_')
        report_path = Path(output_dir) / f"translation_report_{safe_title}_{timestamp}.html"

        # ── بناء صفوف جدول الفصول ──────────────────────────────────────
        chapter_rows = []
        for i, ch in enumerate(successful_chapters, 1):
            q      = ch.get('quality_score', 0)
            q_color = "#27ae60" if q >= 8 else ("#f39c12" if q >= 6 else "#e74c3c")
            retry  = f'<span style="color:#e67e22">({ch.get("retry_strategy","—")})</span>' \
                     if ch.get('retry_strategy') else '—'
            chapter_rows.append(f"""
                <tr>
                    <td style="text-align:center">{i}</td>
                    <td dir="rtl">{ch.get('title','—')}</td>
                    <td style="text-align:center">{ch.get('word_count',0):,}</td>
                    <td style="text-align:center">{ch.get('genre','—')}</td>
                    <td style="text-align:center">{ch.get('tone','—')}</td>
                    <td style="text-align:center;color:{q_color};font-weight:bold">{q:.2f}</td>
                    <td style="text-align:center">{ch.get('translation_time',0):.1f}s</td>
                    <td style="text-align:center">{'⚠️' if ch.get('foreign_content_detected') else '✅'}</td>
                    <td style="text-align:center">{retry}</td>
                </tr>""")

        for ch in failed_chapters:
            chapter_rows.append(f"""
                <tr style="background:#fdecea">
                    <td style="text-align:center">—</td>
                    <td dir="rtl">{ch.get('title','—')}</td>
                    <td colspan="7" style="text-align:center;color:#e74c3c">
                        ❌ Failed — {ch.get('status','error')}
                    </td>
                </tr>""")

        # ── بناء توزيع الأنواع ──────────────────────────────────────────
        def _dist_bars(counts: Dict[str, int]) -> str:
            total = max(1, sum(counts.values()))
            rows  = []
            for label, n in sorted(counts.items(), key=lambda x: -x[1]):
                pct = n / total * 100
                rows.append(
                    f'<div style="margin:4px 0"><span style="display:inline-block;width:130px">'
                    f'{label}</span>'
                    f'<span style="display:inline-block;background:#3498db;height:14px;'
                    f'width:{pct:.0f}%;max-width:200px;border-radius:3px"></span>'
                    f' <small>{n} ({pct:.0f}%)</small></div>'
                )
            return '\n'.join(rows)

        quality_dist = ""
        if successful_chapters:
            buckets = {"Excellent (9-10)": 0, "Good (7-8.9)": 0,
                       "Fair (5-6.9)": 0, "Poor (<5)": 0}
            for ch in successful_chapters:
                q = ch.get('quality_score', 0)
                if q >= 9:    buckets["Excellent (9-10)"] += 1
                elif q >= 7:  buckets["Good (7-8.9)"]     += 1
                elif q >= 5:  buckets["Fair (5-6.9)"]      += 1
                else:         buckets["Poor (<5)"]          += 1
            quality_dist = _dist_bars(buckets)

        smart_retry_stats = f"""
            <li>🔁 Smart Retries Triggered: <b>{self.translation_stats['smart_retries']}</b></li>
            <li>&nbsp;&nbsp;↳ Split Recoveries: <b>{self.translation_stats['split_recoveries']}</b></li>
            <li>&nbsp;&nbsp;↳ Rephrasing Recoveries: <b>{self.translation_stats['rephrasing_recoveries']}</b></li>
        """

        translated_words = sum(ch.get('word_count', 0) for ch in successful_chapters)
        html = f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
  <meta charset="UTF-8">
  <title>Translation Report — {book_title}</title>
  <style>
    body   {{ font-family: 'Segoe UI', Arial, sans-serif; background:#f4f6f8;
              color:#333; margin:0; padding:20px }}
    .card  {{ background:#fff; border-radius:10px; padding:20px; margin-bottom:20px;
              box-shadow:0 2px 8px rgba(0,0,0,.08) }}
    h1     {{ color:#2c3e50; border-bottom:3px solid #3498db; padding-bottom:10px }}
    h2     {{ color:#2980b9; margin-top:0 }}
    table  {{ width:100%; border-collapse:collapse; font-size:.9em }}
    th     {{ background:#2c3e50; color:#fff; padding:8px 10px }}
    td     {{ padding:7px 10px; border-bottom:1px solid #ecf0f1 }}
    tr:hover {{ background:#f8f9fa }}
    .stat  {{ display:inline-block; background:#3498db; color:#fff; border-radius:8px;
              padding:12px 20px; margin:6px; text-align:center; min-width:120px }}
    .stat b {{ display:block; font-size:1.6em }}
    .good  {{ background:#27ae60 }}
    .warn  {{ background:#f39c12 }}
    .bad   {{ background:#e74c3c }}
  </style>
</head>
<body>
  <div class="card">
    <h1>📊 Translation Report</h1>
    <p>
      <b>📖 Title:</b> {book_title} &nbsp;|&nbsp;
      <b>✍️ Author:</b> {author} &nbsp;|&nbsp;
      <b>📅 Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}
    </p>
  </div>

  <div class="card">
    <h2>📈 Summary Statistics</h2>
    <div>
      <div class="stat"><b>{len(successful_chapters)}</b>Chapters OK</div>
      <div class="stat {'bad' if failed_chapters else 'good'}"><b>{len(failed_chapters)}</b>Failed</div>
      <div class="stat"><b>{translated_words:,}</b>Words</div>
      <div class="stat"><b>{total_time/60:.1f} min</b>Total Time</div>
      <div class="stat"><b>{words_per_minute:.0f}</b>Words/min</div>
      <div class="stat {'good' if avg_quality>=8 else ('warn' if avg_quality>=6 else 'bad')}">
        <b>{avg_quality:.2f}/10</b>Avg Quality
      </div>
    </div>
  </div>

  <div class="card" style="display:flex;gap:20px">
    <div style="flex:1">
      <h2>🎭 Genre Distribution</h2>
      {_dist_bars(genre_counts) if genre_counts else "<p>No data</p>"}
    </div>
    <div style="flex:1">
      <h2>💫 Tone Distribution</h2>
      {_dist_bars(tone_counts) if tone_counts else "<p>No data</p>"}
    </div>
    <div style="flex:1">
      <h2>⭐ Quality Distribution</h2>
      {quality_dist if quality_dist else "<p>No data</p>"}
    </div>
  </div>

  <div class="card">
    <h2>🔧 Smart Retry & Parallel Stats</h2>
    <ul>{smart_retry_stats}</ul>
    <ul>
      <li>🌍 Foreign Content Corrections: <b>{self.translation_stats['foreign_content_corrections']}</b></li>
      <li>📖 Contextual Adaptations: <b>{self.translation_stats['contextual_adaptations']}</b></li>
      <li>🔑 API Keys Used: <b>{len(self.api_manager.api_keys)}</b></li>
      <li>📚 Terminology Entries Saved: <b>{len(self.translation_engine.terminology_database)}</b></li>
    </ul>
  </div>

  <div class="card">
    <h2>📋 Chapter Details</h2>
    <table>
      <tr>
        <th>#</th><th>Title</th><th>Words</th><th>Genre</th><th>Tone</th>
        <th>Quality</th><th>Time</th><th>Foreign</th><th>Retry</th>
      </tr>
      {''.join(chapter_rows)}
    </table>
  </div>
</body>
</html>"""

        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(html)
            logger.info(f"HTML report saved: {report_path}")
        except Exception as e:
            logger.warning(f"Could not save HTML report: {e}")

        return str(report_path)


def validate_input_paths(input_path: str, output_dir: str) -> Tuple[bool, str]:
    """التحقق من صحة مسارات الإدخال والإخراج"""

    if not os.path.exists(input_path):
        return False, f"Input file not found: {input_path}"

    if not input_path.lower().endswith('.pdf'):
        return False, "File must be a PDF"

    try:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return False, f"Cannot create output directory: {str(e)}"

    return True, "Paths are valid"


def _build_cli_parser() -> argparse.ArgumentParser:
    """
    بناء محلل سطر الأوامر الاحترافي.

    الاستخدام الأساسي:
        python PF_5_.py --input /path/to/book.pdf --output /path/to/output/

    أمثلة:
        python PF_5_.py -i book.pdf -o ./output --title "اسم الرواية" --author "المؤلف"
        python PF_5_.py -i book.pdf -o ./output --skip-api-test
        python PF_5_.py -i book.pdf -o ./output --api-keys KEY1 KEY2 KEY3
    """
    parser = argparse.ArgumentParser(
        prog="translation_system",
        description=(
            "🚀 نظام الترجمة الشامل عالي الجودة — يُترجم روايات PDF إلى العربية\n"
            "Complete High-Quality Translation System using Gemini AI"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python PF_5_.py -i /path/to/book.pdf -o /path/to/output
  python PF_5_.py -i book.pdf -o ./out --title "رواية الحرب والسلام" --author "تولستوي"
  python PF_5_.py -i book.pdf -o ./out --skip-api-test --api-keys KEY1 KEY2
        """
    )

    # ── مجموعة الإدخال/الإخراج (إلزامية) ──────────────────────────────
    io_group = parser.add_argument_group("📁 Input / Output  (required)")
    io_group.add_argument(
        "-i", "--input",
        metavar="PDF_PATH",
        required=True,
        help="مسار ملف PDF المراد ترجمته (Path to the source PDF file)"
    )
    io_group.add_argument(
        "-o", "--output",
        metavar="OUTPUT_DIR",
        required=True,
        help="مجلد حفظ الملفات الناتجة (Output directory for translated files)"
    )

    # ── معلومات الكتاب (اختيارية) ───────────────────────────────────────
    book_group = parser.add_argument_group("📚 Book Metadata  (optional)")
    book_group.add_argument(
        "-t", "--title",
        metavar="TITLE",
        default=None,
        help="عنوان الرواية (Novel title; auto-detected from PDF if omitted)"
    )
    book_group.add_argument(
        "-a", "--author",
        metavar="AUTHOR",
        default=None,
        help="اسم المؤلف (Author name; auto-detected if omitted)"
    )

    # ── إعدادات API ──────────────────────────────────────────────────────
    api_group = parser.add_argument_group("🔑 API Configuration  (optional)")
    api_group.add_argument(
        "--api-keys",
        metavar="KEY",
        nargs="+",
        default=[],
        help="مفاتيح Gemini API إضافية تُضاف إلى المفاتيح المدمجة "
             "(Extra Gemini API keys appended to built-in keys)"
    )
    api_group.add_argument(
        "--skip-api-test",
        action="store_true",
        default=False,
        help="تخطي اختبار صحة المفاتيح قبل بدء الترجمة (Skip API key validation)"
    )

    return parser


async def main():
    """
    الدالة الرئيسية للنظام المحسن — واجهة سطر أوامر احترافية.

    تدعم التشغيل التفاعلي القديم (بدون معاملات) للتوافق العكسي،
    وتستخدم argparse عند تمرير معاملات سطر الأوامر.
    """
    parser = _build_cli_parser()

    # ── اكتشاف وضع التشغيل ──────────────────────────────────────────────
    # إذا لم يُمرَّر أي معامل (تشغيل تفاعلي قديم) نستخدم stdin
    interactive_mode = len(sys.argv) == 1

    if interactive_mode:
        # ── الوضع التفاعلي (للتوافق مع السلوك القديم) ──────────────────
        console.print(
            "\n[bold cyan]🚀 Enhanced Comprehensive Translation System[/bold cyan]\n"
            "[dim]Tip: Run with --help to see all CLI options for non-interactive use.[/dim]\n"
        )
        input_path = input("📂 PDF file path: ").strip()
        output_dir = input("📁 Output directory: ").strip()
        book_title_in = input("📖 Novel title (Enter to skip): ").strip() or None
        author_in     = input("✍️  Author name (Enter to skip): ").strip()  or None
        skip_api_test = False
        extra_keys    = []
    else:
        # ── وضع سطر الأوامر ──────────────────────────────────────────────
        args = parser.parse_args()
        input_path    = args.input
        output_dir    = args.output
        book_title_in = args.title
        author_in     = args.author
        skip_api_test = args.skip_api_test
        extra_keys    = args.api_keys

    # ── عرض لافتة المعلومات ───────────────────────────────────────────────
    console.rule("[bold blue]Enhanced Comprehensive Translation System[/bold blue]")
    console.print("[bold green]✨ Features:[/bold green]")
    console.print("  🔑 Multiple API keys  |  🔁 Smart retry      |  📊 Real quality score")
    console.print("  📋 Professional TOC   |  📊 Auto HTML report      |  🛡️  DB integrity check")
    console.rule()

    # ── التحقق من صحة المسارات ───────────────────────────────────────────
    is_valid, validation_message = validate_input_paths(input_path, output_dir)
    if not is_valid:
        console.print(f"\n[bold red]❌ Path Error:[/bold red] {validation_message}")
        sys.exit(1)
    console.print(f"[green]✅ {validation_message}[/green]")
    console.print(f"   [dim]Input :[/dim] {input_path}")
    console.print(f"   [dim]Output:[/dim] {output_dir}")

    # ── إنشاء النظام ──────────────────────────────────────────────────────
    system = MasterTranslationSystem(extra_keys)

    # ── اختبار مفاتيح API (ما لم يُطلب التخطي) ───────────────────────────
    if not skip_api_test:
        try:
            await system.test_all_api_keys()
        except RuntimeError as e:
            console.print(f"\n[bold red]❌ API Key Error:[/bold red] {e}")
            sys.exit(1)
    else:
        console.print("[yellow]⚠️  API key test skipped (--skip-api-test)[/yellow]")

    # ── تشغيل النظام الكامل ───────────────────────────────────────────────
    try:
        console.print("\n[bold]🔄 Starting enhanced comprehensive translation process...[/bold]")

        final_document = await system.process_complete_book(
            input_path, output_dir, book_title_in, author_in
        )

        console.rule("[bold green]✅ Translation Complete[/bold green]")
        console.print(f"[green]📄 Final Novel:[/green] {final_document}")
        console.print(f"[green]📊 HTML Report:[/green] saved in {output_dir}")

        # ملخص التحسينات
        stats = system.translation_stats
        if stats['foreign_content_corrections'] > 0:
            console.print(
                f"[cyan]🔧 Applied {stats['foreign_content_corrections']} foreign content corrections[/cyan]"
            )
        if stats['contextual_adaptations'] > 0:
            console.print(
                f"[cyan]📖 Applied {stats['contextual_adaptations']} contextual adaptations[/cyan]"
            )
        if stats['smart_retries'] > 0:
            console.print(
                f"[yellow]🔁 Smart retries: {stats['smart_retries']} "
                f"(split:{stats['split_recoveries']}, "
                f"rephrase:{stats['rephrasing_recoveries']})[/yellow]"
            )
        console.print(
            f"[blue]📚 Saved {len(system.translation_engine.terminology_database)} terms to database[/blue]"
        )
        console.print(
            f"[magenta]⭐ Average quality score: "
            f"{stats['average_quality_score']:.2f}/10[/magenta]"
        )

    except KeyboardInterrupt:
        console.print("\n[yellow]⏹️ Process stopped by user. Progress has been saved.[/yellow]")

    except Exception as e:
        console.print(f"\n[bold red]❌ Unexpected fatal error:[/bold red] {e}")
        logger.error(f"Error in main: {e}")
        logger.error(traceback.format_exc())
        sys.exit(1)

    finally:
        # تنظيف موارد الشبكة — محمي في حال فشل إنشاء النظام
        try:
            if 'system' in dir() and hasattr(system, 'api_manager'):
                await system.api_manager.cleanup()
                logger.info("Network resources cleaned up successfully.")
        except Exception:
            pass


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n👋 Program terminated")
    except Exception as e:
        print(f"Unexpected top-level error: {str(e)}")
